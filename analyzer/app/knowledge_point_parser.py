from __future__ import annotations

import copy
import json
import logging
import os
import re
import shutil
from dataclasses import dataclass, field, replace as dc_replace
from pathlib import Path
from statistics import median
from typing import Any, Callable, Dict, List, Optional, Sequence, Set, Tuple


import fitz
from docx import Document
from sqlalchemy import func
from sqlalchemy.orm import Session

from analyzer.app.llm_client import call_llm, supports_vision_model
from shared.llm_step_config import resolve_step_llm_config, sync_llm_step_configs
from shared.prompt_step_config import resolve_step_prompt, sync_prompt_step_configs

from shared import models

from .config import (
    KNOWLEDGE_BLOCK_LLM_MAX_IMAGE_BYTES,
    KNOWLEDGE_BLOCK_LLM_MAX_IMAGES_PER_CALL,
    KNOWLEDGE_PACKAGE_KEYWORDS,
    KNOWLEDGE_POINTS_DIR,
    KNOWLEDGE_RAG_ENABLED,
    KNOWLEDGE_TOPIC_BLOCK_POINTS_MULTIMODAL,
    QUESTION_BANK_ASSET_DIR,
)
from .knowledge_point_provenance import (
    SOURCE_KIND_KNOWLEDGE_BLOCK,
    record_knowledge_point_provenance,
)
from .knowledge_block_multimodal import (
    append_topic_batch_multimodal_image_parts,
    openai_user_content_for_call,
    redact_openai_messages_for_audit,
    summarize_multimodal_content_for_log,
)


logger = logging.getLogger(__name__)

TOPIC_DOCX_BLOCK_POINTS_STEP_KEY = "analyzer.topic_docx_block_points"
TOPIC_DOCX_QUESTION_BRIDGE_STEP_KEY = "analyzer.topic_docx_question_bridge"
_TOPIC_BRIDGE_LLM_POOL_CAP = 72
_TOPIC_BRIDGE_QUERY_MAX = 1200


def _topic_bridge_llm_max_links() -> int:
    """按题 LLM 桥接最多采纳条数（模型可返回 0～本上限）。"""
    try:
        v = int(os.environ.get("KNOWLEDGE_POINT_BRIDGE_LLM_MAX_LINKS", "5"))
    except (TypeError, ValueError):
        v = 5
    return max(0, min(v, 8))


def _topic_bridge_llm_min_relevance() -> float:
    try:
        v = float(os.environ.get("KNOWLEDGE_POINT_BRIDGE_LLM_MIN_RELEVANCE", "0.38"))
    except (TypeError, ValueError):
        v = 0.38
    return max(0.0, min(v, 0.95))


def _topic_bridge_llm_strong_relevance_threshold() -> float:
    try:
        v = float(os.environ.get("KNOWLEDGE_POINT_BRIDGE_LLM_STRONG_RELEVANCE", "0.78"))
    except (TypeError, ValueError):
        v = 0.78
    return max(0.55, min(v, 0.99))
# 单次写入 run 目录下单个文本文件的字节上限（避免极端大响应撑爆磁盘）
_INGEST_VERBOSE_MAX_FILE_BYTES = int(os.environ.get("KNOWLEDGE_INGEST_VERBOSE_MAX_FILE_BYTES", str(3 * 1024 * 1024)))
_TOPIC_DOCX_LLM_MAX_BLOCK_TEXT = 4000
_TOPIC_DOCX_LLM_BATCH_JSON_CHARS = 20000
_TOPIC_DOCX_LLM_MAX_NAMES_PER_BLOCK = 12
_TOPIC_DOCX_LLM_WEIGHT = 0.6
_TOPIC_DOCX_LLM_CONFIDENCE = 0.72


def _env_flag_true(key: str) -> bool:
    return os.environ.get(key, "").strip().lower() in ("1", "true", "yes", "on")


def _topic_docx_point_mode() -> str:
    """
    专题 DOCX 知识点来源（互斥/组合开关）：
    - regex：仅正则/标题抽取链路（不写 LLM 包↔点）。
    - llm：仅大模型链路（不写正则产生的 KnowledgePackagePoint；块主知识点在 LLM 写入后对齐到首条 LLM 点）。
    - both：正则/标题与块级 LLM 两条链路都执行（需显式设置，不再作为默认）。

    环境变量：KNOWLEDGE_POINT_DOCX_POINT_MODE ∈ {regex, llm, both}；
    未设置时：若 KNOWLEDGE_POINT_LLM_EXTRACT_ENABLED=true 则默认为 llm（纯 LLM），否则为 regex。
    """
    raw = (os.environ.get("KNOWLEDGE_POINT_DOCX_POINT_MODE") or "").strip().lower()
    if raw in ("regex", "llm", "both"):
        return raw
    if _env_flag_true("KNOWLEDGE_POINT_LLM_EXTRACT_ENABLED"):
        return "llm"
    return "regex"


def _docx_formula_image_layout_stats(blocks: List[Dict[str, Any]]) -> Tuple[int, int]:
    """统计 extract_docx_blocks 结果里公式栅格图数量及是否带上 Word 占位尺寸（写入 run.log 用，不依赖 logging INFO）。"""
    formula_images = 0
    with_layout = 0

    def walk(node: Any) -> None:
        nonlocal formula_images, with_layout
        if isinstance(node, dict):
            if node.get("type") == "image" and node.get("omml_raster"):
                formula_images += 1
                lw = node.get("layout_width_px")
                lh = node.get("layout_height_px")
                if isinstance(lw, (int, float)) and isinstance(lh, (int, float)) and lw > 0 and lh > 0:
                    with_layout += 1
            for value in node.values():
                walk(value)
        elif isinstance(node, list):
            for item in node:
                walk(item)

    for block in blocks:
        walk(block.get("render"))
    return formula_images, with_layout


def delete_knowledge_ingest_run_folders_for_source_document(source_document_id: int) -> tuple[list[str], list[str]]:
    """
    删除知识点 WebSocket 摄入在 _ingest_runs 下留下的运行目录（通过 ingestion_detail.json 中的 source_document_id 匹配）。
    在 SourceDocument 已从数据库删除成功后调用，避免误删进行中的任务目录。
    """
    removed: list[str] = []
    failed: list[str] = []
    root = Path(KNOWLEDGE_POINTS_DIR) / "_ingest_runs"
    if not root.is_dir():
        return removed, failed
    sid = int(source_document_id)
    for child in sorted(root.iterdir(), key=lambda p: p.name, reverse=True):
        if not child.is_dir():
            continue
        detail_path = child / "ingestion_detail.json"
        if not detail_path.is_file():
            continue
        try:
            payload = json.loads(detail_path.read_text(encoding="utf-8"))
        except Exception:
            continue
        matches = False
        for item in payload.get("processed") or []:
            if not isinstance(item, dict):
                continue
            try:
                if int(item.get("source_document_id") or 0) == sid:
                    matches = True
                    break
            except (TypeError, ValueError):
                continue
        if not matches:
            continue
        try:
            shutil.rmtree(child)
            removed.append(str(child))
        except Exception as exc:
            logger.warning("Failed to remove knowledge ingest run dir %s: %s", child, exc)
            failed.append(str(child))
    return removed, failed


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
# 不含「行首单独数字」（易与页码 2 / 397 混淆）与「裸 专题 01」（多为目录行）
HEADING_PATTERN = re.compile(
    r"^(?:第[一二三四五六七八九十百零〇0-9]+(?:章|节|讲|专题)|[一二三四五六七八九十百零〇]+[、.．])"
)
_BARE_TOPIC_INDEX_LINE = re.compile(r"^专题\s*\d+\s*$")
_PAGE_FRACTION_FOOTER = re.compile(r"^\d{1,4}\s*[/／]\s*\d{1,4}\s*$")
_PAGE_DASH_RANGE = re.compile(r"^\d{1,4}\s*[-–—]\s*\d{1,4}\s*$")
# 题号 + 卷种标签，整段标题不应登记为独立「知识点」名称
_EXAM_PAPER_TAG_IN_NAME = re.compile(r"【\s*(?:高考|新课标|\d{4}\s*年)")
SECTION_ROLE_PATTERNS: Sequence[Tuple[str, str]] = (
    ("summary", r"总结|归纳|结论|小结"),
    ("exam_focus", r"考向|高频|考试|考查"),
    ("pitfall", r"易错|误区|陷阱"),
    ("example_bridge", r"例题|真题|变式|训练"),
    ("definition", r"定义|概念|性质"),
    ("explainer", r"方法|技巧|思路|讲解|解析"),
)
ATOM_TYPE_BY_BLOCK_ROLE = {
    "summary": "conclusion",
    "exam_focus": "exam_pattern",
    "pitfall": "pitfall",
    "example_bridge": "method",
    "definition": "definition",
    "explainer": "method",
}


@dataclass
class ParsedSection:
    heading: str
    text: str
    page_start: int
    page_end: int


@dataclass
class TopicMarker:
    title: str
    start_page: int
    source_page: int


@dataclass
class ParsedPackage:
    title: str
    sections: List[ParsedSection]


@dataclass
class LayoutLine:
    page_no: int
    text: str
    x0: float
    y0: float
    y_center: float
    font_size: float


# ---------------------------------------------------------------------------
# DOCX 专题段落分类器（状态机）
# ---------------------------------------------------------------------------

@dataclass
class TopicContentSegment:
    """一段连续的专题知识内容（非题目）。"""
    section_title: str
    blocks: List[Dict[str, Any]] = field(default_factory=list)
    plain_text: str = ""
    block_index_start: int = 0
    block_index_end: int = 0
    adjacent_question_nos: List[str] = field(default_factory=list)


@dataclass
class TopicQuestionSegment:
    """一道或一组题目（含题干+选项+答案+解析）。"""
    question_no: str
    section_title: str
    blocks: List[Dict[str, Any]] = field(default_factory=list)
    block_index_start: int = 0
    block_index_end: int = 0


_QUESTION_HEAD_RE = re.compile(
    r"^\s*(?:"
    r"第\s*(\d{1,3})\s*题|"
    r"(\d{1,3})[\.．、](?!\d)|"
    r"(\d{1,3})\s*(?:\(|（)"
    r")\s*"
)
_QUESTION_BODY_HINT_RE = re.compile(
    r"[（(]\s*[　 ]*\s*[）)]|"                       # (　　) 选择留空
    r"[（(]\s*(?:多选|单选|判断|填空)\s*[）)]|"         # (多选) 等题型括号
    r"(?:已知|若集合|设集合|等于)|"
    r"[A-HＡ-Ｈ][\.．、]|"                            # 选项
    r"(?:答案|解析)|"
    r"_{2,}|"                                         # 填空线
    r"=\s*[（(]|"                                     # =( 选择
    r"\d{4}\s*[·•年]"                                 # 年份标记（真题来源）
)
_OPTION_RE = re.compile(r"^\s*[A-HＡ-Ｈ][\.．、:：]\s*")
# Word 自动编号里的「2.」常不在 runs 中，extract 后首段可能直接以卷种括号开头。
_LOOSE_TOPIC_QUESTION_HEAD_RE = re.compile(
    r"^\s*[（(]\s*\d{4}\s*[·•・年][^。\n]{0,70}?(?:设集合|若集合|已知)"
)
_SUB_QUESTION_HEAD_RE = re.compile(
    r"^\s*\((\d{1,2})\)[\.．、]?\s*"
)
_ANSWER_LABEL_RE = re.compile(
    r"^\s*(?:【|\[|\()?\s*(?:参考答案|答案解析|答案|解析|详解|解答|思路导引|思路引导|解法|过程|点评|评注|点拨|点睛)"
    r"\s*(?:】|\]|\))?\s*[:：]?\s*"
)
_KNOWLEDGE_TAG_RE = re.compile(
    r"^\s*(?:【课程标准】|【常用结论】|【学法指导】|【重点提示】|\[微提醒\]|\[真题再现\]|\[教材呈现\]|\[变式探究\])"
)
# 题头识别：行首「[真题再现]」等不是知识块标签，应剥掉再判题号（否则 1. 永远不在 ^）
_LEADING_BRACKET_TOPIC_STRIP_RE = re.compile(
    r"^\s*(?:\[真题再现\]|\[教材呈现\]|\[变式探究\])\s*",
)
_METHOD_SUMMARY_KEYWORDS = re.compile(
    r"方法|技巧|注意|关键|规律|策略|易错|总结|归纳|结论|常用|重点|要点|步骤"
)

_SECTION_TABLE_KEYWORDS_RE = re.compile(
    r"考点[一二三四五六七八九十\d]|自主检测|自主练透|师生共研|课时测评|真题再现|教材呈现|学生用书"
)


def _is_section_marker_table(block: Dict[str, Any]) -> bool:
    render = block.get("render")
    if not isinstance(render, dict) or render.get("type") != "table":
        return False
    rows = render.get("rows") or []
    if len(rows) != 1:
        return False
    cells = (rows[0].get("cells") or [])
    if len(cells) != 1:
        return False
    text = (block.get("text") or "").strip()
    return bool(_SECTION_TABLE_KEYWORDS_RE.search(text))


def _is_knowledge_summary_table(block: Dict[str, Any]) -> bool:
    render = block.get("render")
    if not isinstance(render, dict) or render.get("type") != "table":
        return False
    rows = render.get("rows") or []
    if len(rows) < 2:
        return False
    first_row_cells = (rows[0].get("cells") or [])
    return len(first_row_cells) >= 2


def _extract_section_table_title(block: Dict[str, Any]) -> str:
    text = (block.get("text") or "").strip()
    text = re.sub(r"学生用书[⬇↓]?第?\d*页?", "", text).strip()
    return text[:80] if text else "未命名区域"


def _block_text(block: Dict[str, Any]) -> str:
    return (block.get("text") or "").strip()


def _render_node_plain_text(node: Any) -> str:
    if not isinstance(node, dict):
        return ""
    t = node.get("type", "")
    if t == "text":
        return str(node.get("text") or "")
    if t == "formula":
        return str(node.get("text") or "")
    if t == "image":
        return str(node.get("alt_text") or "[图片]")
    if t == "line_break":
        return "\n"
    parts = []
    for v in node.values():
        if isinstance(v, list):
            for item in v:
                parts.append(_render_node_plain_text(item))
        elif isinstance(v, dict):
            parts.append(_render_node_plain_text(v))
    return "".join(parts)


_LOOSE_STEM_BLANK_RE = re.compile(
    r"[（(]\s*[　\s]{2,10}\s*[）)]|（\s*(?:多选|单选|填空)\s*）|\(\s*(?:多选|单选|填空)\s*\)|_{3,}",
)


def _same_block_suggests_question_stem(stem_text: str) -> bool:
    """题干内常见「题」信号：选择留空、填空线、证明/解答起句等。"""
    if not stem_text:
        return False
    if _LOOSE_STEM_BLANK_RE.search(stem_text):
        return True
    if re.search(r"(?:证明|求证)[:：]", stem_text[:160]):
        return True
    return False


def _lookahead_suggests_question_tail(
    blocks: List[Dict[str, Any]],
    idx: int,
    stem_text: str,
    window: int = 28,
) -> bool:
    """非选择题可能没有 A/B；用答案/解析区、证明起段、或题干留空等佐证独立成题。"""
    if _same_block_suggests_question_stem(stem_text):
        return True
    hi = min(idx + window, len(blocks))
    for j in range(idx + 1, hi):
        jt = _block_text(blocks[j])
        if _OPTION_RE.match(jt):
            return True
        if _ANSWER_LABEL_RE.match(jt):
            return True
        render = blocks[j].get("render")
        if isinstance(render, dict) and render.get("type") == "table" and jt:
            if re.search(r"[A-HＡ-Ｈ][\.．、:：]", jt):
                return True
        if re.match(r"^\s*(?:证明|求证|解[:：]|解答[:：])", jt):
            return True
    return False


def _looks_like_real_question(
    block_text: str,
    q_match: re.Match,
    blocks: List[Dict[str, Any]],
    idx: int,
) -> bool:
    """区分「真题号」（如 "1.(多选)下列..."）和「知识小标题」（如 "1.集合与元素"）。"""
    after_num = block_text[q_match.end():]
    if _QUESTION_BODY_HINT_RE.search(after_num):
        return True
    lookahead = min(idx + 5, len(blocks))
    for j in range(idx + 1, lookahead):
        jt = _block_text(blocks[j])
        if _QUESTION_HEAD_RE.match(jt):
            break
        if _OPTION_RE.match(jt):
            return True
        if _ANSWER_LABEL_RE.match(jt):
            return True
    return False


def _expand_block_group_paragraphs(
    blocks: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """Expand block_group blocks that contain embedded question starts into
    individual paragraph-level blocks so the FSM can detect boundaries.

    Only expands when a NON-FIRST sub-paragraph passes both the question
    head regex AND the body-hint regex (i.e. looks like a real exam question,
    not a knowledge heading like '1.集合与元素').
    """
    expanded: List[Dict[str, Any]] = []
    for block in blocks:
        render = block.get("render")
        if not isinstance(render, dict) or render.get("type") != "block_group":
            expanded.append(block)
            continue
        sub_blocks = render.get("blocks") or []
        if len(sub_blocks) <= 1:
            expanded.append(block)
            continue
        has_embedded_q = False
        for i, sb in enumerate(sub_blocks):
            if i == 0:
                continue
            sb_text = _render_node_plain_text(sb).strip()
            if not sb_text:
                continue
            head_m = _QUESTION_HEAD_RE.match(sb_text)
            if not head_m:
                continue
            after_num = sb_text[head_m.end():]
            if _QUESTION_BODY_HINT_RE.search(after_num):
                has_embedded_q = True
                break
        if not has_embedded_q:
            expanded.append(block)
            continue
        for sb in sub_blocks:
            sb_text = _render_node_plain_text(sb).strip()
            expanded.append({"text": sb_text, "render": sb})
    return expanded


def _gather_embedded_question_split_positions(t: str) -> List[int]:
    """句号/换行之后、下一题题号之前的切分点（不含于前一段）。"""
    poses: set[int] = set()
    pat_sentence = re.compile(
        r"(?<=[。！？；….\uff0e])(?=\s*(?:第\s*\d{1,3}\s*题|\d{1,3}[\.．、](?!\d)))",
    )
    pat_newline = re.compile(
        r"(?<=[\n\r\u2029])(?=\s*(?:第\s*\d{1,3}\s*题|\d{1,3}[\.．、](?!\d)))",
    )
    for m in pat_sentence.finditer(t):
        if m.start() > 0:
            poses.add(m.start())
    for m in pat_newline.finditer(t):
        if m.start() > 0:
            poses.add(m.start())
    return sorted(poses)


def _piece_looks_like_question_opening(px: str) -> bool:
    s = (px or "").strip()
    if not s:
        return False
    if _QUESTION_HEAD_RE.match(s):
        return True
    if _LOOSE_TOPIC_QUESTION_HEAD_RE.match(s):
        return True
    return False


def _expand_embedded_question_after_sentence(
    blocks: List[Dict[str, Any]],
) -> Tuple[List[Dict[str, Any]], int]:
    """
    同一 Word 段落内「解析末句 + 下一题题号」合并时，题号不在块首，专题 FSM 认不出新题。
    在句号/换行后按题头模式切开，并同步截取 paragraph render。
    """
    from analyzer.app.question_bank_parser import QuestionBankIngestionService

    qbs = QuestionBankIngestionService()
    out: List[Dict[str, Any]] = []
    split_paragraph_count = 0
    for block in blocks:
        t = block.get("text") or ""
        render = block.get("render")
        if (
            not isinstance(render, dict)
            or render.get("type") != "paragraph"
            or len(t) < 36
        ):
            out.append(block)
            continue
        rplain = qbs._structured_block_text_from_render(copy.deepcopy(render))
        if abs(len(qbs._normalize_text(rplain)) - len(qbs._normalize_text(t))) > 5:
            out.append(block)
            continue
        cut_points = _gather_embedded_question_split_positions(t)
        if not cut_points:
            out.append(block)
            continue
        bounds: List[Tuple[int, int]] = []
        prev = 0
        for cp in cut_points:
            if cp > prev:
                bounds.append((prev, cp))
            prev = cp
        if prev < len(t):
            bounds.append((prev, len(t)))
        non_empty_bounds = [(a, b) for a, b in bounds if t[a:b].strip()]
        if len(non_empty_bounds) < 2:
            out.append(block)
            continue
        if not any(_piece_looks_like_question_opening(t[a:b]) for a, b in non_empty_bounds[1:]):
            out.append(block)
            continue
        split_paragraph_count += 1
        for a, b in bounds:
            if not t[a:b].strip():
                continue
            nb = copy.deepcopy(block)
            nb["render"] = qbs.slice_paragraph_render_range(render, a, b)
            synced = (qbs._structured_block_text_from_render(nb["render"]) or "").strip()
            nb["text"] = synced if synced else t[a:b].strip()
            out.append(nb)
    if split_paragraph_count:
        logger.info(
            "topic_docx_embedded_q_split: 合并段按题号切开 count=%s",
            split_paragraph_count,
        )
    return out, split_paragraph_count


_SUB_Q_NO_RE = re.compile(r"^\((\d+)\)$")


def _merge_sub_questions_with_shared_answers(
    question_segments: List[TopicQuestionSegment],
) -> List[TopicQuestionSegment]:
    """合并共享答案区的连续子题 ((1)(2)(3))。"""
    if len(question_segments) < 2:
        return question_segments

    merged: List[TopicQuestionSegment] = []
    skip: set[int] = set()
    i = 0
    while i < len(question_segments):
        seg = question_segments[i]
        m = _SUB_Q_NO_RE.match(seg.question_no)
        if not m:
            merged.append(seg)
            i += 1
            continue

        group: List[TopicQuestionSegment] = [seg]
        group_indices = [i]
        expected = int(m.group(1)) + 1
        j = i + 1
        while j < len(question_segments):
            m2 = _SUB_Q_NO_RE.match(question_segments[j].question_no)
            if m2 and int(m2.group(1)) == expected:
                group.append(question_segments[j])
                group_indices.append(j)
                expected += 1
                j += 1
            else:
                break

        if len(group) < 2:
            merged.append(seg)
            i += 1
            continue

        last_text = "\n".join(_block_text(b) for b in group[-1].blocks)
        has_shared = bool(re.search(r"答案[：:].*\(1\)", last_text))
        if not has_shared:
            merged.extend(group)
            i = j
            continue

        all_blocks: List[Dict[str, Any]] = []
        for g in group:
            all_blocks.extend(g.blocks)
        merged.append(
            TopicQuestionSegment(
                question_no=group[0].question_no,
                section_title=group[0].section_title,
                blocks=all_blocks,
                block_index_start=group[0].block_index_start,
                block_index_end=group[-1].block_index_end,
            )
        )
        for idx in group_indices:
            skip.add(idx)
        i = j

    return merged


def _filter_answerless_questions(
    question_segments: List[TopicQuestionSegment],
) -> List[TopicQuestionSegment]:
    """过滤无答案的题目。"""
    kept: List[TopicQuestionSegment] = []
    for seg in question_segments:
        has_answer = False
        for b in seg.blocks:
            text = _block_text(b)
            if _ANSWER_LABEL_RE.match(text):
                has_answer = True
                break
            if len(seg.blocks) > 1 and re.match(r"^\s*证明[:：]", text) and len(text) > 30:
                has_answer = True
                break
        if has_answer:
            kept.append(seg)
    return kept


def classify_topic_docx_blocks(
    blocks: List[Dict[str, Any]],
) -> Tuple[List[TopicContentSegment], List[TopicQuestionSegment], int]:
    """
    将 extract_docx_blocks 输出的 block 列表划分为【专题内容段】和【题目段】。

    采用有限状态机：CONTENT / QUESTION / ANSWER_ZONE 三态。
    """

    blocks = _expand_block_group_paragraphs(blocks)
    blocks, embedded_q_split = _expand_embedded_question_after_sentence(blocks)

    content_segments: List[TopicContentSegment] = []
    question_segments: List[TopicQuestionSegment] = []

    state = "CONTENT"
    current_section_title = "专题正文"

    cur_content_blocks: List[Dict[str, Any]] = []
    cur_content_start = 0

    cur_q_no: Optional[str] = None
    cur_q_blocks: List[Dict[str, Any]] = []
    cur_q_start = 0
    plain_after_answer = 0

    last_question_no_int: Optional[int] = None

    def _flush_content(end_idx: int) -> None:
        nonlocal cur_content_blocks, cur_content_start
        if not cur_content_blocks:
            return
        text_parts = [_block_text(b) for b in cur_content_blocks]
        plain = "\n".join(p for p in text_parts if p)
        if plain.strip():
            seg = TopicContentSegment(
                section_title=current_section_title,
                blocks=cur_content_blocks,
                plain_text=plain,
                block_index_start=cur_content_start,
                block_index_end=end_idx,
            )
            content_segments.append(seg)
        cur_content_blocks = []
        cur_content_start = end_idx

    def _flush_question(end_idx: int) -> None:
        nonlocal cur_q_no, cur_q_blocks, cur_q_start, last_question_no_int, plain_after_answer
        if cur_q_no and cur_q_blocks:
            seg = TopicQuestionSegment(
                question_no=cur_q_no,
                section_title=current_section_title,
                blocks=cur_q_blocks,
                block_index_start=cur_q_start,
                block_index_end=end_idx,
            )
            question_segments.append(seg)
            if content_segments:
                content_segments[-1].adjacent_question_nos.append(cur_q_no)
        cur_q_no = None
        cur_q_blocks = []
        cur_q_start = end_idx
        plain_after_answer = 0

    def _extract_q_no(match: re.Match) -> Tuple[str, Optional[int]]:
        raw = match.group(1) or match.group(2) or match.group(3) or "?"
        try:
            return raw, int(raw)
        except (ValueError, TypeError):
            return raw, None

    def _is_reasonable_transition(
        prev_int: Optional[int],
        cur_int: Optional[int],
        *,
        prev_fsm_state: str,
    ) -> bool:
        """
        题号递增（含从 1 重开）视为合理。
        另：教辅多段「练习」各自从 1、2 编号；上一题已离开 QUESTION 态后，可出现与上一道相同的题号
        （如 [真题再现] 内连续两道「2.」），此时必须仍切分为新题。
        """
        if prev_int is None or cur_int is None:
            return True
        if 1 <= cur_int - prev_int <= 5 or cur_int == 1:
            return True
        if cur_int <= prev_int and cur_int <= 5 and (prev_int - cur_int) >= 4:
            return True
        if cur_int == prev_int and prev_fsm_state != "QUESTION":
            return True
        return False

    for idx, block in enumerate(blocks):
        text = _block_text(block)
        render = block.get("render")
        is_table = isinstance(render, dict) and render.get("type") == "table"

        if _is_section_marker_table(block):
            if state == "QUESTION" or state == "ANSWER_ZONE":
                _flush_question(idx)
            elif state == "CONTENT":
                _flush_content(idx)
            current_section_title = _extract_section_table_title(block)
            state = "CONTENT"
            cur_content_start = idx + 1
            continue

        clean_text = re.sub(r"^(?:\[图片\]|公式图片|公式|formula|\[IMG\]|【图片】)+\s*", "", text) if text else text
        clean_text = _LEADING_BRACKET_TOPIC_STRIP_RE.sub("", clean_text or "")
        raw_q_match = _QUESTION_HEAD_RE.match(clean_text) if clean_text else None
        std_head = bool(
            raw_q_match and _looks_like_real_question(clean_text, raw_q_match, blocks, idx),
        )
        q_match = raw_q_match if std_head else None
        q_no_override: Optional[Tuple[str, Optional[int]]] = None
        if (
            not std_head
            and clean_text
            and _LOOSE_TOPIC_QUESTION_HEAD_RE.match(clean_text)
            and _lookahead_suggests_question_tail(blocks, idx, clean_text)
        ):
            nxt = (last_question_no_int or 0) + 1
            q_no_override = (str(nxt), nxt)
        has_topic_q_head = bool(q_match) or bool(q_no_override)
        is_option = bool(_OPTION_RE.match(text)) if text else False
        is_answer = bool(_ANSWER_LABEL_RE.match(text)) if text else False
        is_ktag = bool(_KNOWLEDGE_TAG_RE.match(text)) if text else False
        is_ksummary_table = _is_knowledge_summary_table(block)
        is_method = (
            not has_topic_q_head
            and not is_option
            and not is_answer
            and not is_table
            and len(text) > 30
            and bool(_METHOD_SUMMARY_KEYWORDS.search(text))
            and not text[:6].strip().startswith(tuple("0123456789"))
        )

        if state == "CONTENT":
            if has_topic_q_head:
                q_str, q_int = (
                    _extract_q_no(q_match) if q_match else (q_no_override[0], q_no_override[1])
                )
                if _is_reasonable_transition(last_question_no_int, q_int, prev_fsm_state=state):
                    _flush_content(idx)
                    state = "QUESTION"
                    cur_q_no = q_str
                    cur_q_blocks = [block]
                    cur_q_start = idx
                    last_question_no_int = q_int
                    plain_after_answer = 0
                    continue
            sub_q_match = _SUB_QUESTION_HEAD_RE.match(clean_text) if clean_text and not has_topic_q_head else None
            if sub_q_match and _lookahead_suggests_question_tail(blocks, idx, clean_text):
                q_str = sub_q_match.group(0).strip()
                try:
                    q_int = int(sub_q_match.group(1))
                except ValueError:
                    q_int = None
                _flush_content(idx)
                state = "QUESTION"
                cur_q_no = q_str
                cur_q_blocks = [block]
                cur_q_start = idx
                last_question_no_int = q_int
                plain_after_answer = 0
                continue
            if not has_topic_q_head and text and _LOOSE_STEM_BLANK_RE.search(text) and _lookahead_suggests_question_tail(blocks, idx, text):
                _flush_content(idx)
                state = "QUESTION"
                cur_q_no = f"Q_bare_{idx}"
                cur_q_blocks = [block]
                cur_q_start = idx
                last_question_no_int = None
                plain_after_answer = 0
                continue
            cur_content_blocks.append(block)

        elif state == "QUESTION":
            if has_topic_q_head:
                q_str, q_int = (
                    _extract_q_no(q_match) if q_match else (q_no_override[0], q_no_override[1])
                )
                if _is_reasonable_transition(last_question_no_int, q_int, prev_fsm_state=state):
                    _flush_question(idx)
                    cur_q_no = q_str
                    cur_q_blocks = [block]
                    cur_q_start = idx
                    last_question_no_int = q_int
                    plain_after_answer = 0
                    continue
            if is_answer:
                cur_q_blocks.append(block)
                state = "ANSWER_ZONE"
                plain_after_answer = 0
                continue
            if is_ktag or is_ksummary_table or is_method:
                _flush_question(idx)
                state = "CONTENT"
                cur_content_blocks = [block]
                cur_content_start = idx
                continue
            cur_q_blocks.append(block)

        elif state == "ANSWER_ZONE":
            if has_topic_q_head:
                q_str, q_int = (
                    _extract_q_no(q_match) if q_match else (q_no_override[0], q_no_override[1])
                )
                if _is_reasonable_transition(last_question_no_int, q_int, prev_fsm_state=state):
                    _flush_question(idx)
                    state = "QUESTION"
                    cur_q_no = q_str
                    cur_q_blocks = [block]
                    cur_q_start = idx
                    last_question_no_int = q_int
                    plain_after_answer = 0
                    continue
            if is_answer:
                cur_q_blocks.append(block)
                plain_after_answer = 0
                continue
            if is_ktag or is_ksummary_table:
                _flush_question(idx)
                state = "CONTENT"
                cur_content_blocks = [block]
                cur_content_start = idx
                continue
            if is_method:
                plain_after_answer += 1
                if plain_after_answer >= 2:
                    _flush_question(idx - 1)
                    state = "CONTENT"
                    cur_content_blocks = [block]
                    cur_content_start = idx
                    continue
            cur_q_blocks.append(block)
            if not is_option and not is_answer and text:
                plain_after_answer += 1

    total = len(blocks)
    if state == "QUESTION" or state == "ANSWER_ZONE":
        _flush_question(total)
    elif state == "CONTENT":
        _flush_content(total)

    for qs in question_segments:
        for cs in content_segments:
            if abs(cs.block_index_end - qs.block_index_start) <= 2:
                if qs.question_no not in cs.adjacent_question_nos:
                    cs.adjacent_question_nos.append(qs.question_no)
            if abs(qs.block_index_end - cs.block_index_start) <= 2:
                if qs.question_no not in cs.adjacent_question_nos:
                    cs.adjacent_question_nos.append(qs.question_no)

    question_segments = _merge_sub_questions_with_shared_answers(question_segments)
    question_segments = _filter_answerless_questions(question_segments)

    return content_segments, question_segments, embedded_q_split


class KnowledgePointIngestionService:


    """知识点专题资料摄入服务（独立于题目主链路）。"""

    def __init__(self, knowledge_points_dir: Optional[str] = None) -> None:
        self.knowledge_points_dir = Path(knowledge_points_dir or KNOWLEDGE_POINTS_DIR)
        self.knowledge_points_dir.mkdir(parents=True, exist_ok=True)
        self.package_keywords = tuple(keyword.strip() for keyword in KNOWLEDGE_PACKAGE_KEYWORDS if keyword and keyword.strip())
        # WebSocket/管理端摄入时由 ingest_files_from_knowledge_points_dir 注入，用于落盘详细审计
        self._ingest_run_log_dir: Optional[Path] = None

    def _ingest_verbose_enabled(self) -> bool:
        return getattr(self, "_ingest_run_log_dir", None) is not None

    def _ingest_verbose_write(self, relative_path: str, text: str, *, max_bytes: Optional[int] = None) -> None:
        """将长文本写入本次摄入 run 目录（相对路径如 llm/foo.txt）。未设置 _ingest_run_log_dir 时 no-op。"""
        root = getattr(self, "_ingest_run_log_dir", None)
        if root is None:
            return
        cap = max_bytes if max_bytes is not None else _INGEST_VERBOSE_MAX_FILE_BYTES
        raw = text if isinstance(text, str) else str(text)
        encoded = raw.encode("utf-8")
        if len(encoded) > cap:
            raw = encoded[:cap].decode("utf-8", errors="ignore") + f"\n\n...[truncated, original {len(encoded)} bytes, cap={cap}]"
        path = root / relative_path
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(raw, encoding="utf-8")

    @staticmethod
    def _notify_progress(callback: Optional[Callable[[str], None]], message: str) -> None:
        if not callback:
            return
        try:
            callback(message)
        except Exception:
            pass

    def _export_pdf_embedded_images(self, source_path: Path, dest_dir: Path, max_images: int = 48) -> int:
        """将 PDF 内嵌位图导出到本次运行 assets 子目录（便于对照解析问题）。"""
        dest_dir.mkdir(parents=True, exist_ok=True)
        count = 0
        try:
            with fitz.open(source_path) as doc:
                for page_index in range(len(doc)):
                    page = doc[page_index]
                    for img in page.get_images(full=True):
                        if count >= max_images:
                            return count
                        xref = img[0]
                        base = doc.extract_image(xref)
                        ext = (base.get("ext") or "png").lower()
                        if ext not in ("png", "jpeg", "jpg", "jpx", "jp2", "webp"):
                            ext = "png"
                        name = f"page{page_index + 1}_xref{xref}.{ext}"
                        (dest_dir / name).write_bytes(base["image"])
                        count += 1
        except Exception:
            return count
        return count

    def ingest_source_document(
        self,
        db: Session,
        source_document_id: int,
        force_reingest: bool = False,
        progress_callback: Optional[Callable[[str], None]] = None,
    ) -> Dict[str, object]:
        source_document = (
            db.query(models.SourceDocument)
            .filter(models.SourceDocument.id == source_document_id)
            .first()
        )
        if not source_document:
            raise ValueError(f"SourceDocument {source_document_id} 不存在")

        self._notify_progress(
            progress_callback,
            f"文档就绪：id={source_document_id} file={source_document.file_name} force_reingest={force_reingest}",
        )

        source_path = Path(source_document.storage_url)
        if not source_path.exists() or not source_path.is_file():
            raise ValueError(f"文件不存在：{source_path}")

        extension = source_path.suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的知识点资料格式：{extension or 'unknown'}")

        # 与题库摄入保持一致：SourceDocument.parse_status 作为「内容管理/试卷列表」展示状态
        # - running: 正在解析
        # - success: 解析完成（无论产出是知识块还是题目）
        # - failed: 解析失败
        source_document.parse_status = "running"
        db.add(source_document)

        parse_job = self._start_parse_job(db, source_document_id)
        db.commit()
        db.refresh(parse_job)

        try:

            if force_reingest:
                self._notify_progress(progress_callback, "强制重摄入：清理该文档已有专题包与关联块…")
                self._clear_existing_package_artifacts(db, source_document_id)

            if extension == ".docx":
                result = self._ingest_docx_topic(
                    db, source_document, source_path, parse_job, fallback_point=None,
                    progress_callback=progress_callback,
                )
                # _ingest_docx_topic 已完成 parse_job，但也需要同步更新 SourceDocument.parse_status
                source_document.parse_status = "success"
                db.add(source_document)
                db.commit()
                return result

            self._notify_progress(progress_callback, f"抽取页面文本：{source_path.name}")
            pages = self._extract_pages(source_path)
            sections = self._split_into_sections(pages)
            parsed_packages = self._split_into_packages(source_document, source_path, pages, sections)


            fallback_point = self._get_or_create_knowledge_point(
                db,
                canonical_name="未归类知识点",
                subject=source_document.subject,
                grade_scope=source_document.grade,
                source_origin="model",
            )

            package_results: List[Dict[str, object]] = []
            total_block_count = 0
            total_atom_count = 0
            total_point_count = 0

            for parsed_package in parsed_packages:
                package_title = parsed_package.title
                package_sections = parsed_package.sections

                package = models.KnowledgePackage(
                    source_document_id=source_document.id,
                    tenant_id=source_document.tenant_id,
                    package_title=package_title,
                    package_type="topic",
                    subject=source_document.subject,
                    grade=source_document.grade,
                    page_range_json=self._build_page_range_from_sections(package_sections),
                    outline_json=self._build_outline(package_sections),
                    summary_text=self._build_summary_text(package_sections),
                    parse_status="running",
                    review_status="draft",
                )
                db.add(package)
                db.flush()

                package_point_links: Dict[int, models.KnowledgePackagePoint] = {}
                package_block_count = 0
                package_atom_count = 0
                topic_key = re.sub(r"\s+", "", package_title)

                for order, section in enumerate(package_sections, start=1):
                    point_names = [
                        name
                        for name in self._extract_knowledge_point_names(section.heading, section.text)
                        if re.sub(r"\s+", "", name or "") != topic_key
                    ]
                    points = [
                        self._get_or_create_knowledge_point(
                            db,
                            canonical_name=name,
                            subject=source_document.subject,
                            grade_scope=source_document.grade,
                            source_origin="model",
                        )
                        for name in point_names
                    ]
                    points = [point for point in points if re.sub(r"\s+", "", point.canonical_name or "") != topic_key]
                    if not points:
                        points = [fallback_point]

                    primary_point = points[0]
                    block_role = self._infer_block_role(section.heading)
                    normalized_text = self._normalize_text(section.text)
                    block = models.KnowledgeBlock(
                        package_id=package.id,
                        knowledge_point_id=primary_point.id,
                        parent_block_id=None,
                        block_order=order,
                        section_path=f"{package_title}/{section.heading}",
                        block_role=block_role,
                        content_format="plain_text",
                        raw_text=section.text,
                        normalized_text=normalized_text,
                        rich_content_json={"heading": section.heading},
                        source_page_no=section.page_start,
                        source_anchor_json={"page_start": section.page_start, "page_end": section.page_end},
                        source_origin="model",
                        confidence=0.8,
                        is_primary=(order == 1),
                    )
                    db.add(block)
                    db.flush()
                    package_block_count += 1

                    for point in points:
                        record_knowledge_point_provenance(
                            db,
                            knowledge_point_id=point.id,
                            source_kind=SOURCE_KIND_KNOWLEDGE_BLOCK,
                            source_id=block.id,
                            package_id=package.id,
                            origin_step="pdf_topic_section",
                        )

                    for index, point in enumerate(points):
                        if point.id not in package_point_links:
                            link = models.KnowledgePackagePoint(
                                package_id=package.id,
                                knowledge_point_id=point.id,
                                relation_type="core" if index == 0 else "supplement",
                                weight_score=1.0 if index == 0 else 0.7,
                                order_in_package=len(package_point_links) + 1,
                                source_origin="model",
                                confidence=0.8 if index == 0 else 0.7,
                                approved_status="pending",
                            )
                            db.add(link)
                            package_point_links[point.id] = link

                    sentences = self._split_sentences(normalized_text)
                    atom_type = ATOM_TYPE_BY_BLOCK_ROLE.get(block_role, "conclusion")
                    for sentence in sentences[:3]:
                        if len(sentence) < 10:
                            continue
                        atom = models.KnowledgeAtom(
                            knowledge_point_id=primary_point.id,
                            package_id=package.id,
                            atom_type=atom_type,
                            canonical_text=sentence,
                            normalized_json={"source": "section_sentence", "section_heading": section.heading},
                            evidence_block_id=block.id,
                            source_origin="model",
                            confidence=0.75,
                            review_status="draft",
                        )
                        db.add(atom)
                        package_atom_count += 1

                self._enforce_topic_title_boundary(
                    db=db,
                    package_id=package.id,
                    package_title=package_title,
                    fallback_point_id=fallback_point.id,
                )

                package.parse_status = "success"
                total_block_count += package_block_count
                total_atom_count += package_atom_count
                total_point_count += len(package_point_links)
                package_results.append(
                    {
                        "package_id": package.id,
                        "package_title": package.package_title,
                        "section_count": len(package_sections),
                        "block_count": package_block_count,
                        "atom_count": package_atom_count,
                        "knowledge_point_count": len(package_point_links),
                    }
                )

            db.commit()

            topic_question_metrics: Dict[str, object] = {"status": "skipped", "question_count": 0, "papers_created": 0}
            if package_results and source_path.suffix.lower() == ".pdf":
                from analyzer.app.question_bank_parser import QuestionBankIngestionService

                topic_question_metrics = QuestionBankIngestionService().ingest_topic_packages_questions(
                    db,
                    source_document.id,
                    progress_callback=progress_callback,
                )
                self._notify_progress(
                    progress_callback,
                    f"专题题完成：papers={topic_question_metrics.get('papers_created')} questions={topic_question_metrics.get('question_count')}",
                )

            source_document.parse_status = "success"
            metrics_json = {
                "package_count": len(package_results),
                "section_count": sum(int(item["section_count"]) for item in package_results),
                "block_count": total_block_count,
                "atom_count": total_atom_count,
                "knowledge_point_count": total_point_count,
                "topic_question_papers": topic_question_metrics.get("papers_created", 0),
                "topic_question_count": topic_question_metrics.get("question_count", 0),
            }
            self._finish_parse_job(
                db,
                parse_job,
                status="success",
                output_location=f"knowledge_packages:{','.join(str(item['package_id']) for item in package_results)}",
                metrics_json=metrics_json,
            )
            db.commit()

            self._notify_progress(
                progress_callback,
                f"摄入成功：packages={len(package_results)} blocks={total_block_count} atoms={total_atom_count}",
            )
            primary_package = package_results[0] if package_results else None
            return {
                "status": "success",
                "source_document_id": source_document.id,
                "package_count": len(package_results),
                "package_id": primary_package["package_id"] if primary_package else None,
                "package_title": primary_package["package_title"] if primary_package else None,
                "section_count": sum(int(item["section_count"]) for item in package_results),
                "block_count": total_block_count,
                "atom_count": total_atom_count,
                "knowledge_point_count": total_point_count,
                "packages": package_results,
                "topic_question_metrics": topic_question_metrics,
            }

        except Exception as exc:
            self._notify_progress(progress_callback, f"摄入失败：{exc.__class__.__name__}: {exc}")
            db.rollback()
            try:
                source_document.parse_status = "failed"
                db.add(source_document)
                db.commit()
            except Exception:
                db.rollback()
            parse_job_in_db = db.query(models.DocumentParseJob).filter(models.DocumentParseJob.id == parse_job.id).first()
            if parse_job_in_db:
                self._fail_parse_job(db, parse_job_in_db, str(exc))
                db.commit()
            raise

    # ------------------------------------------------------------------
    # DOCX 专题摄入 —— 富文本提取 + 段落分类 + 知识块/题目入库
    # ------------------------------------------------------------------

    def _ingest_docx_topic(
        self,
        db: Session,
        source_document: models.SourceDocument,
        source_path: Path,
        parse_job: models.DocumentParseJob,
        *,
        fallback_point: Optional[models.KnowledgePoint] = None,
        progress_callback: Optional[Callable[[str], None]] = None,
    ) -> Dict[str, object]:
        notify = lambda msg: self._notify_progress(progress_callback, msg)

        if fallback_point is None:
            fallback_point = self._get_or_create_knowledge_point(
                db, canonical_name="未归类知识点",
                subject=source_document.subject,
                grade_scope=source_document.grade,
                source_origin="model",
            )

        notify("DOCX 富文本提取（段落/表格/图片/公式）…")
        from .question_bank_rich_content import extract_docx_blocks

        asset_dir = Path(QUESTION_BANK_ASSET_DIR) / f"document_{source_document.id}"
        asset_dir.mkdir(parents=True, exist_ok=True)
        blocks = extract_docx_blocks(source_path, asset_dir)
        fi, fl = _docx_formula_image_layout_stats(blocks)
        notify(f"  提取完成：{len(blocks)} 个结构化块")
        notify(
            f"  公式栅格图(omml_raster)={fi} 张，含 Word 占位 layout_*={fl} 张"
            f"（layout 为 0 时浏览器仍按「一行半」压等高，与前端版本无关）"
        )

        if self._ingest_verbose_enabled():
            self._ingest_verbose_write(
                "docx/source_meta.txt",
                "\n".join(
                    [
                        f"storage_path={source_path}",
                        f"source_document_id={source_document.id}",
                        f"file_name={source_document.file_name}",
                        f"blocks_total={len(blocks)}",
                        f"formula_omml_raster={fi}",
                        f"formula_layout_placeholders={fl}",
                    ]
                )
                + "\n",
            )
            block_index_rows: List[Dict[str, Any]] = []
            for i, blk in enumerate(blocks):
                bt = str(blk.get("type") or blk.get("block_type") or "unknown")
                block_index_rows.append(
                    {
                        "index": i,
                        "type": bt,
                        "plain_text_len": len(_block_text(blk) or ""),
                    }
                )
            self._ingest_verbose_write(
                "docx/blocks_index.json",
                json.dumps(block_index_rows, ensure_ascii=False, indent=2),
            )
            preview_blocks: List[Dict[str, Any]] = []
            for i, blk in enumerate(blocks[:50]):
                preview_blocks.append(
                    {
                        "index": i,
                        "plain_text": (_block_text(blk) or "")[:6000],
                    }
                )
            self._ingest_verbose_write(
                "docx/blocks_plain_preview_first50.json",
                json.dumps(preview_blocks, ensure_ascii=False, indent=2),
            )

        notify("段落分类（知识内容 vs 题目）…")
        content_segments, question_segments, embedded_q_split = classify_topic_docx_blocks(blocks)
        notify(f"  专题内容段={len(content_segments)} 题目段={len(question_segments)}")
        if embedded_q_split:
            notify(
                f"  同段「解析/答案」后紧跟题号已拆开段落数={embedded_q_split}"
                f"（避免题号不在块首导致漏题；详见 topic_docx_embedded_q_split 日志）",
            )

        if self._ingest_verbose_enabled():
            cs_rows = [
                {
                    "section_title": seg.section_title,
                    "block_index_start": seg.block_index_start,
                    "block_index_end": seg.block_index_end,
                    "plain_text_len": len(seg.plain_text or ""),
                    "plain_text": (seg.plain_text or "")[:12000],
                }
                for seg in content_segments
            ]
            qs_rows = [
                {
                    "question_no": seg.question_no,
                    "section_title": seg.section_title,
                    "block_index_start": seg.block_index_start,
                    "block_index_end": seg.block_index_end,
                    "plain_text_len": len("\n".join(_block_text(b) for b in seg.blocks)),
                    "plain_text": "\n".join(_block_text(b) for b in seg.blocks)[:12000],
                }
                for seg in question_segments
            ]
            self._ingest_verbose_write(
                "docx/segmentation.json",
                json.dumps(
                    {
                        "embedded_q_split": embedded_q_split,
                        "content_segments": cs_rows,
                        "question_segments": qs_rows,
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
            )

        # 同一 SourceDocument 的 DOCX 专题只保留「本次摄入」生成的包。未勾选强制重摄入时，
        # 若不在此清理，会反复堆积多个 KnowledgePackage，知识点管理页点到旧包即像「考点遗留」。
        notify("DOCX 专题：清理本来源文档已有专题包与关联块（单次摄入唯一包）…")
        self._clear_existing_package_artifacts(db, source_document.id)

        doc_title = source_path.stem
        package_title = re.sub(r"^\d+\s*", "", doc_title).strip() or doc_title
        package = models.KnowledgePackage(
            source_document_id=source_document.id,
            tenant_id=source_document.tenant_id,
            package_title=package_title,
            package_type="topic",
            subject=source_document.subject,
            grade=source_document.grade,
            page_range_json={"start": 1, "end": 1},
            outline_json=self._build_docx_outline(content_segments, question_segments),
            summary_text=self._build_docx_summary(content_segments),
            parse_status="running",
            review_status="draft",
        )
        db.add(package)
        db.flush()

        notify(f"创建 KnowledgePackage id={package.id} «{package_title}»")

        package_point_links: Dict[int, models.KnowledgePackagePoint] = {}
        topic_key = re.sub(r"\s+", "", package_title)
        block_count = 0
        atom_count = 0
        content_block_ids: List[int] = []
        topic_point_mode = _topic_docx_point_mode()
        if topic_point_mode == "llm":
            notify("知识点来源=llm：跳过正则/标题写入包↔知识点，仅保留未归类主块占位，稍后用大模型填充。")
        elif topic_point_mode == "both":
            notify("知识点来源=both：正则/标题 + 大模型两条链路均会执行。")
        else:
            notify("知识点来源=regex：仅正则/标题链路（不调用大模型块级标注）。")

        for order, seg in enumerate(content_segments, start=1):
            if topic_point_mode == "llm":
                points = [fallback_point]
            else:
                point_names = [
                    name
                    for name in self._extract_knowledge_point_names(seg.section_title, seg.plain_text)
                    if re.sub(r"\s+", "", name or "") != topic_key
                ]
                points = [
                    self._get_or_create_knowledge_point(
                        db, canonical_name=name,
                        subject=source_document.subject,
                        grade_scope=source_document.grade,
                        source_origin="model",
                    )
                    for name in point_names
                ]
                points = [p for p in points if re.sub(r"\s+", "", p.canonical_name or "") != topic_key]
                if not points:
                    points = [fallback_point]

            primary_point = points[0]
            block_role = self._infer_block_role(seg.section_title)
            normalized_text = self._normalize_text(seg.plain_text)

            rich_json = {
                "type": "block_group",
                "role": "topic_content",
                "section_title": seg.section_title,
                "plain_text": normalized_text,
                "blocks": [copy.deepcopy(b.get("render") or {}) for b in seg.blocks if b.get("render")],
            }

            kb = models.KnowledgeBlock(
                package_id=package.id,
                knowledge_point_id=primary_point.id,
                parent_block_id=None,
                block_order=order,
                section_path=f"{package_title}/{seg.section_title}",
                block_role=block_role,
                content_format="rich_docx",
                raw_text=seg.plain_text,
                normalized_text=normalized_text,
                rich_content_json=rich_json,
                source_page_no=1,
                source_anchor_json={
                    "block_index_start": seg.block_index_start,
                    "block_index_end": seg.block_index_end,
                },
                source_origin="structured_extraction",
                confidence=0.9,
                is_primary=(order == 1),
            )
            db.add(kb)
            db.flush()
            content_block_ids.append(kb.id)
            block_count += 1

            for point in points:
                record_knowledge_point_provenance(
                    db,
                    knowledge_point_id=point.id,
                    source_kind=SOURCE_KIND_KNOWLEDGE_BLOCK,
                    source_id=kb.id,
                    package_id=package.id,
                    origin_step="docx_topic_section",
                )

            if topic_point_mode != "llm":
                for idx_p, point in enumerate(points):
                    if point.id not in package_point_links:
                        link = models.KnowledgePackagePoint(
                            package_id=package.id,
                            knowledge_point_id=point.id,
                            relation_type="core" if idx_p == 0 else "supplement",
                            weight_score=1.0 if idx_p == 0 else 0.7,
                            order_in_package=len(package_point_links) + 1,
                            source_origin="model",
                            confidence=0.8 if idx_p == 0 else 0.7,
                            approved_status="pending",
                        )
                        db.add(link)
                        package_point_links[point.id] = link

            sentences = self._split_sentences(normalized_text)
            atom_type = ATOM_TYPE_BY_BLOCK_ROLE.get(block_role, "conclusion")
            for sentence in sentences[:3]:
                if len(sentence) < 10:
                    continue
                atom = models.KnowledgeAtom(
                    knowledge_point_id=primary_point.id,
                    package_id=package.id,
                    atom_type=atom_type,
                    canonical_text=sentence,
                    normalized_json={"source": "docx_section", "section_title": seg.section_title},
                    evidence_block_id=kb.id,
                    source_origin="model",
                    confidence=0.75,
                    review_status="draft",
                )
                db.add(atom)
                atom_count += 1

        self._enforce_topic_title_boundary(
            db=db, package_id=package.id,
            package_title=package_title, fallback_point_id=fallback_point.id,
        )

        package_point_links.clear()
        package_point_links.update(self._refresh_package_point_links_map(db, package.id))

        if topic_point_mode in ("both", "llm"):
            self._enrich_topic_docx_with_llm_block_points(
                db,
                package,
                content_segments,
                source_document,
                package_point_links,
                progress_callback=progress_callback,
                topic_point_mode=topic_point_mode,
            )

        notify(f"知识块写入完成：blocks={block_count} atoms={atom_count} points={len(package_point_links)}")

        topic_question_metrics: Dict[str, object] = {"status": "skipped", "question_count": 0, "papers_created": 0}
        if question_segments:
            notify(f"DOCX 专题题目入库：{len(question_segments)} 道题…")
            topic_question_metrics = self._persist_docx_topic_questions(
                db, source_document, package, question_segments,
                content_block_ids=content_block_ids,
                content_segments=content_segments,
                progress_callback=progress_callback,
            )
            notify(
                f"专题题完成：papers={topic_question_metrics.get('papers_created')} "
                f"questions={topic_question_metrics.get('question_count')}"
            )

        package.parse_status = "success"
        db.commit()

        graph_projection_summary: Dict[str, Any] = {"status": "skipped"}
        try:
            from .config import KNOWLEDGE_GRAPH_ENABLED as _KG_ENABLED
            from . import knowledge_graph_projection as _kg_proj

            if _KG_ENABLED:
                notify("知识图谱投影：把本包的业务关系投影到 entity_graph_edges…")
                graph_projection_summary = _kg_proj.project_package(db, package.id)
                notify(
                    "知识图谱投影完成："
                    f"inserted={graph_projection_summary.get('inserted')} "
                    f"deleted={graph_projection_summary.get('deleted')}"
                )
            else:
                notify("知识图谱投影：KNOWLEDGE_GRAPH_ENABLED=false，跳过。")
        except Exception as exc:
            logger.warning("Graph projection after ingest failed: %s", exc, exc_info=True)
            graph_projection_summary = {"status": "error", "reason": str(exc)}

        metrics_json = {
            "package_count": 1,
            "block_count": block_count,
            "atom_count": atom_count,
            "knowledge_point_count": len(package_point_links),
            "topic_question_papers": topic_question_metrics.get("papers_created", 0),
            "topic_question_count": topic_question_metrics.get("question_count", 0),
            "graph_projection": graph_projection_summary,
        }
        self._finish_parse_job(
            db, parse_job, status="success",
            output_location=f"knowledge_packages:{package.id}",
            metrics_json=metrics_json,
        )
        db.commit()

        notify(f"DOCX 专题摄入成功：package_id={package.id} blocks={block_count} questions={topic_question_metrics.get('question_count', 0)}")
        return {
            "status": "success",
            "source_document_id": source_document.id,
            "package_count": 1,
            "package_id": package.id,
            "package_title": package_title,
            "block_count": block_count,
            "atom_count": atom_count,
            "knowledge_point_count": len(package_point_links),
            "packages": [
                {
                    "package_id": package.id,
                    "package_title": package_title,
                    "section_count": len(content_segments),
                    "block_count": block_count,
                    "atom_count": atom_count,
                    "knowledge_point_count": len(package_point_links),
                }
            ],
            "topic_question_metrics": topic_question_metrics,
        }

    def _pick_content_segment_index_for_question(
        self,
        qseg: TopicQuestionSegment,
        content_segments: List[TopicContentSegment],
    ) -> Optional[int]:
        """按块下标：取「题」之前最近的一段专题正文，用于挂接知识点。"""
        if not content_segments:
            return None
        q_start = qseg.block_index_start
        best: Optional[int] = None
        for k, c in enumerate(content_segments):
            if c.block_index_start <= q_start <= c.block_index_end:
                return k
            if c.block_index_end <= q_start:
                best = k
        return best

    def _point_ids_for_evidence_block(self, db: Session, block_id: int) -> List[int]:
        kb = db.query(models.KnowledgeBlock).filter(models.KnowledgeBlock.id == block_id).first()
        if not kb:
            return []
        out: List[int] = []
        seen: Set[int] = set()
        if kb.knowledge_point_id is not None:
            pid0 = int(kb.knowledge_point_id)
            if pid0 not in seen:
                seen.add(pid0)
                out.append(pid0)
        anchor = kb.source_anchor_json if isinstance(kb.source_anchor_json, dict) else {}
        for row in anchor.get("llm_knowledge_points") or []:
            if not isinstance(row, dict):
                continue
            try:
                pid = int(row["knowledge_point_id"])
            except (TypeError, ValueError, KeyError):
                continue
            if pid not in seen:
                seen.add(pid)
                out.append(pid)
        return out

    def _collect_all_content_block_point_ids(self, db: Session, content_block_ids: List[int]) -> List[int]:
        seen: Set[int] = set()
        ordered: List[int] = []
        for bid in content_block_ids:
            for pid in self._point_ids_for_evidence_block(db, bid):
                if pid not in seen:
                    seen.add(pid)
                    ordered.append(pid)
        return ordered

    def _package_linked_point_ids(self, db: Session, package_id: int) -> List[int]:
        rows = (
            db.query(models.KnowledgePackagePoint.knowledge_point_id)
            .filter(models.KnowledgePackagePoint.package_id == package_id)
            .order_by(models.KnowledgePackagePoint.order_in_package.asc(), models.KnowledgePackagePoint.id.asc())
            .all()
        )
        out: List[int] = []
        seen: Set[int] = set()
        for (pid,) in rows:
            if pid is None or pid in seen:
                continue
            seen.add(int(pid))
            out.append(int(pid))
        return out

    def _package_representative_point_id(self, db: Session, package_id: int) -> Optional[int]:
        """本包「代表考点」：优先 core，其次 order_in_package 最小/id 最小。"""
        row = (
            db.query(models.KnowledgePackagePoint.knowledge_point_id)
            .filter(
                models.KnowledgePackagePoint.package_id == package_id,
                models.KnowledgePackagePoint.relation_type == "core",
            )
            .order_by(models.KnowledgePackagePoint.order_in_package.asc(), models.KnowledgePackagePoint.id.asc())
            .first()
        )
        if row and row[0] is not None:
            return int(row[0])
        row = (
            db.query(models.KnowledgePackagePoint.knowledge_point_id)
            .filter(models.KnowledgePackagePoint.package_id == package_id)
            .order_by(models.KnowledgePackagePoint.order_in_package.asc(), models.KnowledgePackagePoint.id.asc())
            .first()
        )
        if row and row[0] is not None:
            return int(row[0])
        return None

    @staticmethod
    def _normalize_text_for_scoring(text: str) -> str:
        if not text:
            return ""
        return re.sub(r"\s+", "", text)

    @staticmethod
    def _score_point_against_stem(canonical_name: str, stem_norm: str) -> float:
        """基于字符覆盖的轻量打分，范围 [0, 1]。canonical_name 为空或 stem 为空返回 0。"""
        if not canonical_name or not stem_norm:
            return 0.0
        name_norm = KnowledgePointIngestionService._normalize_text_for_scoring(canonical_name)
        if not name_norm:
            return 0.0
        if name_norm in stem_norm:
            return 1.0
        chars = [c for c in name_norm if re.match(r"[\u4e00-\u9fffA-Za-z0-9]", c)]
        if not chars:
            return 0.0
        hit = sum(1 for c in set(chars) if c in stem_norm)
        return hit / max(1, len(set(chars)))

    def _rank_points_for_question(
        self,
        db: Session,
        candidate_point_ids: List[int],
        stem_text: str,
        top_k: int,
        min_score: float,
    ) -> List[Tuple[int, float, str]]:
        """
        对候选知识点按「知识点名在题干中的字符覆盖」打分，返回 (pid, score, tier) 列表。
        tier ∈ {"strong","adjacent"}。strong=名称整体命中题干；adjacent=部分字符覆盖通过阈值。
        不命中的点返回 score=0, tier="adjacent"（供上层决定是否启用保底）。
        """
        if not candidate_point_ids:
            return []
        name_rows = (
            db.query(models.KnowledgePoint.id, models.KnowledgePoint.canonical_name)
            .filter(models.KnowledgePoint.id.in_(candidate_point_ids))
            .all()
        )
        name_map = {int(pid): (nm or "") for pid, nm in name_rows}
        stem_norm = self._normalize_text_for_scoring(stem_text)

        scored: List[Tuple[int, float, str]] = []
        for pid in candidate_point_ids:
            name = name_map.get(int(pid), "")
            score = self._score_point_against_stem(name, stem_norm)
            if score >= 1.0:
                scored.append((int(pid), score, "strong"))
            elif score >= min_score:
                scored.append((int(pid), score, "adjacent"))
        # 按分数降序截 Top-K
        scored.sort(key=lambda r: (-(r[1]), r[0]))
        if top_k > 0:
            scored = scored[:top_k]
        return scored

    def _bridge_collect_vector_point_scores(
        self,
        db: Session,
        package_id: int,
        query_text: str,
        allowed_point_ids: Set[int],
        *,
        max_results: int = 36,
    ) -> Dict[int, float]:
        """对本包检索索引做 hybrid 查询，按 knowledge_point_id 聚合最高分（需 KNOWLEDGE_RAG_ENABLED）。"""
        if not KNOWLEDGE_RAG_ENABLED or not allowed_point_ids:
            return {}
        q = (query_text or "").strip()
        if not q:
            return {}
        try:
            from analyzer.app.knowledge_point_retriever import search_knowledge_documents
        except Exception as exc:  # noqa: BLE001
            logger.warning("bridge vector import failed: %s", exc)
            return {}
        try:
            data = search_knowledge_documents(
                db,
                query=q[:_TOPIC_BRIDGE_QUERY_MAX],
                top_k=max(max_results * 3, 24),
                package_id=package_id,
            )
        except Exception as exc:  # noqa: BLE001
            logger.warning("bridge vector search failed: %s", exc)
            return {}
        scores: Dict[int, float] = {}
        for row in data.get("results") or []:
            meta = row.get("metadata") or {}
            et = str(meta.get("entity_type") or "")
            pid: Optional[int] = None
            if et == "knowledge_point":
                try:
                    pid = int(meta.get("entity_id") or meta.get("knowledge_point_id") or 0)
                except (TypeError, ValueError):
                    pid = None
            else:
                raw = meta.get("knowledge_point_id")
                if raw is not None:
                    try:
                        pid = int(raw)
                    except (TypeError, ValueError):
                        pid = None
            if pid is None or pid not in allowed_point_ids:
                continue
            if et == "knowledge_question_bridge":
                continue
            sc = float(row.get("score") or 0.0)
            scores[pid] = max(scores.get(pid, 0.0), sc)
        return scores

    def _bridge_rank_from_vector(
        self,
        db: Session,
        package_id: int,
        stem_text: str,
        allowed_point_ids: Set[int],
        top_k: int,
    ) -> List[Tuple[int, float, str]]:
        scores = self._bridge_collect_vector_point_scores(
            db, package_id, stem_text, allowed_point_ids, max_results=max(24, top_k * 6)
        )
        if not scores:
            return []
        max_s = max(scores.values()) or 1.0
        ranked: List[Tuple[int, float, str]] = []
        for pid, s in sorted(scores.items(), key=lambda x: -x[1])[: max(top_k, 1)]:
            norm = s / max_s
            tier = "strong" if norm >= 0.82 else "adjacent"
            ranked.append((int(pid), norm, tier))
        return ranked

    def _bridge_rank_hybrid(
        self,
        db: Session,
        package_id: int,
        candidate_ids: List[int],
        stem_text: str,
        allowed_point_ids: Set[int],
        top_k: int,
        min_score: float,
    ) -> List[Tuple[int, float, str]]:
        """字符覆盖 + 检索向量分数加权融合，再截 Top-K。"""
        relax = min(0.35, float(min_score))
        overlap = self._rank_points_for_question(
            db, candidate_ids, stem_text, max(top_k * 4, 16), relax
        )
        overlap_map = {pid: (sc, tier) for pid, sc, tier in overlap}
        vec = self._bridge_collect_vector_point_scores(
            db, package_id, stem_text, allowed_point_ids, max_results=max(32, top_k * 8)
        )
        max_v = max(vec.values()) if vec else 0.0
        max_v = max_v if max_v > 0 else 1.0
        pool: Set[int] = set(candidate_ids) | set(vec.keys())
        pool &= allowed_point_ids
        merged: List[Tuple[int, float, str]] = []
        for pid in pool:
            o_sc, o_tier = overlap_map.get(pid, (0.0, "adjacent"))
            vn = (vec.get(pid, 0.0) / max_v) if vec else 0.0
            combined = 0.5 * float(o_sc) + 0.5 * float(vn)
            if combined <= 0 and pid not in vec and pid not in overlap_map:
                continue
            if o_sc >= float(min_score):
                tier = o_tier
            elif vn >= 0.82:
                tier = "strong"
            elif combined >= 0.25 or vn > 0.15:
                tier = "adjacent"
            else:
                continue
            merged.append((int(pid), combined, tier))
        merged.sort(key=lambda r: (-(r[1]), 0 if r[2] == "strong" else 1, -r[0]))
        out: List[Tuple[int, float, str]] = []
        seen: Set[int] = set()
        for row in merged:
            if row[0] in seen:
                continue
            seen.add(row[0])
            out.append(row)
            if len(out) >= top_k:
                break
        return out

    def _bridge_rank_with_llm(
        self,
        db: Session,
        package: models.KnowledgePackage,
        stem_text: str,
        pool_point_ids: List[int],
        max_links: int,
        progress_callback: Optional[Callable[[str], None]] = None,
        *,
        question_item_id: Optional[int] = None,
    ) -> List[Tuple[int, float, Any]]:
        """按题调用大模型，从候选中挑选 0～max_links 条；成功时返回 (pid, relevance, confidence) 三数值元组。

        失败或旧版 JSON 回退时仍可能返回 (pid, score, tier_str) 供 overlap 式分支消费（第三项为 str）。"""
        notify = lambda msg: self._notify_progress(progress_callback, msg)
        q = (stem_text or "").strip()

        def _audit(
            model_name: Optional[str],
            raw: Optional[str],
            parse_ok: bool,
            accepted: List[int],
            skipped_reason: Optional[str] = None,
        ) -> None:
            if question_item_id is None:
                return
            self._package_append_question_bridge_llm_debug_entry(
                package,
                question_item_id=question_item_id,
                model_name=model_name,
                response_text=raw,
                parse_ok=parse_ok,
                accepted_point_ids=accepted,
                skipped_reason=skipped_reason,
            )

        if not q or not pool_point_ids:
            _audit(None, None, False, [], "empty_stem_or_pool")
            return []
        try:
            sync_llm_step_configs(db)
            sync_prompt_step_configs(db)
        except Exception as exc:  # noqa: BLE001
            logger.warning("topic question bridge LLM: sync configs failed: %s", exc)
            _audit(None, None, False, [], f"sync_configs_failed:{exc}")
            return []

        llm_cfg = resolve_step_llm_config(
            db, TOPIC_DOCX_QUESTION_BRIDGE_STEP_KEY, allow_generic_fallback=True
        )
        if not llm_cfg or not llm_cfg.get("api_url") or not llm_cfg.get("api_key"):
            llm_cfg = resolve_step_llm_config(
                db, TOPIC_DOCX_BLOCK_POINTS_STEP_KEY, allow_generic_fallback=True
            )
        if not llm_cfg or not llm_cfg.get("api_url") or not llm_cfg.get("api_key"):
            notify("  题-知识点桥接 LLM：未解析到模型配置，跳过。")
            _audit(None, None, False, [], "no_llm_config")
            return []

        model_name = llm_cfg.get("model_name")

        uniq: List[int] = []
        seen: Set[int] = set()
        for p in pool_point_ids:
            if p in seen:
                continue
            seen.add(p)
            uniq.append(int(p))
            if len(uniq) >= _TOPIC_BRIDGE_LLM_POOL_CAP:
                break

        name_rows = (
            db.query(models.KnowledgePoint.id, models.KnowledgePoint.canonical_name)
            .filter(models.KnowledgePoint.id.in_(uniq))
            .all()
        )
        id_to_name = {int(i): (n or "") for i, n in name_rows}
        cand = [
            {"knowledge_point_id": pid, "canonical_name": (id_to_name.get(pid, "") or "")[:160]}
            for pid in uniq
        ]
        candidates_json = json.dumps({"candidates": cand}, ensure_ascii=False)
        max_links_eff = max(0, int(max_links))
        prompt_cfg = resolve_step_prompt(
            db,
            TOPIC_DOCX_QUESTION_BRIDGE_STEP_KEY,
            variables={
                "candidates_json": candidates_json,
                "question_text": q[:4000],
                "package_title": (package.package_title or "")[:240],
                "max_pick": str(max(1, max_links_eff or 1)),
                "max_links": str(max_links_eff),
            },
        )
        prompt_text = ((prompt_cfg or {}).get("prompt_text") or "").strip()
        if not prompt_text:
            prompt_text = (
                "你是一名中学教研助手。请根据题干与候选知识点判断关联，只输出 JSON：\n"
                '{"links":[{"knowledge_point_id":整数,"relevance":0到1,"confidence":0到1},...]}\n'
                "ID 必须全部来自候选；links 可空；最多 "
                f"{max_links_eff} 条；按 relevance 降序。\n\n"
                f"专题：{(package.package_title or '')[:200]}\n\n题干：\n{q[:4000]}\n\n候选：\n{candidates_json}"
            )

        if self._ingest_verbose_enabled() and question_item_id is not None:
            qslug = int(question_item_id)
            self._ingest_verbose_write(
                f"llm/topic_docx_question_bridge/q_{qslug}_request.json",
                json.dumps(
                    {
                        "step_key": TOPIC_DOCX_QUESTION_BRIDGE_STEP_KEY,
                        "model_name": model_name,
                        "question_item_id": qslug,
                        "package_id": package.id,
                        "max_links": max_links_eff,
                        "candidate_pool_size": len(uniq),
                        "messages": [{"role": "user", "content": prompt_text}],
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
            )
            self._ingest_verbose_write(
                f"llm/topic_docx_question_bridge/q_{qslug}_user_prompt.txt",
                prompt_text,
            )

        raw = call_llm(
            [{"role": "user", "content": prompt_text}],
            llm_cfg,
            json_mode=True,
        )
        if self._ingest_verbose_enabled() and question_item_id is not None:
            self._ingest_verbose_write(
                f"llm/topic_docx_question_bridge/q_{int(question_item_id)}_response_raw.txt",
                raw if raw is not None else "(null response)",
            )
        data = self._parse_llm_json_object_response(raw)
        if not data:
            notify("  题-知识点桥接 LLM：JSON 解析失败，回退其它排序。")
            _audit(model_name, raw, False, [], "json_parse_failed")
            return []

        allow = set(uniq)
        min_rel = _topic_bridge_llm_min_relevance()

        def _clamp01(x: Any, default: float) -> float:
            try:
                v = float(x)
            except (TypeError, ValueError):
                v = default
            return max(0.0, min(1.0, v))

        ranked: List[Tuple[int, float, Any]] = []
        links_raw = data.get("links")
        if isinstance(links_raw, list) and links_raw:
            scored: List[Tuple[int, float, float]] = []
            for item in links_raw:
                if max_links_eff > 0 and len(scored) >= max_links_eff:
                    break
                if not isinstance(item, dict):
                    continue
                try:
                    pid = int(item.get("knowledge_point_id"))
                except (TypeError, ValueError):
                    continue
                if pid not in allow:
                    continue
                rel = _clamp01(item.get("relevance"), 0.5)
                conf = item.get("confidence", rel)
                conf_f = _clamp01(conf, rel)
                if rel < min_rel:
                    continue
                scored.append((pid, rel, conf_f))
            scored.sort(key=lambda r: (-r[1], -r[2], r[0]))
            seen_pid: Set[int] = set()
            for pid, rel, conf_f in scored:
                if pid in seen_pid:
                    continue
                seen_pid.add(pid)
                ranked.append((pid, rel, conf_f))
                if max_links_eff > 0 and len(ranked) >= max_links_eff:
                    break
            accepted = [r[0] for r in ranked]
            _audit(model_name, raw, True, accepted, None if accepted else "empty_links_after_filter")
            return ranked

        ids_raw = data.get("knowledge_point_ids")
        if not isinstance(ids_raw, list):
            _audit(model_name, raw, False, [], "no_links_and_ids_not_list")
            return []
        legacy: List[Tuple[int, float, float]] = []
        for i, x in enumerate(ids_raw):
            try:
                pid = int(x)
            except (TypeError, ValueError):
                continue
            if pid not in allow:
                continue
            rel = max(0.0, min(1.0, 0.9 - i * 0.07))
            if rel < min_rel:
                continue
            conf_f = max(0.35, min(0.98, rel * 0.96 + 0.02))
            legacy.append((pid, rel, conf_f))
            if max_links_eff > 0 and len(legacy) >= max_links_eff:
                break
        accepted = [r[0] for r in legacy]
        _audit(model_name, raw, True, accepted, "legacy_knowledge_point_ids_shape" if accepted else "empty_legacy_ids")
        return legacy

    def _sync_docx_topic_question_knowledge_links(
        self,
        db: Session,
        package: models.KnowledgePackage,
        content_segments: List[TopicContentSegment],
        content_block_ids: List[int],
        question_item_ids: List[int],
        paired_segments: List[Optional[TopicQuestionSegment]],
        progress_callback: Optional[Callable[[str], None]] = None,
    ) -> Dict[str, int]:
        """
        写入 KnowledgeQuestionLink，使「专题包 -> 知识点 -> 题」与 list_package_related_questions 对齐。
        专题材料题在 persist_questions 中只写了 KnowledgePackageQuestion；这里补齐知识点↔题链。

        产品不变量（与 plan 对齐）：
        - 每道本包 KPQ 题目至少挂一条 KnowledgeQuestionLink；
        - 所有链接的 knowledge_point_id 必须属于本包 KnowledgePackagePoint 集合；
        - 启发式无命中时用 topic_fallback + 低 confidence 保底，报告侧可降权折叠。
        """
        notify = lambda msg: self._notify_progress(progress_callback, msg)
        stats = {
            "questions": len(question_item_ids or []),
            "new_links_strong": 0,
            "new_links_adjacent": 0,
            "new_links_fallback": 0,
            "questions_bridged": 0,
            "questions_fallback_only": 0,
            "rank_mode": "overlap",
        }
        if not question_item_ids:
            return stats

        if len(paired_segments) != len(question_item_ids):
            logger.warning(
                "topic docx question bridge: paired_segments=%s qids=%s; padding with None",
                len(paired_segments),
                len(question_item_ids),
            )
            paired_segments = (list(paired_segments) + [None] * len(question_item_ids))[: len(question_item_ids)]

        allowed_point_ids = set(self._package_linked_point_ids(db, package.id))
        representative_point_id = self._package_representative_point_id(db, package.id)

        fallback_point_ids = [
            pid for pid in self._collect_all_content_block_point_ids(db, content_block_ids)
            if pid in allowed_point_ids
        ]
        if not fallback_point_ids and allowed_point_ids:
            fallback_point_ids = list(allowed_point_ids)

        if not allowed_point_ids:
            notify("  知识点↔题桥接：本包尚无 KnowledgePackagePoint，跳过。")
            return stats

        top_k = max(1, int(os.environ.get("KNOWLEDGE_POINT_BRIDGE_TOPK", "3")))
        try:
            min_score = float(os.environ.get("KNOWLEDGE_POINT_BRIDGE_MIN_SCORE", "0.5"))
        except ValueError:
            min_score = 0.5

        rank_mode = (os.environ.get("KNOWLEDGE_POINT_BRIDGE_RANK_MODE") or "overlap").strip().lower()
        if rank_mode not in ("overlap", "vector", "vector_then_overlap", "llm", "llm_then_overlap", "hybrid"):
            rank_mode = "overlap"
        stats["rank_mode"] = rank_mode
        if rank_mode != "overlap":
            notify(
                f"  题-知识点桥接排序模式：{rank_mode}"
                f"（RAG={'on' if KNOWLEDGE_RAG_ENABLED else 'off'}）。"
            )

        if rank_mode in ("llm", "llm_then_overlap"):
            self._package_reset_question_bridge_llm_debug(package, rank_mode=rank_mode)

        def _filter_allowed(pids: List[int]) -> List[int]:
            seen: Set[int] = set()
            out: List[int] = []
            for p in pids:
                if p in allowed_point_ids and p not in seen:
                    seen.add(p)
                    out.append(p)
            return out

        def _tier_to_scores(tier: str) -> Tuple[str, float, float]:
            if tier == "strong":
                return ("topic_strong", 0.85, 0.88)
            return ("topic_adjacent", 0.6, 0.65)

        for qid, qseg in zip(question_item_ids, paired_segments):
            if qseg is None:
                candidate_ids = list(fallback_point_ids)
                stem_text = ""
            else:
                idx = self._pick_content_segment_index_for_question(qseg, content_segments)
                if idx is None:
                    idx = 0
                if idx < 0 or idx >= len(content_block_ids):
                    candidate_ids = list(fallback_point_ids)
                else:
                    candidate_ids = self._point_ids_for_evidence_block(db, content_block_ids[idx])
                stem_text = "\n".join(_block_text(b) for b in (qseg.blocks or []))

            candidate_ids = _filter_allowed(candidate_ids)
            if not candidate_ids:
                candidate_ids = _filter_allowed(list(fallback_point_ids))

            stem_ok = bool((stem_text or "").strip())
            ranked: List[Tuple[int, float, Any]] = []
            pool_for_llm = list(dict.fromkeys([*candidate_ids, *fallback_point_ids]))

            if rank_mode == "overlap":
                ranked = self._rank_points_for_question(db, candidate_ids, stem_text, top_k, min_score)
            elif rank_mode == "vector":
                ranked = self._bridge_rank_from_vector(
                    db, package.id, stem_text, allowed_point_ids, top_k
                )
                if not ranked:
                    ranked = self._rank_points_for_question(db, candidate_ids, stem_text, top_k, min_score)
            elif rank_mode == "vector_then_overlap":
                ranked = self._bridge_rank_from_vector(
                    db, package.id, stem_text, allowed_point_ids, top_k
                )
                if not ranked:
                    ranked = self._rank_points_for_question(db, candidate_ids, stem_text, top_k, min_score)
            elif rank_mode == "llm":
                if stem_ok:
                    ranked = self._bridge_rank_with_llm(
                        db,
                        package,
                        stem_text,
                        pool_for_llm,
                        _topic_bridge_llm_max_links(),
                        progress_callback,
                        question_item_id=qid,
                    )
                if not ranked:
                    ranked = self._rank_points_for_question(db, candidate_ids, stem_text, top_k, min_score)
            elif rank_mode == "llm_then_overlap":
                if stem_ok:
                    ranked = self._bridge_rank_with_llm(
                        db,
                        package,
                        stem_text,
                        pool_for_llm,
                        _topic_bridge_llm_max_links(),
                        progress_callback,
                        question_item_id=qid,
                    )
                if not ranked:
                    ranked = self._rank_points_for_question(db, candidate_ids, stem_text, top_k, min_score)
            elif rank_mode == "hybrid":
                ranked = self._bridge_rank_hybrid(
                    db, package.id, candidate_ids, stem_text, allowed_point_ids, top_k, min_score
                )
                if not ranked:
                    ranked = self._rank_points_for_question(db, candidate_ids, stem_text, top_k, min_score)
            else:
                ranked = self._rank_points_for_question(db, candidate_ids, stem_text, top_k, min_score)

            added_for_qid = 0
            strong_th = _topic_bridge_llm_strong_relevance_threshold()
            for row in ranked:
                if not isinstance(row, tuple) or len(row) != 3:
                    continue
                try:
                    pid = int(row[0])
                except (TypeError, ValueError):
                    continue
                third = row[2]
                if isinstance(third, str):
                    relation_type, relevance, confidence = _tier_to_scores(str(third))
                    tier_key = "strong" if third == "strong" else "adjacent"
                elif isinstance(third, (int, float)):
                    relevance_m = max(0.0, min(1.0, float(row[1])))
                    confidence_m = max(0.05, min(1.0, float(third)))
                    relation_type = "topic_strong" if relevance_m >= strong_th else "topic_adjacent"
                    relevance = round(relevance_m, 4)
                    confidence = round(confidence_m, 2)
                    tier_key = "strong" if relation_type == "topic_strong" else "adjacent"
                else:
                    continue
                exists = (
                    db.query(models.KnowledgeQuestionLink)
                    .filter(
                        models.KnowledgeQuestionLink.knowledge_point_id == pid,
                        models.KnowledgeQuestionLink.question_item_id == qid,
                    )
                    .first()
                )
                if exists:
                    added_for_qid += 1
                    continue
                db.add(
                    models.KnowledgeQuestionLink(
                        knowledge_point_id=pid,
                        question_item_id=qid,
                        relation_type=relation_type,
                        relevance_score=relevance,
                        entry_point_text=None,
                        source_origin="model",
                        confidence=confidence,
                        approved_status="pending",
                    )
                )
                if tier_key == "strong":
                    stats["new_links_strong"] += 1
                else:
                    stats["new_links_adjacent"] += 1
                added_for_qid += 1

            if added_for_qid == 0 and representative_point_id is not None:
                exists = (
                    db.query(models.KnowledgeQuestionLink)
                    .filter(
                        models.KnowledgeQuestionLink.knowledge_point_id == representative_point_id,
                        models.KnowledgeQuestionLink.question_item_id == qid,
                    )
                    .first()
                )
                if not exists:
                    db.add(
                        models.KnowledgeQuestionLink(
                            knowledge_point_id=representative_point_id,
                            question_item_id=qid,
                            relation_type="topic_fallback",
                            relevance_score=0.3,
                            entry_point_text=None,
                            source_origin="model",
                            confidence=0.4,
                            approved_status="pending",
                        )
                    )
                    stats["new_links_fallback"] += 1
                stats["questions_fallback_only"] += 1
                stats["questions_bridged"] += 1
            elif added_for_qid > 0:
                stats["questions_bridged"] += 1

        total_new = stats["new_links_strong"] + stats["new_links_adjacent"] + stats["new_links_fallback"]
        if total_new or stats["questions_fallback_only"]:
            notify(
                "  知识点↔题桥接：新建链="
                f"{total_new}（strong={stats['new_links_strong']}, "
                f"adjacent={stats['new_links_adjacent']}, "
                f"fallback={stats['new_links_fallback']}），"
                f"覆盖题数 {stats['questions_bridged']}/{stats['questions']}（"
                f"仅保底 {stats['questions_fallback_only']} 道）。"
                f"Top-K={top_k}，min_score={min_score}，rank_mode={stats.get('rank_mode', 'overlap')}。"
            )
        db.flush()
        return stats

    def _persist_docx_topic_questions(
        self,
        db: Session,
        source_document: models.SourceDocument,
        package: models.KnowledgePackage,
        question_segments: List[TopicQuestionSegment],
        *,
        content_block_ids: List[int],
        content_segments: List[TopicContentSegment],
        progress_callback: Optional[Callable[[str], None]] = None,
    ) -> Dict[str, object]:
        from analyzer.app.question_bank_parser import QuestionBankIngestionService

        qbs = QuestionBankIngestionService()
        notify = lambda msg: self._notify_progress(progress_callback, msg)

        file_stem = Path(source_document.file_name or "document").stem
        title = f"{file_stem} · {package.package_title}"[:255]
        paper = models.Paper(
            source_document_id=source_document.id,
            knowledge_package_id=package.id,
            title=title,
            subject=source_document.subject,
            grade=source_document.grade,
            year=source_document.year,
            region=source_document.region,
            exam_type="topic_material",
            source_type="topic_material",
            is_canonical=False,
            review_status="draft",
            total_questions=0,
        )
        db.add(paper)
        db.flush()

        # 按 TopicQuestionSegment 逐段解析，避免把所有题目块拼成一条流导致选项/答案区边界错乱
        extracted_questions: List[object] = []
        paired_question_segments: List[TopicQuestionSegment] = []
        sub_question_nos: List[Optional[str]] = []
        built_from_per_segment = False
        for order, seg in enumerate(question_segments, start=1):
            if not seg.blocks:
                continue
            eq = qbs._parse_structured_question_segment(seg.question_no, seg.blocks)
            extracted_questions.append(
                dc_replace(
                    eq,
                    question_no=str(order),
                    original_question_label=(seg.question_no or "").strip() or None,
                )
            )
            paired_question_segments.append(seg)
            sub_question_nos.append(None)
        built_from_per_segment = True
        if not extracted_questions:
            built_from_per_segment = False
            paired_question_segments.clear()
            notify("  按段结构化切题为 0，尝试整段结构化 / 纯文本切题…")
            all_q_blocks: List[Dict[str, Any]] = []
            for seg in question_segments:
                all_q_blocks.extend(seg.blocks)
            extracted_questions = qbs._segment_structured_questions(all_q_blocks)
            if not extracted_questions:
                combined_text = "\n".join(
                    _block_text(b) for seg in question_segments for b in seg.blocks
                )
                extracted_questions = qbs.segment_questions(combined_text)
            if extracted_questions:
                extracted_questions = [
                    dc_replace(
                        eq,
                        question_no=str(i),
                        original_question_label=eq.original_question_label or eq.question_no,
                    )
                    for i, eq in enumerate(extracted_questions, start=1)
                ]

        if not extracted_questions:
            notify("  题目切分为 0，跳过")
            return {"status": "skipped", "question_count": 0, "papers_created": 0}

        if self._ingest_verbose_enabled():
            seg_meta = [
                {
                    "order_in_doc": i + 1,
                    "question_no": seg.question_no,
                    "section_title": seg.section_title,
                    "block_index_start": seg.block_index_start,
                    "block_index_end": seg.block_index_end,
                    "block_count": len(seg.blocks or []),
                }
                for i, seg in enumerate(question_segments)
            ]
            self._ingest_verbose_write(
                "questions/topic_question_segments.json",
                json.dumps(seg_meta, ensure_ascii=False, indent=2),
            )
            extracted_rows: List[Dict[str, Any]] = []
            for eq in extracted_questions:
                opts = [
                    {
                        "option_key": getattr(o, "option_key", ""),
                        "option_text_preview": (getattr(o, "option_text", "") or "")[:1200],
                    }
                    for o in (getattr(eq, "options", None) or [])[:20]
                ]
                extracted_rows.append(
                    {
                        "question_no": getattr(eq, "question_no", None),
                        "original_question_label": getattr(eq, "original_question_label", None),
                        "question_type": getattr(eq, "question_type", None),
                        "has_formula": getattr(eq, "has_formula", None),
                        "stem_text": (getattr(eq, "stem_text", "") or "")[:12000],
                        "text": (getattr(eq, "text", "") or "")[:12000],
                        "answer_text": (getattr(eq, "answer_text", "") or "")[:8000],
                        "analysis_text": (getattr(eq, "analysis_text", "") or "")[:8000],
                        "solution_text": (getattr(eq, "solution_text", "") or "")[:8000],
                        "options": opts,
                    }
                )
            self._ingest_verbose_write(
                "questions/extracted_questions.json",
                json.dumps(
                    {
                        "built_from_per_segment": built_from_per_segment,
                        "count": len(extracted_questions),
                        "questions": extracted_rows,
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
            )

        metrics = qbs.persist_questions(
            db=db,
            paper=paper,
            source_document=source_document,
            extracted_text="\n".join(_block_text(b) for seg in question_segments for b in seg.blocks),
            extracted_questions=extracted_questions,
            document_asset_count=0,
            knowledge_package_id=package.id,
        )

        question_count = int(metrics.get("question_count") or 0)
        qids = metrics.get("question_item_ids") or []

        if built_from_per_segment and len(paired_question_segments) == len(qids):
            paired_for_link: List[Optional[TopicQuestionSegment]] = list(paired_question_segments)
        else:
            paired_for_link = [None] * len(qids)

        bridge_stats = self._sync_docx_topic_question_knowledge_links(
            db=db,
            package=package,
            content_segments=content_segments,
            content_block_ids=content_block_ids,
            question_item_ids=qids,
            paired_segments=paired_for_link,
            progress_callback=progress_callback,
        )

        if self._ingest_verbose_enabled():
            self._ingest_verbose_write(
                "questions/persist_metrics.json",
                json.dumps(metrics, ensure_ascii=False, indent=2, default=str),
            )
            map_lines = ["order\tquestion_item_id\textracted_question_no\toriginal_question_label\tstem_preview"]
            for idx, qid in enumerate(qids):
                eq = extracted_questions[idx] if idx < len(extracted_questions) else None
                stem_pv = ((getattr(eq, "stem_text", "") or "")[:200].replace("\n", " ").replace("\t", " ")) if eq else ""
                map_lines.append(
                    "\t".join(
                        [
                            str(idx + 1),
                            str(qid),
                            str(getattr(eq, "question_no", "") if eq else ""),
                            str(getattr(eq, "original_question_label", "") if eq else ""),
                            stem_pv,
                        ]
                    )
                )
            self._ingest_verbose_write("questions/question_item_id_map.tsv", "\n".join(map_lines) + "\n")
            self._ingest_verbose_write(
                "bridge/bridge_sync_stats.json",
                json.dumps(bridge_stats, ensure_ascii=False, indent=2, default=str),
            )
            oj = package.outline_json if isinstance(package.outline_json, dict) else {}
            qbd = oj.get("question_bridge_llm_debug")
            if isinstance(qbd, dict):
                self._ingest_verbose_write(
                    "bridge/question_bridge_llm_debug.json",
                    json.dumps(qbd, ensure_ascii=False, indent=2, default=str),
                )
            llm_dbg = oj.get("llm_ingest_debug")
            if isinstance(llm_dbg, dict):
                self._ingest_verbose_write(
                    "llm/block_points_outline_llm_ingest_debug.json",
                    json.dumps(llm_dbg, ensure_ascii=False, indent=2, default=str),
                )

        db.commit()

        try:
            index_metrics = qbs.index_document_questions(db, source_document.id)
            notify(f"  向量索引完成：indexed={index_metrics.get('indexed_documents', 0)}")
        except Exception as exc:
            logger.warning("DOCX topic question vector index failed: %s", exc)
            index_metrics = {"indexed_documents": 0}

        if self._ingest_verbose_enabled():
            self._ingest_verbose_write(
                "questions/vector_index_metrics.json",
                json.dumps(index_metrics, ensure_ascii=False, indent=2, default=str),
            )

        return {
            "status": "success",
            "papers_created": 1,
            "question_count": question_count,
            "question_item_ids": qids,
            "paper_id": paper.id,
            "indexed_documents": index_metrics.get("indexed_documents", 0),
            "bridge_metrics": bridge_stats,
        }

    def _build_docx_outline(
        self,
        content_segments: List[TopicContentSegment],
        question_segments: List[TopicQuestionSegment],
    ) -> List[Dict[str, object]]:
        items: List[Dict[str, object]] = []
        order = 0
        for seg in content_segments:
            order += 1
            items.append({
                "order": order,
                "type": "content",
                "heading": seg.section_title,
                "block_range": [seg.block_index_start, seg.block_index_end],
            })
        for seg in question_segments:
            order += 1
            items.append({
                "order": order,
                "type": "question",
                "question_no": seg.question_no,
                "section": seg.section_title,
                "block_range": [seg.block_index_start, seg.block_index_end],
            })
        return items

    def _build_docx_summary(self, content_segments: List[TopicContentSegment]) -> Optional[str]:
        if not content_segments:
            return None
        snippets = [seg.plain_text[:80] for seg in content_segments[:5] if seg.plain_text]
        summary = "；".join(snippets)
        return summary[:1000] if summary else None

    def ingest_files_from_knowledge_points_dir(
        self,
        db: Session,
        files: List[str],
        force_reingest: bool = False,
        progress_callback: Optional[Callable[[str], None]] = None,
        ingest_run_assets_dir: Optional[Path] = None,
        ingest_run_dir: Optional[Path] = None,
    ) -> Dict[str, object]:
        if not files:
            return {"status": "skipped", "reason": "未指定文件", "processed": []}

        source = self._ensure_local_content_source(db)
        processed: List[Dict[str, object]] = []
        prev_log_dir = getattr(self, "_ingest_run_log_dir", None)
        self._ingest_run_log_dir = ingest_run_dir
        if ingest_run_dir:
            self._ingest_verbose_write(
                "VERBOSE_README.txt",
                "本目录由知识点摄入写入的详细审计文件（与 run.log 摘要配合使用）。\n"
                "子目录约定：\n"
                "  docx/        — DOCX 块统计、段落分类、正文/题段摘要\n"
                "  llm/         — 每次大模型调用的完整提示词与原始响应文本\n"
                "  questions/   — 题目切分结果、持久化结果、题号与 question_item_id 对照\n"
                "  bridge/      — 按题桥接写入 outline_json 的 question_bridge_llm_debug 快照（若有）\n"
                f"单文件大小上限（字节）：{_INGEST_VERBOSE_MAX_FILE_BYTES}，可用环境变量 KNOWLEDGE_INGEST_VERBOSE_MAX_FILE_BYTES 调整。\n",
            )
            self._notify_progress(
                progress_callback,
                f"详细审计目录：{ingest_run_dir}（见 VERBOSE_README.txt 与 docx/、llm/、questions/、bridge/）",
            )

        try:
            for raw_name in files:
                safe_name = os.path.basename(raw_name)
                source_path = self.knowledge_points_dir / safe_name
                if not source_path.exists() or not source_path.is_file():
                    self._notify_progress(progress_callback, f"[跳过] 文件不存在：{safe_name}")
                    continue
                if source_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                    self._notify_progress(progress_callback, f"[跳过] 不支持的扩展名：{safe_name}")
                    continue

                self._notify_progress(progress_callback, f"—— 开始处理文件：{safe_name} ——")

                source_document = (
                    db.query(models.SourceDocument)
                    .filter(models.SourceDocument.storage_url == str(source_path))
                    .first()
                )
                if not source_document:
                    source_document = models.SourceDocument(
                        source_id=source.id,
                        tenant_id=source.tenant_id,
                        file_name=safe_name,
                        file_ext=source_path.suffix.lower().lstrip("."),
                        mime_type=self._guess_mime_type(source_path.suffix.lower()),
                        storage_url=str(source_path),
                        parse_profile="knowledge_point",
                        title=None,
                        visibility_scope="tenant_private",
                        parse_status="pending",
                    )
                    db.add(source_document)
                    db.commit()
                    db.refresh(source_document)

                result = self.ingest_source_document(
                    db,
                    source_document_id=source_document.id,
                    force_reingest=force_reingest,
                    progress_callback=progress_callback,
                )
                processed.append({"file": safe_name, **result})

                if ingest_run_assets_dir is not None and source_path.suffix.lower() == ".pdf":
                    sub = re.sub(r'[<>:"/\\|?*]', "_", source_path.stem) or "pdf"
                    asset_subdir = ingest_run_assets_dir / sub
                    n_img = self._export_pdf_embedded_images(source_path, asset_subdir)
                    self._notify_progress(
                        progress_callback,
                        f"已导出 PDF 内嵌图 {n_img} 张到 assets/{sub}/（若内嵌为 0 则目录可能为空）",
                    )

        finally:
            self._ingest_run_log_dir = prev_log_dir

        return {
            "status": "complete",
            "processed_count": len(processed),
            "processed": processed,
        }

    def _extract_pages(self, source_path: Path) -> List[Tuple[int, str]]:
        suffix = source_path.suffix.lower()
        if suffix == ".pdf":
            from .pdf_structured_extractor import extract_page_structured, is_page_text_based
            pages: List[Tuple[int, str]] = []
            with fitz.open(source_path) as doc:
                for index, page in enumerate(doc, start=1):
                    if is_page_text_based(page, min_spans=5):
                        content = extract_page_structured(page, page_no=index)
                        text = content.plain_text.strip()
                    else:
                        text = (page.get_text() or "").strip()
                    if text:
                        pages.append((index, text))
            return pages
        if suffix == ".docx":
            doc = Document(str(source_path))
            lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            return [(1, "\n".join(lines))]

        text = source_path.read_text(encoding="utf-8", errors="ignore")
        return [(1, text)]

    def _split_into_sections(self, pages: List[Tuple[int, str]]) -> List[ParsedSection]:
        sections: List[ParsedSection] = []
        current_heading = "专题正文"
        current_lines: List[str] = []
        current_start_page = pages[0][0] if pages else 1
        current_end_page = current_start_page

        def flush_section() -> None:
            nonlocal current_lines, current_heading, current_start_page, current_end_page
            content = "\n".join(line for line in current_lines if line).strip()
            if content:
                sections.append(
                    ParsedSection(
                        heading=current_heading,
                        text=content,
                        page_start=current_start_page,
                        page_end=current_end_page,
                    )
                )
            current_lines = []

        for page_no, page_text in pages:
            lines = [line.strip() for line in page_text.splitlines() if line and line.strip()]
            for line in lines:
                if self._is_page_index_footer(line):
                    continue
                if self._looks_like_heading(line):
                    flush_section()
                    current_heading = line[:120]
                    current_start_page = page_no
                    current_end_page = page_no
                    continue
                current_lines.append(line)
                current_end_page = page_no

        flush_section()

        if not sections and pages:
            full_text = "\n".join(text for _, text in pages).strip()
            if full_text:
                sections.append(
                    ParsedSection(
                        heading="专题正文",
                        text=full_text,
                        page_start=pages[0][0],
                        page_end=pages[-1][0],
                    )
                )

        return sections

    def _is_page_index_footer(self, line: str) -> bool:
        t = (line or "").strip()
        if _PAGE_FRACTION_FOOTER.fullmatch(t) or _PAGE_DASH_RANGE.fullmatch(t):
            return True
        if re.fullmatch(r"\d{1,3}\s*", t):
            return True
        return False

    def _looks_like_numbered_section_heading(self, cleaned: str) -> bool:
        if re.match(r"^\d+(?:\.\d+){1,3}[、.。．]?\s*\S", cleaned):
            return True
        if re.match(r"^\d{1,2}[、.．]\s*\S", cleaned):
            return True
        return False

    def _looks_like_heading(self, line: str) -> bool:
        candidate = (line or "").strip()
        if not candidate:
            return False
        if self._is_page_index_footer(candidate):
            return False
        if _BARE_TOPIC_INDEX_LINE.fullmatch(candidate):
            return False

        cleaned = re.sub(r"\.{2,}.*$", "", candidate).strip()
        if re.search(r"\.{2,}", candidate):
            cleaned = re.sub(r"\s+\d{1,4}\s*$", "", cleaned).strip()
        if not cleaned:
            return False
        if len(cleaned) > 60:
            return False
        if cleaned.endswith(("。", "；", "，", "：", ":", "?", "？", "!", "！")):
            return False

        topic_num = re.match(r"^专题\s*\d+", cleaned)
        if topic_num:
            remainder = cleaned[topic_num.end() :].strip()
            return bool(remainder)

        if HEADING_PATTERN.match(cleaned):
            return True
        if self._looks_like_numbered_section_heading(cleaned):
            return True
        return bool(
            re.search(
                r"^(?:考向分析|考点[一二三四五六七八九十百0-9]+|知识点|类型[一二三四五六七八九十0-9]+|"
                r"题型[一二三四五六七八九十0-9]+|专题|方法总结|易错点|例题解析)",
                cleaned,
            )
        )

    def _extract_topic_title_from_heading(self, heading: str) -> Optional[str]:
        candidate = (heading or "").strip()
        if not candidate:
            return None

        candidate = re.sub(r"\.{2,}\s*\d+\s*$", "", candidate).strip()
        candidate = re.sub(r"\s{2,}", " ", candidate).strip("-_—· ")

        number_pattern = r"(?:\d{1,3}|[一二三四五六七八九十百零〇]{1,6})"
        for keyword in self.package_keywords or ("专题",):
            escaped = re.escape(keyword)
            matched = re.match(rf"^{escaped}\s*{number_pattern}\s*[：:.、\-\s]*(.+)$", candidate)
            if not matched:
                matched = re.match(rf"^{escaped}\s*[：:.、\-\s]+(.+)$", candidate)
            if not matched:
                continue

            title = matched.group(1).strip()
            title = re.sub(r"\.{2,}\s*\d+\s*$", "", title).strip()
            title = re.sub(r"\s+\d+\s*$", "", title).strip()
            title = re.sub(r"\s{2,}", " ", title).strip("-_—· ")
            if len(title) < 2:
                continue
            return title[:255]
        return None

    def _extract_page_hint(self, line: str) -> Optional[int]:
        candidate = (line or "").strip()
        if not candidate:
            return None
        if self._is_topic_index_line(candidate):
            return None

        dotted_match = re.search(r"\.{2,}\s*(\d{1,4})\s*$", candidate)
        if dotted_match:
            return int(dotted_match.group(1))

        plain_match = re.search(r"\s(\d{1,4})\s*$", candidate)
        if plain_match and "/" not in candidate:
            return int(plain_match.group(1))
        return None


    def _is_topic_index_line(self, line: str) -> bool:
        candidate = (line or "").strip()
        if not candidate:
            return False

        number_pattern = r"(?:\d{1,3}|[一二三四五六七八九十百零〇]{1,6})"
        for keyword in self.package_keywords or ("专题",):
            escaped = re.escape(keyword)
            if re.match(rf"^{escaped}\s*{number_pattern}\s*$", candidate):
                return True
        return False

    def _clean_topic_title_candidate(self, line: str) -> Optional[str]:
        candidate = (line or "").strip()
        if not candidate:
            return None

        candidate = re.sub(r"\.{2,}.*$", "", candidate).strip()
        candidate = re.sub(r"\s+\d+\s*$", "", candidate).strip()
        candidate = candidate.strip("-_—·:： ")
        if len(candidate) < 2 or len(candidate) > 80:
            return None
        if self._is_topic_index_line(candidate):
            return None
        if re.fullmatch(r"[\d\s./,，:：\-]+", candidate):
            return None
        return candidate

    def _extract_pdf_layout_lines(self, source_path: Path) -> List[LayoutLine]:
        layout_lines: List[LayoutLine] = []
        if source_path.suffix.lower() != ".pdf":
            return layout_lines

        with fitz.open(source_path) as doc:
            for page_index, page in enumerate(doc, start=1):
                raw_items: List[LayoutLine] = []
                text_dict = page.get_text("dict") or {}
                blocks = text_dict.get("blocks") or []
                for block in blocks:
                    if block.get("type") != 0:
                        continue
                    for line in block.get("lines") or []:
                        spans = line.get("spans") or []
                        if not spans:
                            continue
                        span_texts = [(span.get("text") or "").strip() for span in spans]
                        text = "".join(part for part in span_texts if part).strip()
                        if not text:
                            continue

                        x0 = min(float((span.get("bbox") or [0, 0, 0, 0])[0]) for span in spans)
                        y0 = min(float((span.get("bbox") or [0, 0, 0, 0])[1]) for span in spans)
                        y1 = max(float((span.get("bbox") or [0, 0, 0, 0])[3]) for span in spans)
                        font_sizes = [float(span.get("size") or 0.0) for span in spans if span.get("size")]
                        font_size = median(font_sizes) if font_sizes else 0.0
                        raw_items.append(
                            LayoutLine(
                                page_no=page_index,
                                text=text,
                                x0=x0,
                                y0=y0,
                                y_center=(y0 + y1) / 2.0,
                                font_size=font_size,
                            )
                        )

                if not raw_items:
                    continue

                raw_items.sort(key=lambda item: (item.y_center, item.x0))
                font_ref = median([item.font_size for item in raw_items if item.font_size > 0] or [10.0])
                y_tolerance = max(2.5, font_ref * 0.45)

                clusters: List[List[LayoutLine]] = []
                for item in raw_items:
                    if not clusters:
                        clusters.append([item])
                        continue
                    last_cluster = clusters[-1]
                    cluster_center = sum(part.y_center for part in last_cluster) / len(last_cluster)
                    if abs(item.y_center - cluster_center) <= y_tolerance:
                        last_cluster.append(item)
                    else:
                        clusters.append([item])

                for cluster in clusters:
                    cluster.sort(key=lambda part: part.x0)
                    merged_text = " ".join(part.text for part in cluster if part.text).strip()
                    if not merged_text:
                        continue
                    layout_lines.append(
                        LayoutLine(
                            page_no=page_index,
                            text=merged_text,
                            x0=min(part.x0 for part in cluster),
                            y0=min(part.y0 for part in cluster),
                            y_center=sum(part.y_center for part in cluster) / len(cluster),
                            font_size=median([part.font_size for part in cluster if part.font_size > 0] or [font_ref]),
                        )
                    )

        layout_lines.sort(key=lambda item: (item.page_no, item.y0, item.x0))
        return layout_lines

    def _build_plain_text_lines(self, pages: List[Tuple[int, str]]) -> List[LayoutLine]:
        plain_lines: List[LayoutLine] = []
        for page_no, page_text in pages:
            lines = [line.strip() for line in page_text.splitlines() if line and line.strip()]
            for row_index, line in enumerate(lines):
                plain_lines.append(
                    LayoutLine(
                        page_no=page_no,
                        text=line,
                        x0=0.0,
                        y0=float(row_index),
                        y_center=float(row_index),
                        font_size=0.0,
                    )
                )
        return plain_lines

    def _collect_topic_markers(self, source_path: Path, pages: List[Tuple[int, str]]) -> List[TopicMarker]:
        lines = self._extract_pdf_layout_lines(source_path)
        if not lines:
            lines = self._build_plain_text_lines(pages)

        markers_by_key: Dict[Tuple[str, int], TopicMarker] = {}
        for index, line_item in enumerate(lines):
            line = line_item.text
            title = self._extract_topic_title_from_heading(line)
            hinted_page = self._extract_page_hint(line)

            if not title and self._is_topic_index_line(line):
                lookahead = lines[index + 1 : index + 4]
                for candidate_item in lookahead:
                    if candidate_item.page_no != line_item.page_no:
                        break
                    candidate_line = candidate_item.text
                    title = self._clean_topic_title_candidate(candidate_line)
                    if title:
                        hinted_page = hinted_page or self._extract_page_hint(candidate_line)
                        break

            if not title:
                continue

            start_page = hinted_page if hinted_page and hinted_page > 0 else line_item.page_no
            key = (title, start_page)
            if key not in markers_by_key:
                markers_by_key[key] = TopicMarker(title=title, start_page=start_page, source_page=line_item.page_no)

        markers = sorted(markers_by_key.values(), key=lambda item: (item.start_page, item.source_page, item.title))
        deduped: List[TopicMarker] = []
        seen_titles = set()
        for marker in markers:
            title_key = marker.title.strip()
            if title_key in seen_titles:
                continue
            seen_titles.add(title_key)
            deduped.append(marker)
        return deduped

    def _split_into_packages(
        self,
        source_document: models.SourceDocument,
        source_path: Path,
        pages: List[Tuple[int, str]],
        sections: List[ParsedSection],
    ) -> List[ParsedPackage]:
        markers = self._collect_topic_markers(source_path, pages)

        if not markers:
            fallback_title = self._resolve_package_title(source_document, sections)
            return [ParsedPackage(title=fallback_title, sections=sections)]

        buckets: List[List[ParsedSection]] = [[] for _ in markers]
        for section in sections:
            marker_index = 0
            for index, marker in enumerate(markers):
                if section.page_start >= marker.start_page:
                    marker_index = index
                else:
                    break
            buckets[marker_index].append(section)

        packages: List[ParsedPackage] = []
        for index, marker in enumerate(markers):
            marker_sections = buckets[index]
            if not marker_sections:
                continue
            packages.append(ParsedPackage(title=marker.title, sections=marker_sections))

        if not packages:
            fallback_title = self._resolve_package_title(source_document, sections)
            return [ParsedPackage(title=fallback_title, sections=sections)]
        return packages

    def _resolve_package_title(
        self,
        source_document: models.SourceDocument,
        sections: List[ParsedSection],
    ) -> str:
        for section in sections:
            topic_title = self._extract_topic_title_from_heading(section.heading)
            if topic_title:
                return topic_title

        if source_document.title:
            normalized_title = source_document.title.strip()
            file_stem = Path(source_document.file_name or "knowledge_package").stem.strip()
            if normalized_title and normalized_title != file_stem:
                return normalized_title[:255]

        if sections:
            first_heading = sections[0].heading.strip()
            if first_heading and first_heading != "专题正文":
                return first_heading[:255]
        return Path(source_document.file_name or "knowledge_package").stem[:255]

    def _build_page_range(self, pages: List[Tuple[int, str]]) -> Optional[Dict[str, int]]:
        if not pages:
            return None
        return {"start": pages[0][0], "end": pages[-1][0]}

    def _build_page_range_from_sections(self, sections: List[ParsedSection]) -> Optional[Dict[str, int]]:
        if not sections:
            return None
        return {"start": sections[0].page_start, "end": sections[-1].page_end}


    def _build_outline(self, sections: List[ParsedSection]) -> List[Dict[str, object]]:
        return [
            {
                "order": index,
                "heading": section.heading,
                "page_start": section.page_start,
                "page_end": section.page_end,
            }
            for index, section in enumerate(sections, start=1)
        ]

    def _build_summary_text(self, sections: List[ParsedSection]) -> Optional[str]:
        if not sections:
            return None
        snippets: List[str] = []
        for section in sections[:5]:
            if section.text:
                snippets.append(section.text[:80])
        summary = "；".join(snippets)
        return summary[:1000] if summary else None

    def _normalize_text(self, text: str) -> str:
        return re.sub(r"\s+", " ", (text or "")).strip()

    def _extract_knowledge_point_names(self, heading: str, text: str) -> List[str]:
        names: List[str] = []

        if not self._extract_topic_title_from_heading(heading):
            heading_candidate = heading.strip()
            if self._is_valid_point_name(heading_candidate):
                names.append(heading_candidate)

        for match in re.findall(r"(?:知识点|考点)\s*[:：]\s*([^\n。；;]{2,80})", text):
            for part in re.split(r"[、,，/|；;]", match):
                candidate = part.strip()
                if self._is_valid_point_name(candidate):
                    names.append(candidate)

        deduped: List[str] = []
        seen = set()
        for item in names:
            normalized = item.strip()
            if normalized and normalized not in seen:
                seen.add(normalized)
                deduped.append(normalized)
        return deduped[:5]

    def _is_valid_point_name(self, value: str) -> bool:
        candidate = (value or "").strip()
        if not candidate:
            return False
        if len(candidate) < 2 or len(candidate) > 48:
            return False
        if candidate in {"专题正文", "未归类知识点"}:
            return False
        if _BARE_TOPIC_INDEX_LINE.fullmatch(candidate):
            return False
        if re.search(r"目录|第\d+页|图\d+", candidate):
            return False
        if self._extract_topic_title_from_heading(candidate):
            return False
        if re.fullmatch(r"[\d\s./,，:：\-]+", candidate):
            return False
        if re.search(r"\d+\s*/\s*\d+", candidate):
            return False
        if re.search(r"\.{3,}\s*\d+\s*$", candidate):
            return False
        # 必须含至少一个汉字，避免 PDF 碎片（如 40．182、0．025、希腊字母乱码）入库为知识点
        if not re.search(r"[\u4e00-\u9fff]", candidate):
            return False
        # 「2．【高考…」类整题题干标题不作为知识点名
        if re.match(r"^\d{1,2}[．.\u3001]", candidate) and _EXAM_PAPER_TAG_IN_NAME.search(candidate):
            return False
        return True


    def _package_append_llm_raw_debug(
        self,
        package: models.KnowledgePackage,
        *,
        batch_index: int,
        total_batches: int,
        model_name: Optional[str],
        response_text: Optional[str],
        parse_ok: bool,
        max_stored_chars: int = 150000,
    ) -> None:
        """将每批 LLM 原始返回写入 package.outline_json，便于验收与排错（仅块级知识点 LLM 路径调用）。"""
        oj = package.outline_json
        if isinstance(oj, dict):
            outline = dict(oj)
        elif isinstance(oj, list):
            outline = {"outline_items": list(oj)}
        else:
            outline = {}
        dbg = outline.get("llm_ingest_debug")
        if not isinstance(dbg, dict):
            dbg = {}
        raw_list = dbg.get("raw_responses")
        if not isinstance(raw_list, list):
            raw_list = []
        text = response_text if response_text is not None else ""
        if len(text) > max_stored_chars:
            text = text[:max_stored_chars] + "\n...[truncated for storage]"
        raw_list.append(
            {
                "batch_index": batch_index,
                "total_batches": total_batches,
                "model_name": model_name,
                "parse_ok": parse_ok,
                "response_text": text,
            }
        )
        dbg["raw_responses"] = raw_list
        outline["llm_ingest_debug"] = dbg
        package.outline_json = outline

    def _package_reset_question_bridge_llm_debug(
        self,
        package: models.KnowledgePackage,
        *,
        rank_mode: str,
    ) -> None:
        """新一轮桥接写入前清空 question_bridge_llm_debug，避免重复摄入时与旧记录混杂。"""
        oj = package.outline_json
        if isinstance(oj, dict):
            outline = dict(oj)
        elif isinstance(oj, list):
            outline = {"outline_items": list(oj)}
        else:
            outline = {}
        outline["question_bridge_llm_debug"] = {
            "step_key": TOPIC_DOCX_QUESTION_BRIDGE_STEP_KEY,
            "rank_mode": rank_mode,
            "raw_responses": [],
        }
        package.outline_json = outline

    def _package_append_question_bridge_llm_debug_entry(
        self,
        package: models.KnowledgePackage,
        *,
        question_item_id: int,
        model_name: Optional[str],
        response_text: Optional[str],
        parse_ok: bool,
        accepted_point_ids: List[int],
        skipped_reason: Optional[str] = None,
        max_stored_chars: int = 80000,
    ) -> None:
        """追加一条按题桥接 LLM 审计记录到 outline_json.question_bridge_llm_debug.raw_responses。"""
        oj = package.outline_json
        if isinstance(oj, dict):
            outline = dict(oj)
        elif isinstance(oj, list):
            outline = {"outline_items": list(oj)}
        else:
            outline = {}
        dbg = outline.get("question_bridge_llm_debug")
        if not isinstance(dbg, dict):
            dbg = {
                "step_key": TOPIC_DOCX_QUESTION_BRIDGE_STEP_KEY,
                "rank_mode": "",
                "raw_responses": [],
            }
        raw_list = dbg.get("raw_responses")
        if not isinstance(raw_list, list):
            raw_list = []
        text = response_text if response_text is not None else ""
        if len(text) > max_stored_chars:
            text = text[:max_stored_chars] + "\n...[truncated for storage]"
        raw_list.append(
            {
                "question_item_id": int(question_item_id),
                "model_name": model_name,
                "parse_ok": bool(parse_ok),
                "response_text": text,
                "accepted_knowledge_point_ids": [int(x) for x in accepted_point_ids],
                "skipped_reason": skipped_reason,
            }
        )
        dbg["raw_responses"] = raw_list
        outline["question_bridge_llm_debug"] = dbg
        package.outline_json = outline

    def _parse_llm_json_object_response(self, raw: Optional[str]) -> Optional[Dict[str, Any]]:
        if not raw or not str(raw).strip():
            return None
        text = str(raw).strip()
        if text.startswith("```"):
            text = re.sub(r"^```[a-zA-Z]*\s*", "", text)
            text = re.sub(r"\s*```\s*$", "", text).strip()
        try:
            out = json.loads(text)
        except json.JSONDecodeError:
            return None
        return out if isinstance(out, dict) else None

    def _topic_docx_llm_segment_batches(
        self,
        segment_dicts: List[Dict[str, Any]],
        max_chars: int = _TOPIC_DOCX_LLM_BATCH_JSON_CHARS,
    ) -> List[List[Dict[str, Any]]]:
        batches: List[List[Dict[str, Any]]] = []
        current: List[Dict[str, Any]] = []
        size = 0
        for item in segment_dicts:
            piece = json.dumps(item, ensure_ascii=False)
            step = len(piece) + 2
            if current and size + step > max_chars:
                batches.append(current)
                current = []
                size = 0
            current.append(item)
            size += step
        if current:
            batches.append(current)
        return batches

    def _refresh_package_point_links_map(
        self, db: Session, package_id: int
    ) -> Dict[int, models.KnowledgePackagePoint]:
        rows = (
            db.query(models.KnowledgePackagePoint)
            .filter(models.KnowledgePackagePoint.package_id == package_id)
            .all()
        )
        return {row.knowledge_point_id: row for row in rows}

    def _enrich_topic_docx_with_llm_block_points(
        self,
        db: Session,
        package: models.KnowledgePackage,
        content_segments: List[TopicContentSegment],
        source_document: models.SourceDocument,
        package_point_links: Dict[int, models.KnowledgePackagePoint],
        progress_callback: Optional[Callable[[str], None]] = None,
        *,
        topic_point_mode: str = "both",
    ) -> None:
        notify = lambda msg: self._notify_progress(progress_callback, msg)
        subject = source_document.subject
        grade = source_document.grade
        topic_key = re.sub(r"\s+", "", package.package_title or "")

        try:
            sync_llm_step_configs(db)
            sync_prompt_step_configs(db)
        except Exception as exc:
            logger.warning("topic_docx LLM sync step configs failed: %s", exc)
            notify(f"块级知识点 LLM：同步步骤配置失败，跳过。{exc}")
            return

        llm_cfg = resolve_step_llm_config(
            db, TOPIC_DOCX_BLOCK_POINTS_STEP_KEY, allow_generic_fallback=True
        )
        if not llm_cfg or not llm_cfg.get("api_url") or not llm_cfg.get("api_key"):
            notify("块级知识点 LLM：未解析到模型配置，跳过。")
            return

        segment_payloads: List[Dict[str, Any]] = []
        for order, seg in enumerate(content_segments, start=1):
            norm = self._normalize_text(seg.plain_text or "")
            segment_payloads.append(
                {
                    "block_order": order,
                    "section_title": (seg.section_title or "")[:240],
                    "plain_text": norm[:_TOPIC_DOCX_LLM_MAX_BLOCK_TEXT],
                }
            )

        batches = self._topic_docx_llm_segment_batches(segment_payloads)
        if not batches:
            return

        max_order_row = (
            db.query(func.max(models.KnowledgePackagePoint.order_in_package))
            .filter(models.KnowledgePackagePoint.package_id == package.id)
            .scalar()
        )
        next_pkg_order = int(max_order_row or 0) + 1

        notify(
            f"块级知识点 LLM：已启用，共 {len(batches)} 批请求（step={TOPIC_DOCX_BLOCK_POINTS_STEP_KEY}）…"
        )
        if KNOWLEDGE_TOPIC_BLOCK_POINTS_MULTIMODAL and not supports_vision_model(llm_cfg.get("model_name")):
            notify(
                "块级知识点 LLM：已开启 KNOWLEDGE_TOPIC_BLOCK_POINTS_MULTIMODAL，"
                "但当前步骤模型名不被识别为视觉模型，多模态附件不会发送（仍按纯文本）。"
            )

        for batch_idx, batch in enumerate(batches, start=1):
            blocks_json = json.dumps({"blocks": batch}, ensure_ascii=False)
            try:
                prompt_cfg = resolve_step_prompt(
                    db,
                    TOPIC_DOCX_BLOCK_POINTS_STEP_KEY,
                    variables={"blocks_json": blocks_json},
                )
                if not prompt_cfg or not (prompt_cfg.get("prompt_text") or "").strip():
                    notify(f"块级知识点 LLM：第 {batch_idx}/{len(batches)} 批无提示词，中止后续批。")
                    break
                prompt_body = prompt_cfg["prompt_text"]
                user_content: Any = prompt_body
                mm_image_parts = 0
                if KNOWLEDGE_TOPIC_BLOCK_POINTS_MULTIMODAL and supports_vision_model(llm_cfg.get("model_name")):
                    mm_tail: List[Dict[str, Any]] = []
                    mm_image_parts = append_topic_batch_multimodal_image_parts(
                        mm_tail,
                        batch=batch,
                        content_segments=content_segments,
                        search_roots=[Path(QUESTION_BANK_ASSET_DIR)],
                        max_images=KNOWLEDGE_BLOCK_LLM_MAX_IMAGES_PER_CALL,
                        max_image_bytes=KNOWLEDGE_BLOCK_LLM_MAX_IMAGE_BYTES,
                    )
                    if mm_image_parts > 0:
                        user_content = openai_user_content_for_call(
                            prompt_text=prompt_body,
                            extra_parts=mm_tail,
                        )
                        notify(
                            f"块级知识点 LLM：第 {batch_idx}/{len(batches)} 批已附加多模态图片段={mm_image_parts}"
                        )
                if self._ingest_verbose_enabled():
                    req_messages = [{"role": "user", "content": user_content}]
                    safe_messages = (
                        redact_openai_messages_for_audit(req_messages)
                        if isinstance(user_content, list)
                        else req_messages
                    )
                    self._ingest_verbose_write(
                        f"llm/topic_docx_block_points/batch_{batch_idx:03d}_request.json",
                        json.dumps(
                            {
                                "step_key": TOPIC_DOCX_BLOCK_POINTS_STEP_KEY,
                                "model_name": llm_cfg.get("model_name"),
                                "batch_index": batch_idx,
                                "total_batches": len(batches),
                                "multimodal_image_parts": mm_image_parts,
                                "content_summary": summarize_multimodal_content_for_log(user_content),
                                "messages": safe_messages,
                            },
                            ensure_ascii=False,
                            indent=2,
                        ),
                    )
                    self._ingest_verbose_write(
                        f"llm/topic_docx_block_points/batch_{batch_idx:03d}_user_prompt.txt",
                        prompt_body,
                    )
                raw = call_llm(
                    [{"role": "user", "content": user_content}],
                    llm_cfg,
                    json_mode=True,
                )
                if self._ingest_verbose_enabled():
                    self._ingest_verbose_write(
                        f"llm/topic_docx_block_points/batch_{batch_idx:03d}_response_raw.txt",
                        raw if raw is not None else "(null response)",
                    )
                data = self._parse_llm_json_object_response(raw)
                parse_ok = bool(data and isinstance(data.get("blocks"), list))
                self._package_append_llm_raw_debug(
                    package,
                    batch_index=batch_idx,
                    total_batches=len(batches),
                    model_name=llm_cfg.get("model_name"),
                    response_text=raw,
                    parse_ok=parse_ok,
                )
                if not data:
                    notify(f"块级知识点 LLM：第 {batch_idx}/{len(batches)} 批 JSON 解析失败，已跳过。")
                    continue
                blocks_out = data.get("blocks")
                if not isinstance(blocks_out, list):
                    notify(f"块级知识点 LLM：第 {batch_idx}/{len(batches)} 批缺少 blocks 数组，已跳过。")
                    continue
                for row in blocks_out:
                    if not isinstance(row, dict):
                        continue
                    try:
                        block_order = int(row.get("block_order"))
                    except (TypeError, ValueError):
                        continue
                    names_raw = row.get("knowledge_point_names")
                    if names_raw is None:
                        names_raw = row.get("knowledge_points")
                    if not isinstance(names_raw, list):
                        continue
                    kb = (
                        db.query(models.KnowledgeBlock)
                        .filter(
                            models.KnowledgeBlock.package_id == package.id,
                            models.KnowledgeBlock.block_order == block_order,
                        )
                        .first()
                    )
                    if not kb:
                        continue
                    first_llm_for_block = topic_point_mode == "llm"
                    anchor = dict(kb.source_anchor_json or {})
                    merged_list: List[Dict[str, Any]] = []
                    seen_ids: set[int] = set()
                    old_llm = anchor.get("llm_knowledge_points")
                    if isinstance(old_llm, list):
                        for x in old_llm:
                            if not isinstance(x, dict):
                                continue
                            pid_val = x.get("knowledge_point_id")
                            if pid_val is None:
                                merged_list.append(x)
                                continue
                            try:
                                pid = int(pid_val)
                            except (TypeError, ValueError):
                                merged_list.append(x)
                                continue
                            if pid not in seen_ids:
                                seen_ids.add(pid)
                                merged_list.append(x)

                    for raw_name in names_raw[:_TOPIC_DOCX_LLM_MAX_NAMES_PER_BLOCK]:
                        name = str(raw_name).strip()
                        if len(name) > 48:
                            name = name[:48]
                        if not self._is_valid_point_name(name):
                            continue
                        if re.sub(r"\s+", "", name) == topic_key:
                            continue
                        pt = self._get_or_create_knowledge_point(
                            db,
                            canonical_name=name,
                            subject=subject,
                            grade_scope=grade,
                            source_origin="llm",
                        )
                        record_knowledge_point_provenance(
                            db,
                            knowledge_point_id=pt.id,
                            source_kind=SOURCE_KIND_KNOWLEDGE_BLOCK,
                            source_id=kb.id,
                            package_id=package.id,
                            origin_step="llm_topic_block_points",
                        )
                        if pt.id == kb.knowledge_point_id and topic_point_mode != "llm":
                            continue
                        if pt.id in seen_ids:
                            continue
                        seen_ids.add(pt.id)
                        merged_list.append(
                            {
                                "canonical_name": pt.canonical_name,
                                "knowledge_point_id": pt.id,
                                "confidence": float(_TOPIC_DOCX_LLM_CONFIDENCE),
                            }
                        )
                        if pt.id not in package_point_links:
                            use_llm_core = topic_point_mode == "llm" and first_llm_for_block
                            if use_llm_core:
                                first_llm_for_block = False
                                rel_type = "core"
                                weight_v = 1.0
                                conf_v = 0.85
                            else:
                                rel_type = "supplement"
                                weight_v = _TOPIC_DOCX_LLM_WEIGHT
                                conf_v = _TOPIC_DOCX_LLM_CONFIDENCE
                            link = models.KnowledgePackagePoint(
                                package_id=package.id,
                                knowledge_point_id=pt.id,
                                relation_type=rel_type,
                                weight_score=weight_v,
                                order_in_package=next_pkg_order,
                                source_origin="llm",
                                confidence=conf_v,
                                approved_status="pending",
                            )
                            db.add(link)
                            package_point_links[pt.id] = link
                            next_pkg_order += 1
                            if use_llm_core:
                                kb.knowledge_point_id = pt.id
                    anchor["llm_knowledge_points"] = merged_list
                    anchor["llm_meta"] = {
                        "step_key": TOPIC_DOCX_BLOCK_POINTS_STEP_KEY,
                        "model_name": llm_cfg.get("model_name"),
                        "batch_index": batch_idx,
                        "total_batches": len(batches),
                    }
                    kb.source_anchor_json = anchor
            except Exception as exc:
                logger.exception("topic_docx LLM batch %s failed", batch_idx)
                notify(
                    f"块级知识点 LLM：第 {batch_idx}/{len(batches)} 批异常已跳过："
                    f"{exc.__class__.__name__}: {exc}"
                )

        db.flush()

    def _enforce_topic_title_boundary(
        self,
        db: Session,
        package_id: int,
        package_title: str,
        fallback_point_id: int,
    ) -> None:
        matched_point_ids = [
            row[0]
            for row in db.query(models.KnowledgePoint.id)
            .join(models.KnowledgePackagePoint, models.KnowledgePackagePoint.knowledge_point_id == models.KnowledgePoint.id)
            .filter(models.KnowledgePackagePoint.package_id == package_id)
            .filter(models.KnowledgePoint.canonical_name == package_title)
            .all()
        ]
        if not matched_point_ids:
            return

        db.query(models.KnowledgeBlock).filter(
            models.KnowledgeBlock.package_id == package_id,
            models.KnowledgeBlock.knowledge_point_id.in_(matched_point_ids),
        ).update({models.KnowledgeBlock.knowledge_point_id: fallback_point_id}, synchronize_session=False)

        db.query(models.KnowledgeAtom).filter(
            models.KnowledgeAtom.package_id == package_id,
            models.KnowledgeAtom.knowledge_point_id.in_(matched_point_ids),
        ).update({models.KnowledgeAtom.knowledge_point_id: fallback_point_id}, synchronize_session=False)

        db.query(models.KnowledgePackagePoint).filter(
            models.KnowledgePackagePoint.package_id == package_id,
            models.KnowledgePackagePoint.knowledge_point_id.in_(matched_point_ids),
        ).delete(synchronize_session=False)

    def _infer_block_role(self, heading: str) -> str:
        for role, pattern in SECTION_ROLE_PATTERNS:
            if re.search(pattern, heading):
                return role
        return "explainer"


    def _split_sentences(self, text: str) -> List[str]:
        raw_parts = re.split(r"[。！？；\n]", text or "")
        return [part.strip() for part in raw_parts if part and part.strip()]

    def _get_or_create_knowledge_point(
        self,
        db: Session,
        canonical_name: str,
        subject: Optional[str],
        grade_scope: Optional[str],
        source_origin: str = "model",
    ) -> models.KnowledgePoint:
        query = db.query(models.KnowledgePoint).filter(models.KnowledgePoint.canonical_name == canonical_name)
        if subject:
            query = query.filter(models.KnowledgePoint.subject == subject)
        point = query.first()
        if point:
            return point

        point = models.KnowledgePoint(
            canonical_name=canonical_name,
            subject=subject,
            grade_scope=grade_scope,
            knowledge_type="concept",
            source_origin=source_origin,
            review_status="draft",
            is_active=True,
        )
        db.add(point)
        db.flush()
        return point

    def _clear_existing_package_artifacts(self, db: Session, source_document_id: int) -> None:
        package_ids = [
            row[0]
            for row in db.query(models.KnowledgePackage.id)
            .filter(models.KnowledgePackage.source_document_id == source_document_id)
            .all()
        ]
        if not package_ids:
            return

        from analyzer.app.question_bank_parser import QuestionBankIngestionService

        QuestionBankIngestionService().clear_topic_material_papers_for_document(db, source_document_id, package_ids)

        block_ids = [
            row[0]
            for row in db.query(models.KnowledgeBlock.id)
            .filter(models.KnowledgeBlock.package_id.in_(package_ids))
            .all()
        ]

        if block_ids:
            db.query(models.KnowledgePointProvenance).filter(
                models.KnowledgePointProvenance.source_kind == SOURCE_KIND_KNOWLEDGE_BLOCK,
                models.KnowledgePointProvenance.source_id.in_(block_ids),
            ).delete(synchronize_session=False)

        db.query(models.KnowledgeAtom).filter(models.KnowledgeAtom.package_id.in_(package_ids)).delete(synchronize_session=False)
        db.query(models.KnowledgePackagePoint).filter(models.KnowledgePackagePoint.package_id.in_(package_ids)).delete(synchronize_session=False)
        db.query(models.KnowledgeQuestionLink).filter(
            models.KnowledgeQuestionLink.explanation_block_id.in_(block_ids)
            | models.KnowledgeQuestionLink.commentary_block_id.in_(block_ids)
        ).delete(synchronize_session=False)
        db.query(models.KnowledgePointRelation).filter(models.KnowledgePointRelation.evidence_block_id.in_(block_ids)).delete(
            synchronize_session=False
        )
        db.query(models.KnowledgeBlock).filter(models.KnowledgeBlock.package_id.in_(package_ids)).delete(synchronize_session=False)
        db.query(models.KnowledgePackage).filter(models.KnowledgePackage.id.in_(package_ids)).delete(synchronize_session=False)

    def _ensure_local_content_source(self, db: Session) -> models.ContentSource:
        source = (
            db.query(models.ContentSource)
            .filter(models.ContentSource.source_name == "knowledge_points_local")
            .first()
        )
        if source:
            return source

        source = models.ContentSource(
            source_name="knowledge_points_local",
            source_type="knowledge_point_material",
            provider_name="local_filesystem",
            commercial_allowed=False,
            ai_processing_allowed=True,
            training_allowed=False,
            license_scope={"type": "local"},
            remark="Auto-created for knowledge point ingestion from analyzer/knowledge_points",
        )
        db.add(source)
        db.flush()
        return source

    def _guess_mime_type(self, suffix: str) -> Optional[str]:
        suffix = (suffix or "").lower()
        if suffix == ".pdf":
            return "application/pdf"
        if suffix == ".docx":
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if suffix == ".txt":
            return "text/plain"
        return None

    def _start_parse_job(self, db: Session, source_document_id: int) -> models.DocumentParseJob:
        job = models.DocumentParseJob(
            source_document_id=source_document_id,
            job_stage="knowledge_structure_parse",
            tool_name="knowledge_point_parser",
            model_name="heuristic",
            input_version="v1",
            status="running",
        )
        db.add(job)
        db.flush()
        return job

    def _finish_parse_job(
        self,
        db: Session,
        job: models.DocumentParseJob,
        status: str,
        output_location: Optional[str],
        metrics_json: Optional[Dict[str, object]] = None,
    ) -> None:
        job.status = status
        job.output_location = output_location
        job.metrics_json = metrics_json
        job.error_message = None
        job.ended_at = models.func.now()
        db.add(job)

    def _fail_parse_job(self, db: Session, job: models.DocumentParseJob, error_message: str) -> None:
        job.status = "failed"
        job.error_message = error_message[:5000]
        job.ended_at = models.func.now()
        db.add(job)
