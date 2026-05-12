import copy
import hashlib
import logging
import os
import re
import shutil
import subprocess
import tempfile
import threading
import time

from dataclasses import dataclass, field
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional, Sequence, Tuple
from uuid import NAMESPACE_URL, UUID, uuid5




import fitz
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image
from sqlalchemy.orm import Session, joinedload

try:
    import pythoncom
    from win32com.client import DispatchEx, gencache
except Exception:
    pythoncom = None
    DispatchEx = None
    gencache = None

from shared import models
from . import vector_db
from .config import (
    KNOWLEDGE_POINTS_DIR,
    NORMALIZED_DOCUMENTS_DIR,
    QUESTION_BANK_ASSET_DIR,
    QUESTION_BANK_DOCX_SANITIZE_BACKEND,
    QUESTION_BANK_DOCX_SANITIZE_ENABLED,
    QUESTION_BANK_DOCX_SANITIZE_MODE,
    QUESTION_BANK_PDF_MATH_CLIP_IMAGES,
)
from .question_bank_docx_sanitizer import sanitize_docx_document, logger as docx_sanitizer_logger
from .question_bank_rich_content import extract_docx_blocks, logger as rich_content_logger





logger = logging.getLogger(__name__)


class _FrontendProgressLogHandler(logging.Handler):
    def __init__(self, progress_callback: Optional[Callable[[str], None]]):
        super().__init__(level=logging.INFO)
        self._progress_callback = progress_callback
        self._thread_id = threading.get_ident()
        self.setFormatter(logging.Formatter("%(message)s"))

    def emit(self, record: logging.LogRecord) -> None:
        if not self._progress_callback or record.thread != self._thread_id:
            return
        try:
            self._progress_callback(f"[INGEST-{record.levelname}] {self.format(record)}")
        except Exception:
            pass


class _ScopedFileLogHandler(logging.FileHandler):
    def __init__(self, log_path: Path):
        log_path.parent.mkdir(parents=True, exist_ok=True)
        super().__init__(str(log_path), mode="a", encoding="utf-8")
        self._thread_id = threading.get_ident()
        self.setLevel(logging.INFO)
        self.setFormatter(logging.Formatter("[%(asctime)s] [%(levelname)s] %(name)s - %(message)s"))

    def emit(self, record: logging.LogRecord) -> None:
        if record.thread != self._thread_id:
            return
        super().emit(record)



# 第三支：题号后无句点直接接卷标，如「3 【高考…」；限定 lookahead 避免误匹配正文「5 【某词】」
QUESTION_HEADER_PATTERN = re.compile(
    r"(?m)^\s*(?:"
    r"第\s*(\d{1,3})\s*题|"
    r"(\d{1,3})[\.．、](?!\d)|"
    r"(\d{1,3})\s*【(?=\s*(?:高考|新课标|\d{4}\s*年))"
    r")\s*"
)
_LEADING_IMAGE_PLACEHOLDER_RE = re.compile(
    r"^(?:\[图片\]|公式图片|公式|formula|\[IMG\]|【图片】)+\s*"
)

FORMULA_HINT_PATTERN = re.compile(r"[=≤≥∑√^_+\-×÷]|\\frac|\\sqrt|\\int|\\sum")
OPTION_PATTERN = re.compile(r"(?:^|\n)\s*[A-HＡ-Ｈ][\.．、:：]\s*")
OPTION_BLOCK_PATTERN = re.compile(
    r"(?ms)(?:^|\n)\s*(?P<key>[A-HＡ-Ｈ])[\.．、:：]\s*(?P<content>.*?)(?=(?:\n\s*[A-HＡ-Ｈ][\.．、:：]\s*)|\Z)"
)
FIGURE_MARKER_PATTERN = re.compile(r"(图\s*\d+|图[一二三四五六七八九十]+|如图|见图|下图|上图|\[图片\]|【图片】|插图)")

_SUB_Q_NO_PATTERN = re.compile(r"^\s*\(\d{1,3}\)")
LATEX_FORMULA_PATTERN = re.compile(r"(\$[^$]+\$|\\(?:frac|sqrt|int|sum|sin|cos|tan|log|ln)[^\n]{0,120})")

SECTION_LABEL_ALIASES = {
    "参考答案": "answer",
    "答案": "answer",
    "答案解析": "analysis",
    "解析": "analysis",
    "分析": "analysis",
    "详解": "solution",
    "解答": "solution",
    "思路导引": "solution",
    "思路引导": "solution",
    "专家解读": "comment",

    "解法": "solution",
    "过程": "solution",
    "点评": "comment",
    "评注": "comment",
    "点拨": "comment",
    "点睛": "comment",
    "备注": "comment",
    "考点": "knowledge",
    "知识点": "knowledge",
    "专题": "topic",
    "题型": "topic",
    "主题": "topic",
}
SECTION_LABEL_PATTERN = re.compile(
    r"(?im)(?:^|\n)\s*(?:【|\[|\()?\s*(?P<label>"
    + "|".join(sorted((re.escape(label) for label in SECTION_LABEL_ALIASES), key=len, reverse=True))
    + r")\s*(?:】|\]|\))?\s*[:：]?\s*"
)

# 选项正文后误并入「答案/解析」或下一题题号（无换行 A. 时）
_OPTION_TAIL_BLEED = re.compile(
    r"(?:"
    r"(?:\n|^)\s*(?:"
    r"答案\s*[:：]|参考答案\s*[:：]|【答案】|【参考答案】|"
    r"解析\s*[:：]|【解析】|"
    r"详解\s*[:：]|解答\s*[:：]|"
    r"(?:第\s*\d{1,3}\s*题|\d{1,3}[\.．、](?!\d))"
    r")|"
    r"(?:[。.．\n])\s*\(\d{1,2}\)|"
    r"(?:\u3002|[。.．\n])\s*\uFF08\d{1,2}\uFF09"
    r")",
    re.IGNORECASE,
)

# 仅用于「题干+选项」截取：在首个正式答案区【…】前截断，避免【考点】【专题】等把选项切没
OPTION_SCOPE_BRACKET_LABELS = (
    "参考答案",
    "答案解析",
    "答案",
    "解析",
    "详解",
    "解答",
    "专家解读",
    "思路引导",
    "思路导引",
    "解法",
    "点评",
    "评注",
    "点拨",
    "点睛",
    "备注",
)
OPTION_SCOPE_BRACKET_PATTERN = re.compile(
    r"【\s*(?:"
    + "|".join(re.escape(x) for x in sorted(OPTION_SCOPE_BRACKET_LABELS, key=len, reverse=True))
    + r")\s*】",
    re.IGNORECASE,
)

INLINE_KNOWLEDGE_PATTERNS = [
    re.compile(r"(?:考点|知识点|考查内容)\s*[:：]\s*(?P<value>[^\n]+)", re.IGNORECASE),
    re.compile(r"本题考查\s*(?P<value>[^。；\n]+)", re.IGNORECASE),
]
INLINE_TOPIC_PATTERNS = [
    re.compile(r"(?:专题|题型|主题)\s*[:：]\s*(?P<value>[^\n]+)", re.IGNORECASE),
]
INLINE_COMMENT_PATTERNS = [
    re.compile(r"(?:点评|评注|点拨|备注)\s*[:：]\s*(?P<value>[^\n]+)", re.IGNORECASE),
]


class DocumentNormalizationError(RuntimeError):
    pass


@dataclass
class ExtractedOption:
    option_key: str
    option_text: str
    is_correct: Optional[bool] = None


@dataclass
class ExtractedFormula:
    source_text: str
    normalized_signature: str
    block_role: str
    option_key: Optional[str] = None
    source_type: str = "text"


@dataclass
class ExtractedAsset:
    asset_role: str
    storage_url: str
    file_hash: Optional[str] = None
    page_no: Optional[int] = None
    width: Optional[int] = None
    height: Optional[int] = None
    ocr_text: Optional[str] = None
    caption_text: Optional[str] = None


@dataclass
class ExtractedQuestion:
    question_no: str
    text: str
    question_type: str
    has_formula: bool
    stem_text: str
    # 原卷题号（如「2」）；入库 question_no 可能为顺序号以保证向量 id 唯一
    original_question_label: Optional[str] = None
    options: List[ExtractedOption] = field(default_factory=list)
    answer_text: Optional[str] = None
    analysis_text: Optional[str] = None
    solution_text: Optional[str] = None
    comment_text: Optional[str] = None
    knowledge_points: List[str] = field(default_factory=list)
    topics: List[str] = field(default_factory=list)
    formulas: List[ExtractedFormula] = field(default_factory=list)
    figure_markers: List[str] = field(default_factory=list)
    render_payloads: Dict[str, Dict[str, object]] = field(default_factory=dict)
    # 本题在整段 slice_text 中的 [start, end)，供富文本映射时限制 find 范围
    source_slice_span: Optional[Tuple[int, int]] = None


@dataclass
class NormalizedDocument:
    source_path: Path
    normalized_docx_path: Optional[Path] = None
    word_export_docx_path: Optional[Path] = None
    normalized_pdf_path: Optional[Path] = None



@dataclass
class ExtractedDocumentContent:
    text: str
    structured_questions: Optional[List[ExtractedQuestion]] = None


class QuestionBankIngestionService:

    def __init__(
        self,
        normalized_root: Optional[str] = None,
        asset_root: Optional[str] = None,
        docx_sanitize_enabled: Optional[bool] = None,
        docx_sanitize_mode: Optional[str] = None,
        docx_sanitize_backend: Optional[str] = None,
    ):
        self.normalized_root = Path(normalized_root or NORMALIZED_DOCUMENTS_DIR)
        self.asset_root = Path(asset_root or QUESTION_BANK_ASSET_DIR)
        self.docx_sanitize_enabled = QUESTION_BANK_DOCX_SANITIZE_ENABLED if docx_sanitize_enabled is None else bool(docx_sanitize_enabled)
        self.docx_sanitize_mode = str(docx_sanitize_mode or QUESTION_BANK_DOCX_SANITIZE_MODE).strip().lower()
        self.docx_sanitize_backend = str(docx_sanitize_backend or QUESTION_BANK_DOCX_SANITIZE_BACKEND).strip().lower()
        self.normalized_root.mkdir(parents=True, exist_ok=True)
        self.asset_root.mkdir(parents=True, exist_ok=True)


    def _describe_source_document(self, source_document: models.SourceDocument) -> str:
        return (
            f"id={source_document.id} "
            f"file_name={source_document.file_name} "
            f"file_ext={source_document.file_ext} "
            f"parse_status={source_document.parse_status} "
            f"storage_url={source_document.storage_url}"
        )

    def _format_elapsed(self, started_at: float) -> str:
        return f"{time.perf_counter() - started_at:.2f}s"

    def _build_ingest_log_path(self, source_document: models.SourceDocument) -> Path:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        return self.asset_root / "_ingest_logs" / f"document_{source_document.id}_{timestamp}.txt"

    def delete_source_document(self, db: Session, source_document_id: int) -> Dict[str, object]:

        source_document = (
            db.query(models.SourceDocument)
            .options(joinedload(models.SourceDocument.content_source))
            .filter(models.SourceDocument.id == source_document_id)
            .first()
        )
        if not source_document:
            raise ValueError(f"SourceDocument {source_document_id} 不存在")

        content_source = source_document.content_source
        is_knowledge_point_material = (
            (source_document.parse_profile or "").strip().lower() == "knowledge_point"
            or (
                content_source is not None
                and getattr(content_source, "source_name", None) == "knowledge_points_local"
            )
        )

        logger.info(
            "Deleting source document and all artifacts: %s knowledge_point_material=%s",
            self._describe_source_document(source_document),
            is_knowledge_point_material,
        )

        # 凡挂在本 source_document_id 下的 KnowledgePackage 一律先拆（无包时为空操作）。
        # 若仅按 parse_profile / 内容源名判断「专题资料」，漏判时会跳过清理；而 SourceDocument ORM 未级联专题包，
        # 在部分 SQLite 配置下仍可能删掉文档行却留下孤儿包，知识点管理页就会继续看到旧考点。
        from analyzer.app.knowledge_point_parser import KnowledgePointIngestionService

        KnowledgePointIngestionService()._clear_existing_package_artifacts(db, source_document_id)
        db.commit()

        self._clear_existing_document_artifacts(db, source_document_id)
        db.query(models.ExamSession).filter(
            models.ExamSession.source_document_id == source_document_id
        ).update(
            {models.ExamSession.source_document_id: None},
            synchronize_session=False,
        )
        db.commit()

        file_cleanup_summary = self._delete_local_files_for_source_document(source_document)
        db.delete(source_document)
        db.commit()

        ingest_run_dirs_removed: List[str] = []
        ingest_run_dirs_failed: List[str] = []
        if is_knowledge_point_material:
            from analyzer.app.knowledge_point_parser import delete_knowledge_ingest_run_folders_for_source_document

            ingest_run_dirs_removed, ingest_run_dirs_failed = delete_knowledge_ingest_run_folders_for_source_document(
                source_document_id
            )

        logger.info(
            "Source document deleted: source_document_id=%s removed_paths=%s failed_paths=%s ingest_run_dirs_removed=%s",
            source_document_id,
            file_cleanup_summary.get("removed_paths"),
            file_cleanup_summary.get("failed_paths"),
            ingest_run_dirs_removed,
        )
        return {
            "status": "deleted",
            "source_document_id": source_document_id,
            "knowledge_point_material": is_knowledge_point_material,
            "ingest_run_dirs_removed": ingest_run_dirs_removed,
            "ingest_run_dirs_failed": ingest_run_dirs_failed,
            **file_cleanup_summary,
        }

    def _delete_local_files_for_source_document(
        self,
        source_document: models.SourceDocument,
    ) -> Dict[str, object]:
        removed_paths: List[str] = []
        failed_paths: List[str] = []
        preserved_paths: List[str] = []

        # Build a set of paths that must NEVER be deleted (original source file).
        never_delete: set[str] = set()
        if source_document.storage_url:
            try:
                orig = self._resolve_local_path(source_document.storage_url).resolve(strict=False)
                never_delete.add(os.path.normcase(str(orig)))
            except Exception:
                pass

        knowledge_points_root_norm = os.path.normcase(
            str(Path(KNOWLEDGE_POINTS_DIR).resolve())
        )

        def _try_delete_path(path: Path) -> None:
            try:
                resolved_path = path.resolve(strict=False)
            except Exception:
                resolved_path = path
            resolved_norm = os.path.normcase(str(resolved_path))

            # 1) Never delete the original source file itself.
            if resolved_norm in never_delete:
                logger.info("Preserved original source file: %s", resolved_path)
                preserved_paths.append(str(resolved_path))
                return

            # 2) Never delete anything under the knowledge_points directory.
            try:
                if os.path.commonpath([knowledge_points_root_norm, resolved_norm]) == knowledge_points_root_norm:
                    logger.info("Preserved knowledge_points path: %s", resolved_path)
                    preserved_paths.append(str(resolved_path))
                    return
            except ValueError:
                pass

            if not path.exists():
                return
            try:
                if path.is_file():
                    path.unlink()
                    parent = path.parent
                    if parent.exists() and parent.is_dir() and not any(parent.iterdir()):
                        parent.rmdir()
                else:
                    shutil.rmtree(path, ignore_errors=False)
                removed_paths.append(str(path))
            except Exception as exc:
                logger.warning(
                    "Failed to delete path for source_document_id=%s path=%s error=%s",
                    source_document.id,
                    path,
                    exc,
                )
                failed_paths.append(str(path))

        logger.info(
            "File cleanup for source_document_id=%s: storage_url=%r normalized_docx_url=%r normalized_pdf_url=%r never_delete=%s",
            source_document.id,
            source_document.storage_url,
            source_document.normalized_docx_url,
            source_document.normalized_pdf_url,
            never_delete,
        )

        for url in {
            source_document.storage_url,
            source_document.normalized_docx_url,
            source_document.normalized_pdf_url,
        }:
            if not url:
                continue
            try:
                path = self._resolve_local_path(url)
            except Exception:
                continue
            _try_delete_path(path)

        _try_delete_path(self.normalized_root / f"document_{source_document.id}")
        _try_delete_path(self.asset_root / f"document_{source_document.id}")
        return {
            "removed_paths": removed_paths,
            "failed_paths": failed_paths,
            "preserved_paths": preserved_paths,
        }

    def ingest_source_document(

        self,
        db: Session,
        source_document_id: int,
        force_reingest: bool = False,
        progress_callback: Optional[Callable[[str], None]] = None,
    ) -> Dict[str, object]:
        source_document = db.query(models.SourceDocument).filter(models.SourceDocument.id == source_document_id).first()
        if not source_document:
            raise ValueError(f"SourceDocument {source_document_id} 不存在")

        progress_log_handler = _FrontendProgressLogHandler(progress_callback) if progress_callback else None
        scoped_log_path = self._build_ingest_log_path(source_document)
        scoped_file_log_handler = _ScopedFileLogHandler(scoped_log_path)
        progress_loggers = [logger, vector_db.logger, rich_content_logger, docx_sanitizer_logger]


        original_levels = {}
        for progress_logger in progress_loggers:
            original_levels[progress_logger] = progress_logger.level
            progress_logger.setLevel(logging.INFO)
            progress_logger.addHandler(scoped_file_log_handler)
            if progress_log_handler:
                progress_logger.addHandler(progress_log_handler)

        try:

            total_started_at = time.perf_counter()
            logger.info(
                "Question bank ingest scoped log file: source_document_id=%s log_path=%s",
                source_document_id,
                scoped_log_path,
            )

            logger.info(
                "Question bank ingest start: %s force_reingest=%s",
                self._describe_source_document(source_document),
                force_reingest,
            )

            if source_document.parse_status == "success" and not force_reingest:
                logger.info("Question bank ingest skipped: source_document_id=%s reason=already_success", source_document_id)
                return {
                    "status": "skipped",
                    "source_document_id": source_document_id,
                    "reason": "document already ingested",
                    "ingest_log_path": str(scoped_log_path),
                }


            if force_reingest or source_document.parse_status in {"failed", "running"}:
                logger.info(
                    "Question bank ingest clearing existing artifacts: source_document_id=%s previous_status=%s",
                    source_document_id,
                    source_document.parse_status,
                )
                self._clear_existing_document_artifacts(db, source_document_id)

            source_document.parse_status = "running"
            db.commit()

            normalize_job = self._start_job(db, source_document_id, "normalize", tool_name="filesystem")
            extract_job = None
            asset_job = None
            segment_job = None
            index_job = None

            try:
                stage_started_at = time.perf_counter()
                logger.info("Normalize stage start: source_document_id=%s", source_document_id)
                normalized = self.normalize_document(source_document)
                source_document.normalized_docx_url = self._path_to_storage_url(normalized.normalized_docx_path)
                source_document.normalized_pdf_url = self._path_to_storage_url(normalized.normalized_pdf_path)
                self._finish_job(db, normalize_job, output_location=source_document.normalized_docx_url or source_document.normalized_pdf_url)
                logger.info(
                    "Normalize stage done: source_document_id=%s elapsed=%s normalized_docx=%s normalized_pdf=%s",
                    source_document_id,
                    self._format_elapsed(stage_started_at),
                    source_document.normalized_docx_url,
                    source_document.normalized_pdf_url,
                )

                extract_job = self._start_job(db, source_document_id, "extract", tool_name="pymupdf/python-docx")
                stage_started_at = time.perf_counter()
                logger.info("Extract stage start: source_document_id=%s", source_document_id)
                extracted_content = self.extract_document_content(source_document, normalized)
                extracted_text = extracted_content.text
                if not extracted_text.strip():
                    raise ValueError("文档未提取到有效文本，无法切题")
                self._finish_job(db, extract_job, metrics_json={"text_length": len(extracted_text)})
                logger.info(
                    "Extract stage done: source_document_id=%s elapsed=%s text_length=%s structured_question_count=%s",
                    source_document_id,
                    self._format_elapsed(stage_started_at),
                    len(extracted_text),
                    len(extracted_content.structured_questions or []),
                )

                asset_job = self._start_job(db, source_document_id, "assets", tool_name="pymupdf/python-docx")
                stage_started_at = time.perf_counter()
                logger.info("Asset stage start: source_document_id=%s", source_document_id)
                extracted_assets = self.extract_document_assets(source_document, normalized)
                document_asset_count = self.persist_document_assets(db, source_document, extracted_assets)
                self._finish_job(db, asset_job, metrics_json={"asset_count": document_asset_count})
                logger.info(
                    "Asset stage done: source_document_id=%s elapsed=%s extracted_assets=%s persisted_assets=%s",
                    source_document_id,
                    self._format_elapsed(stage_started_at),
                    len(extracted_assets),
                    document_asset_count,
                )

                segment_job = self._start_job(db, source_document_id, "segment", tool_name="rule_based_segmenter")
                stage_started_at = time.perf_counter()
                logger.info("Segment stage start: source_document_id=%s", source_document_id)
                paper = self._upsert_paper(db, source_document)
                extracted_questions = extracted_content.structured_questions or self.segment_questions(extracted_text)
                question_metrics = self.persist_questions(
                    db=db,
                    paper=paper,
                    source_document=source_document,
                    extracted_text=extracted_text,
                    extracted_questions=extracted_questions,
                    document_asset_count=document_asset_count,
                )
                self._finish_job(db, segment_job, metrics_json=question_metrics)
                logger.info(
                    "Segment stage done: source_document_id=%s elapsed=%s paper_id=%s question_metrics=%s",
                    source_document_id,
                    self._format_elapsed(stage_started_at),
                    paper.id,
                    question_metrics,
                )

                index_job = self._start_job(
                    db,
                    source_document_id,
                    "index",
                    tool_name=vector_db.db.index_backend_label,
                )
                stage_started_at = time.perf_counter()
                logger.info("Index stage start: source_document_id=%s backend=%s", source_document_id, vector_db.db.index_backend_label)
                index_metrics = self.index_document_questions(db, source_document_id)
                self._finish_job(db, index_job, metrics_json=index_metrics)
                logger.info(
                    "Index stage done: source_document_id=%s elapsed=%s index_metrics=%s",
                    source_document_id,
                    self._format_elapsed(stage_started_at),
                    index_metrics,
                )

                source_document.parse_status = "success"
                db.commit()
                logger.info(
                    "Question bank ingest success: source_document_id=%s total_elapsed=%s paper_id=%s question_count=%s",
                    source_document_id,
                    self._format_elapsed(total_started_at),
                    paper.id,
                    question_metrics["question_count"],
                )
                return {
                    "status": "success",
                    "source_document_id": source_document_id,
                    "paper_id": paper.id,
                    "question_count": question_metrics["question_count"],
                    "formula_count": question_metrics["formula_count"],
                    "question_asset_count": question_metrics["question_asset_count"],
                    "document_asset_count": document_asset_count,
                    "indexed_documents": index_metrics.get("indexed_documents", 0),
                    "vector_backend": index_metrics.get("vector_backend"),
                    "text_backend": index_metrics.get("text_backend"),
                    "ingest_log_path": str(scoped_log_path),
                }


            except Exception as exc:
                source_document.parse_status = "failed"
                db.commit()
                for job in [normalize_job, extract_job, asset_job, segment_job, index_job]:
                    if job and job.status not in {"success", "failed"}:
                        self._fail_job(db, job, str(exc))
                logger.exception(
                    "Question bank ingest failed: source_document_id=%s total_elapsed=%s",
                    source_document_id,
                    self._format_elapsed(total_started_at),
                )
                return {
                    "status": "failed",
                    "source_document_id": source_document_id,
                    "error": str(exc),
                    "ingest_log_path": str(scoped_log_path),
                }

        finally:
            for progress_logger in progress_loggers:
                if progress_log_handler:
                    progress_logger.removeHandler(progress_log_handler)
                progress_logger.removeHandler(scoped_file_log_handler)
                if progress_logger in original_levels:
                    progress_logger.setLevel(original_levels[progress_logger])
            if progress_log_handler:
                progress_log_handler.close()
            scoped_file_log_handler.close()





    def normalize_document(self, source_document: models.SourceDocument) -> NormalizedDocument:
        source_path = self._resolve_local_path(source_document.storage_url)
        if not source_path.exists():
            raise FileNotFoundError(f"源文件不存在: {source_path}")

        file_ext = self._normalize_extension(source_document.file_ext or source_path.suffix)
        output_dir = self.normalized_root / f"document_{source_document.id}"
        output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(
            "Normalize document routing: source_document_id=%s source_path=%s file_ext=%s output_dir=%s",
            source_document.id,
            source_path,
            file_ext,
            output_dir,
        )

        if file_ext == ".doc":
            logger.info("Normalize document plan: source_document_id=%s .doc -> LibreOffice docx(required) + pdf(optional)", source_document.id)
            word_export_docx_path = self._convert_with_libreoffice(source_path, output_dir, "docx")
            normalized_pdf_path = self._convert_with_libreoffice(source_path, output_dir, "pdf", optional=True)
            if not word_export_docx_path:
                raise DocumentNormalizationError(
                    "当前环境未检测到 LibreOffice，无法自动将 .doc 归一化为 .docx。"
                )
            normalized_docx_path = self._maybe_sanitize_docx(source_document.id, file_ext, word_export_docx_path, output_dir)

            return NormalizedDocument(
                source_path=source_path,
                normalized_docx_path=normalized_docx_path,
                word_export_docx_path=word_export_docx_path,
                normalized_pdf_path=normalized_pdf_path,
            )

        if file_ext == ".docx":
            logger.info("Normalize document plan: source_document_id=%s .docx keeps original docx and tries optional LibreOffice pdf conversion", source_document.id)
            word_export_docx_path = source_path
            normalized_docx_path = self._maybe_sanitize_docx(source_document.id, file_ext, word_export_docx_path, output_dir)

            normalized_pdf_path = self._convert_with_libreoffice(source_path, output_dir, "pdf", optional=True)
            return NormalizedDocument(
                source_path=source_path,
                normalized_docx_path=normalized_docx_path,
                word_export_docx_path=word_export_docx_path,
                normalized_pdf_path=normalized_pdf_path,
            )



        if file_ext == ".pdf":
            logger.info("Normalize document plan: source_document_id=%s .pdf direct pass-through", source_document.id)
            return NormalizedDocument(source_path=source_path, normalized_pdf_path=source_path)

        if file_ext == ".txt":
            logger.info("Normalize document plan: source_document_id=%s .txt direct pass-through", source_document.id)
            return NormalizedDocument(source_path=source_path)

        raise ValueError(f"暂不支持的文件类型: {file_ext}")

    def _resolve_docx_sanitize_backend(self, source_file_ext: str) -> str:
        backend = self.docx_sanitize_backend
        if backend == "auto" and source_file_ext == ".doc":
            return "xml"
        return backend

    def _maybe_sanitize_docx(self, source_document_id: int, source_file_ext: str, docx_path: Optional[Path], output_dir: Path) -> Optional[Path]:
        if not docx_path:
            return docx_path
        if not self.docx_sanitize_enabled:
            logger.info(
                "Docx sanitize layer disabled: source_document_id=%s source_file_ext=%s docx_path=%s",
                source_document_id,
                source_file_ext,
                docx_path,
            )
            return docx_path

        effective_backend = self._resolve_docx_sanitize_backend(source_file_ext)
        if effective_backend != self.docx_sanitize_backend:
            logger.info(
                "Docx sanitize backend override: source_document_id=%s source_file_ext=%s configured_backend=%s effective_backend=%s reason=avoid_word_resave_side_effect_on_libreoffice_docx",
                source_document_id,
                source_file_ext,
                self.docx_sanitize_backend,
                effective_backend,
            )

        sanitized_docx_path = output_dir / "sanitized.docx"
        result = sanitize_docx_document(
            docx_path,
            sanitized_docx_path,
            enabled=self.docx_sanitize_enabled,
            mode=self.docx_sanitize_mode,
            backend=effective_backend,
        )
        final_docx_path = result.output_path if result.output_path.exists() else docx_path
        logger.info(
            "Docx sanitize layer finished: source_document_id=%s source_file_ext=%s input_path=%s output_path=%s final_path=%s enabled=%s applied=%s mode=%s configured_backend=%s effective_backend=%s result_backend=%s removed_total=%s reason_counts=%s error=%s",
            source_document_id,
            source_file_ext,
            docx_path,
            sanitized_docx_path,
            final_docx_path,
            result.enabled,
            result.applied,
            result.mode,
            self.docx_sanitize_backend,
            effective_backend,
            result.backend,
            result.removed_total,
            result.reason_counts,
            result.error,
        )
        return final_docx_path


    def extract_document_content(

        self,
        source_document: models.SourceDocument,
        normalized_document: NormalizedDocument,
    ) -> ExtractedDocumentContent:
        if normalized_document.normalized_docx_path:
            output_dir = self.asset_root / f"document_{source_document.id}"
            output_dir.mkdir(parents=True, exist_ok=True)
            word_export_docx_path = normalized_document.word_export_docx_path or normalized_document.normalized_docx_path
            logger.info(
                "Extract content via docx blocks: source_document_id=%s docx_path=%s word_export_docx_path=%s output_dir=%s",
                source_document.id,
                normalized_document.normalized_docx_path,
                word_export_docx_path,
                output_dir,
            )
            structured_blocks = extract_docx_blocks(
                normalized_document.normalized_docx_path,
                output_dir,
                word_export_docx_path=word_export_docx_path,
            )
            text = self._normalize_text("\n".join(block.get("text") or "" for block in structured_blocks))
            structured_questions = self._segment_structured_questions(structured_blocks)

            logger.info(
                "Extract content via docx blocks done: source_document_id=%s block_count=%s structured_question_count=%s text_length=%s",
                source_document.id,
                len(structured_blocks),
                len(structured_questions),
                len(text),
            )
            return ExtractedDocumentContent(text=text, structured_questions=structured_questions or None)
        logger.info(
            "Extract content via plain text fallback: source_document_id=%s pdf_path=%s source_path=%s",
            source_document.id,
            normalized_document.normalized_pdf_path,
            normalized_document.source_path,
        )
        text = self.extract_text(normalized_document)
        logger.info("Extract content via plain text fallback done: source_document_id=%s text_length=%s", source_document.id, len(text))
        return ExtractedDocumentContent(text=text)


    def extract_text(self, normalized_document: NormalizedDocument) -> str:
        if normalized_document.normalized_docx_path:
            return self._extract_text_from_docx(normalized_document.normalized_docx_path)
        if normalized_document.normalized_pdf_path:
            return self._extract_text_from_pdf(normalized_document.normalized_pdf_path)
        return self._extract_text_from_txt(normalized_document.source_path)


    def extract_document_assets(
        self,
        source_document: models.SourceDocument,
        normalized_document: NormalizedDocument,
    ) -> List[ExtractedAsset]:
        output_dir = self.asset_root / f"document_{source_document.id}"
        output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(
            "Extract assets start: source_document_id=%s docx_path=%s pdf_path=%s output_dir=%s",
            source_document.id,
            normalized_document.normalized_docx_path,
            normalized_document.normalized_pdf_path,
            output_dir,
        )

        try:
            if normalized_document.normalized_docx_path:
                assets = self._extract_docx_assets(normalized_document.normalized_docx_path, output_dir)
                logger.info("Extract assets done via docx: source_document_id=%s asset_count=%s", source_document.id, len(assets))
                return assets
            if normalized_document.normalized_pdf_path:
                assets = self._extract_pdf_assets(normalized_document.normalized_pdf_path, output_dir)
                logger.info("Extract assets done via pdf: source_document_id=%s asset_count=%s", source_document.id, len(assets))
                return assets
        except Exception as exc:
            logger.warning("Extracting document assets failed for source_document=%s: %s", source_document.id, exc)
        logger.info("Extract assets skipped/no assets: source_document_id=%s", source_document.id)
        return []


    def _extract_question_no_from_match(self, match: re.Match[str]) -> Optional[int]:
        raw_value = match.group(1) or match.group(2) or match.group(3)
        if not raw_value:
            return None
        try:
            return int(raw_value)
        except (TypeError, ValueError):
            return None

    def _is_reasonable_question_transition(self, previous_no: Optional[int], current_no: Optional[int]) -> bool:
        if current_no is None:
            return False
        if previous_no is None:
            return True
        return current_no == 1 or current_no == previous_no + 1

    def segment_questions(self, text: str) -> List[ExtractedQuestion]:
        matches = list(QUESTION_HEADER_PATTERN.finditer(text))
        if not matches:
            normalized_text = self._normalize_text(text)
            if not normalized_text:
                return []
            return [
                self._parse_question_segment(
                    question_no="1",
                    question_text=normalized_text,
                    source_slice_span=(0, len(text)),
                ),
            ]

        valid_matches: List[re.Match[str]] = []
        last_question_no: Optional[int] = None
        for match in matches:
            candidate_no = self._extract_question_no_from_match(match)
            if self._is_reasonable_question_transition(last_question_no, candidate_no):
                valid_matches.append(match)
                last_question_no = candidate_no

        if not valid_matches:
            normalized_text = self._normalize_text(text)
            if not normalized_text:
                return []
            return [
                self._parse_question_segment(
                    question_no="1",
                    question_text=normalized_text,
                    source_slice_span=(0, len(text)),
                ),
            ]

        extracted_questions: List[ExtractedQuestion] = []
        for index, match in enumerate(valid_matches):
            question_no = match.group(1) or match.group(2) or match.group(3) or str(index + 1)
            segment_start = match.end()
            segment_end = valid_matches[index + 1].start() if index + 1 < len(valid_matches) else len(text)
            question_text = self._normalize_text(text[segment_start:segment_end])
            if not question_text:
                continue
            extracted_questions.append(
                self._parse_question_segment(
                    question_no=question_no,
                    question_text=question_text,
                    source_slice_span=(segment_start, segment_end),
                ),
            )
        return extracted_questions


    def persist_document_assets(
        self,
        db: Session,
        source_document: models.SourceDocument,
        extracted_assets: List[ExtractedAsset],
    ) -> int:
        db.query(models.Asset).filter(models.Asset.owner_type == "source_document").filter(
            models.Asset.owner_id == source_document.id
        ).delete(synchronize_session=False)

        for asset in extracted_assets:
            db.add(
                models.Asset(
                    tenant_id=source_document.tenant_id,
                    owner_type="source_document",
                    owner_id=source_document.id,
                    asset_role=asset.asset_role,
                    storage_url=asset.storage_url,
                    page_no=asset.page_no,
                    width=asset.width,
                    height=asset.height,
                    ocr_text=asset.ocr_text,
                    caption_text=asset.caption_text,
                    file_hash=asset.file_hash,
                )
            )

        db.commit()
        return len(extracted_assets)

    @staticmethod
    def _attach_rich_payloads_to_questions(
        extracted_questions: List["ExtractedQuestion"],
        slice_text: str,
        pages: List[Tuple[int, str]],
        page_rich_by_no: Dict[int, List[Dict[str, Any]]],
        start_p: int,
        end_p: int,
    ) -> None:
        """为切题后的 ExtractedQuestion 附加 rich_content_json render_payloads。

        原理：先将本 slice 涉及页的 rich_paragraphs 按在 slice_text 中的字符偏移
        排列，然后根据每道题 stem_text / answer_text 等在 slice_text 中的位置，
        找出对应的 rich paragraph 子集，组装成 block_group 写入 render_payloads。
        """
        from .pdf_structured_extractor import build_rich_content_json

        para_offsets: List[Tuple[int, int, Dict[str, Any]]] = []
        offset = 0
        for pn, txt in pages:
            if not (start_p <= pn <= end_p):
                continue
            rich_paras = page_rich_by_no.get(pn, [])
            page_text = txt
            line_cursor = 0
            for para in rich_paras:
                children = para.get("children", [])
                para_plain = "".join(c.get("text", "") for c in children)
                idx = page_text.find(para_plain, line_cursor)
                if idx >= 0:
                    abs_start = offset + idx
                    abs_end = abs_start + len(para_plain)
                    para_offsets.append((abs_start, abs_end, para))
                    line_cursor = idx + len(para_plain)
                else:
                    para_offsets.append((offset + line_cursor, offset + line_cursor + len(para_plain), para))
                    line_cursor += len(para_plain)
            offset += len(page_text) + 1

        for eq in extracted_questions:
            if eq.render_payloads:
                continue

            span = eq.source_slice_span or (0, len(slice_text))
            lo, hi = span[0], span[1]

            def _find_paragraphs(text_segment: Optional[str], role: str) -> Optional[Dict[str, Any]]:
                if not text_segment or not para_offsets:
                    return None
                loc = QuestionBankIngestionService._locate_field_span_for_rich(
                    slice_text, text_segment, role, lo, hi,
                )
                if not loc:
                    return None
                seg_start, seg_end = loc
                matched = [p for (ps, pe, p) in para_offsets if pe > seg_start and ps < seg_end]
                if not matched:
                    return None
                return build_rich_content_json(matched, text_segment)

            payloads: Dict[str, Any] = {}
            stem_rich = _find_paragraphs(eq.stem_text or eq.text, "stem")
            if stem_rich:
                payloads["stem"] = stem_rich
            for role in ("answer", "analysis", "solution", "comment"):
                tv = getattr(eq, f"{role}_text", None)
                if tv:
                    r = _find_paragraphs(tv, role)
                    if r:
                        payloads[role] = r
            if payloads:
                eq.render_payloads = payloads

    def persist_questions(
        self,
        db: Session,
        paper: models.Paper,
        source_document: models.SourceDocument,
        extracted_text: str,
        extracted_questions: List[ExtractedQuestion],
        document_asset_count: int = 0,
        knowledge_package_id: Optional[int] = None,
    ) -> Dict[str, object]:
        question_count = 0
        formula_count = 0
        question_asset_count = 0
        question_item_ids: List[int] = []

        paper.total_questions = len(extracted_questions)
        raw_outline = {
            "text_length": len(extracted_text),
            "question_count": len(extracted_questions),
            "document_asset_count": document_asset_count,
        }
        paper.raw_outline_json = raw_outline


        for display_order, extracted_question in enumerate(extracted_questions, start=1):
            stem_seed_text = extracted_question.stem_text or extracted_question.text
            stem_hash = hashlib.sha256(stem_seed_text.encode("utf-8")).hexdigest()
            question_item = models.QuestionItem(
                tenant_id=source_document.tenant_id,
                subject=source_document.subject,
                grade=source_document.grade,
                question_type=extracted_question.question_type,
                stem_plain_text=extracted_question.stem_text or extracted_question.text,
                stem_normalized_text=self._normalize_text(extracted_question.stem_text or extracted_question.text),
                answer_text=extracted_question.answer_text,
                solution_summary=self._build_solution_summary(extracted_question),
                has_formula=bool(extracted_question.formulas),
                has_figure=bool(extracted_question.figure_markers),
                canonical_hash=stem_hash,
                source_origin="explicit",
                review_status="draft",
            )
            db.add(question_item)
            db.flush()
            question_item_ids.append(question_item.id)

            question_asset_increment, render_asset_lookup = self._persist_question_assets(
                db=db,
                source_document=source_document,
                question_item=question_item,
                extracted_question=extracted_question,
            )
            question_asset_count += question_asset_increment
            render_payloads = self._attach_asset_ids_to_render_payloads(
                extracted_question.render_payloads,
                render_asset_lookup,
            )

            role_blocks: Dict[Tuple[str, Optional[str]], models.QuestionBlock] = {}
            option_rows: Dict[str, models.QuestionOption] = {}
            block_order = 1

            stem_payload = render_payloads.get("stem")
            stem_block, block_order = self._append_question_block(
                db=db,
                question_item_id=question_item.id,
                block_order=block_order,
                block_role="stem",
                text_content=extracted_question.stem_text or extracted_question.text,
                rich_content_json=stem_payload,
                content_format="json" if stem_payload else "plain_text",
                is_primary=True,
            )
            role_blocks[("stem", None)] = stem_block

            options_payload = render_payloads.get("options")
            if options_payload:
                options_block, block_order = self._append_question_block(
                    db=db,
                    question_item_id=question_item.id,
                    block_order=block_order,
                    block_role="options",
                    text_content=options_payload.get("plain_text") or "",
                    rich_content_json=options_payload,
                    content_format="json",
                    parent_block_id=stem_block.id,
                    is_primary=True,
                )
                role_blocks[("options", None)] = options_block

            for option_index, option in enumerate(extracted_question.options, start=1):

                option_row = models.QuestionOption(
                    question_item_id=question_item.id,
                    option_key=option.option_key,
                    option_text=option.option_text,
                    display_order=option_index,
                    is_correct=option.is_correct,
                )
                db.add(option_row)
                db.flush()
                option_rows[option.option_key] = option_row

                option_payload = render_payloads.get(f"option:{option.option_key}") or {"option_key": option.option_key}
                option_block, block_order = self._append_question_block(
                    db=db,
                    question_item_id=question_item.id,
                    block_order=block_order,
                    block_role="option",
                    text_content=option.option_text,
                    rich_content_json=option_payload,
                    content_format="json" if render_payloads.get(f"option:{option.option_key}") else "plain_text",
                    parent_block_id=stem_block.id,
                )
                role_blocks[("option", option.option_key)] = option_block

            for role, text_value in [
                ("answer", extracted_question.answer_text),
                ("analysis", extracted_question.analysis_text),
                ("solution", extracted_question.solution_text),
                ("comment", extracted_question.comment_text),
            ]:
                if not text_value:
                    continue
                role_payload = render_payloads.get(role)
                role_block, block_order = self._append_question_block(
                    db=db,
                    question_item_id=question_item.id,
                    block_order=block_order,
                    block_role=role,
                    text_content=text_value,
                    rich_content_json=role_payload,
                    content_format="json" if role_payload else "plain_text",
                    is_primary=True,
                )
                role_blocks[(role, None)] = role_block

            for knowledge_index, knowledge_point in enumerate(extracted_question.knowledge_points, start=1):
                knowledge_payload = render_payloads.get(f"knowledge:{knowledge_index}") or {"kind": "knowledge_point"}
                knowledge_block, block_order = self._append_question_block(
                    db=db,
                    question_item_id=question_item.id,
                    block_order=block_order,
                    block_role="knowledge",
                    text_content=knowledge_point,
                    rich_content_json=knowledge_payload,
                    content_format="json" if render_payloads.get(f"knowledge:{knowledge_index}") else "plain_text",
                    is_primary=True,
                )
                role_blocks.setdefault(("knowledge", None), knowledge_block)

            for topic_index, topic in enumerate(extracted_question.topics, start=1):
                topic_payload = render_payloads.get(f"topic:{topic_index}") or {"kind": "topic"}
                topic_block, block_order = self._append_question_block(
                    db=db,
                    question_item_id=question_item.id,
                    block_order=block_order,
                    block_role="topic",
                    text_content=topic,
                    rich_content_json=topic_payload,
                    content_format="json" if render_payloads.get(f"topic:{topic_index}") else "plain_text",
                    is_primary=True,
                )
                role_blocks.setdefault(("topic", None), topic_block)


            db.add(
                models.PaperQuestion(
                    paper_id=paper.id,
                    question_item_id=question_item.id,
                    question_no=extracted_question.question_no,
                    display_order=display_order,
                    parse_confidence=0.86,
                    source_question_label=(
                        (extracted_question.original_question_label or extracted_question.question_no) or ""
                    )[:64],
                )
            )
            if knowledge_package_id:
                db.add(
                    models.KnowledgePackageQuestion(
                        package_id=knowledge_package_id,
                        question_item_id=question_item.id,
                        display_order=display_order,
                        relation_type="topic_material",
                        source_origin="model",
                        confidence=0.8,
                        approved_status="pending",
                    )
                )

            formula_count += self._persist_formulas_for_question(
                db=db,
                question_item=question_item,
                extracted_question=extracted_question,
                role_blocks=role_blocks,
                option_rows=option_rows,
            )
            self._persist_retrieval_documents(

                db=db,
                question_item=question_item,
                source_document=source_document,
                paper=paper,
                extracted_question=extracted_question,
                stem_hash=stem_hash,
            )
            question_count += 1

        raw_outline["formula_count"] = formula_count
        raw_outline["question_asset_count"] = question_asset_count
        paper.raw_outline_json = raw_outline
        db.commit()

        return {
            "question_count": question_count,
            "formula_count": formula_count,
            "question_asset_count": question_asset_count,
            "question_item_ids": question_item_ids,
        }

    def index_document_questions(self, db: Session, source_document_id: int) -> Dict[str, object]:
        retrieval_documents = (
            db.query(models.RetrievalDocument)
            .filter(models.RetrievalDocument.metadata_json.isnot(None))
            .filter(models.RetrievalDocument.is_active.is_(True))
            .all()
        )

        payload = []
        point_rows: List[Tuple[int, str, Optional[str], Optional[str]]] = []
        for retrieval_document in retrieval_documents:
            metadata = retrieval_document.metadata_json or {}
            if metadata.get("source_document_id") != source_document_id:
                continue
            question_no = str(metadata.get("question_no") or retrieval_document.entity_id)
            entity_type = metadata.get("entity_type", retrieval_document.entity_type)
            vector_id = self._normalize_backend_vector_id(metadata.get("vector_id")) or self._build_vector_id(
                source_document_id,
                entity_type,
                question_no,
            )

            payload.append(
                {
                    "document": retrieval_document.text_for_embedding,
                    "metadata": {
                        "source": f"source_document:{source_document_id}:question:{question_no}",
                        "entity_type": entity_type,
                        "entity_id": str(retrieval_document.entity_id),
                        "source_document_id": str(source_document_id),
                        "paper_id": str(metadata.get("paper_id") or ""),
                        "question_no": question_no,
                        "subject": metadata.get("subject") or "",
                        "grade": metadata.get("grade") or "",
                        "block_role": metadata.get("block_role") or "",
                    },
                    "id": vector_id,
                }
            )
            point_rows.append((retrieval_document.id, vector_id, retrieval_document.content_hash, retrieval_document.text_for_embedding))

        if not payload:
            return {
                "indexed_documents": 0,
                **vector_db.db.backend_summary,
            }

        sync_result = vector_db.db.upsert_retrieval_documents(payload)
        db.query(models.EmbeddingPoint).filter(models.EmbeddingPoint.retrieval_document_id.in_([row[0] for row in point_rows])).delete(
            synchronize_session=False
        )
        for retrieval_document_id, point_id, content_hash, _ in point_rows:
            db.add(
                models.EmbeddingPoint(
                    retrieval_document_id=retrieval_document_id,
                    backend_type=sync_result["vector_backend"],
                    point_id=point_id,
                    model_name=sync_result["embedding_model"],
                    vector_dim=sync_result["vector_dim"],
                    content_hash=content_hash,
                )
            )
        db.commit()
        return {
            "indexed_documents": len(payload),
            **sync_result,
        }


    def _slice_before_first_option_scope_bracket(self, text: str) -> str:
        """截取到首个【答案】【解析】等正式区块之前，供选项抽取兜底。"""
        if not text:
            return ""
        m = OPTION_SCOPE_BRACKET_PATTERN.search(text)
        return text[: m.start()].strip() if m else text.strip()

    def _parse_question_segment(
        self,
        question_no: str,
        question_text: str,
        source_slice_span: Optional[Tuple[int, int]] = None,
    ) -> ExtractedQuestion:
        normalized_text = self._normalize_text(question_text)
        body_text, sections = self._split_labeled_sections(normalized_text)
        stem_text, options = self._extract_options(body_text)
        if not options:
            scoped = self._slice_before_first_option_scope_bracket(normalized_text)
            if scoped and scoped.strip() != (body_text or "").strip():
                stem_alt, options_alt = self._extract_options(scoped)
                if options_alt:
                    stem_text, options = stem_alt, options_alt
        if not options:
            stem_alt2, options_alt2 = self._extract_options(self._slice_before_first_option_scope_bracket(body_text))
            if options_alt2:
                stem_text, options = stem_alt2, options_alt2

        answer_text = self._sanitize_answer_text(sections.get("answer"))
        analysis_text = sections.get("analysis")
        solution_text = sections.get("solution")
        comment_text = sections.get("comment") or self._extract_inline_comment(normalized_text)
        comment_text = self._trim_topic_comment_bleed(comment_text)
        knowledge_points = self._coalesce_metadata_items(
            self._split_metadata_values(sections.get("knowledge")),
            self._extract_inline_metadata(normalized_text, INLINE_KNOWLEDGE_PATTERNS),
        )
        topics = self._coalesce_metadata_items(
            self._split_metadata_values(sections.get("topic")),
            self._extract_inline_metadata(normalized_text, INLINE_TOPIC_PATTERNS),
        )

        question_type = self._infer_question_type(stem_text or normalized_text, options)
        correct_option_keys = self._extract_correct_option_keys(answer_text) if question_type == "choice" else set()
        for option in options:
            if correct_option_keys:
                option.is_correct = option.option_key in correct_option_keys

        formulas = self._extract_question_formulas(
            stem_text=stem_text or normalized_text,
            options=options,
            answer_text=answer_text,
            analysis_text=analysis_text,
            solution_text=solution_text,
        )
        figure_markers = self._unique_preserve_order(FIGURE_MARKER_PATTERN.findall(normalized_text))

        return ExtractedQuestion(
            question_no=question_no,
            text=self._build_full_question_text(
                stem_text=stem_text or normalized_text,
                options=options,
                answer_text=answer_text,
                analysis_text=analysis_text,
                solution_text=solution_text,
                comment_text=comment_text,
                knowledge_points=knowledge_points,
                topics=topics,
            ),
            question_type=question_type,
            has_formula=bool(formulas),
            stem_text=stem_text or normalized_text,
            options=options,
            answer_text=answer_text,
            analysis_text=analysis_text,
            solution_text=solution_text,
            comment_text=comment_text,
            knowledge_points=knowledge_points,
            topics=topics,
            formulas=formulas,
            figure_markers=figure_markers,
            source_slice_span=source_slice_span,
        )

    def _segment_structured_questions(self, blocks: List[Dict[str, Any]]) -> List[ExtractedQuestion]:
        extracted_questions: List[ExtractedQuestion] = []
        current_question_no: Optional[str] = None
        current_question_no_int: Optional[int] = None
        current_blocks: List[Dict[str, Any]] = []

        for block in blocks:
            text = self._normalize_text(str(block.get("text") or ""))
            clean_text = _LEADING_IMAGE_PLACEHOLDER_RE.sub("", text)
            match = QUESTION_HEADER_PATTERN.match(clean_text)
            if match:
                candidate_no_int = self._extract_question_no_from_match(match)
                if current_question_no is not None and not self._is_reasonable_question_transition(current_question_no_int, candidate_no_int):
                    current_blocks.append(block)
                    continue

                if current_question_no and current_blocks:
                    extracted_questions.append(self._parse_structured_question_segment(current_question_no, current_blocks))
                current_question_no = match.group(1) or match.group(2) or match.group(3) or str(len(extracted_questions) + 1)
                current_question_no_int = candidate_no_int
                trimmed_block = self._trim_structured_block_prefix(block, match.end())
                current_blocks = [trimmed_block] if self._block_has_visible_content(trimmed_block) else []
                continue
            if current_question_no is not None:
                current_blocks.append(block)

        if current_question_no and current_blocks:
            extracted_questions.append(self._parse_structured_question_segment(current_question_no, current_blocks))
        return extracted_questions

    def _parse_structured_question_segment(
        self,
        question_no: str,
        blocks: List[Dict[str, Any]],
    ) -> ExtractedQuestion:
        role_blocks: Dict[str, List[Dict[str, Any]]] = {
            "stem": [],
            "answer": [],
            "analysis": [],
            "solution": [],
            "comment": [],
            "knowledge": [],
            "topic": [],
        }
        role_display_labels: Dict[str, str] = {}
        option_block_map: Dict[str, List[Dict[str, Any]]] = {}
        option_order: List[str] = []
        raw_option_blocks: List[Dict[str, Any]] = []
        current_role = "stem"
        current_option_key: Optional[str] = None

        for raw_block in blocks:
            display_block = copy.deepcopy(raw_block)
            block = copy.deepcopy(raw_block)
            text = self._normalize_text(str(block.get("text") or ""))
            option_match = self._match_option_label(text)
            section_match = SECTION_LABEL_PATTERN.match(text)
            sub_q_match = _SUB_Q_NO_PATTERN.match(text) if current_role == "option" else None

            if option_match:
                current_role = "option"
                current_option_key = self._normalize_option_key(option_match.group("key"))
                if current_option_key not in option_block_map:
                    option_block_map[current_option_key] = []
                    option_order.append(current_option_key)
                block = self._trim_structured_block_prefix(block, option_match.end())
            elif sub_q_match:
                current_role = "stem"
                current_option_key = None
            elif section_match:
                matched_label = self._normalize_text(section_match.group("label"))
                current_role = SECTION_LABEL_ALIASES.get(matched_label, "comment")
                current_option_key = None
                block = self._trim_structured_block_prefix(block, section_match.end())

            if not self._block_has_visible_content(block):
                continue

            if section_match and current_role != "option":
                role_display_labels.setdefault(current_role, matched_label)


            if current_role == "option" and current_option_key:
                if self._block_has_visible_content(display_block):
                    raw_option_blocks.append(display_block)
                option_block_map.setdefault(current_option_key, []).append(block)
            else:
                role_blocks.setdefault(current_role, []).append(block)


        option_block_map, option_order = self._normalize_structured_option_blocks(option_block_map, option_order)
        if option_order:
            role_blocks["stem"] = self._strip_trailing_option_paragraphs_from_stem_blocks(
                role_blocks.get("stem") or [],
                option_order,
            )

        stem_text = self._join_structured_block_texts(role_blocks.get("stem") or [])
        stem_text = self._strip_trailing_option_lines_from_stem_text(stem_text, option_order)

        if option_order:
            option_inline_parts: List[str] = []
            for ok in option_order:
                opt_txt = self._truncate_option_text_bleed(
                    self._join_structured_block_texts(option_block_map.get(ok) or [])
                )
                if opt_txt:
                    option_inline_parts.append("%s.%s" % (ok, opt_txt))
            if option_inline_parts:
                option_inline = "\n".join(option_inline_parts)
                insert_pos = self._find_grouped_subquestion_insert_pos(stem_text)
                if insert_pos > 0:
                    stem_text = stem_text[:insert_pos].rstrip() + "\n" + option_inline + "\n" + stem_text[insert_pos:]
                else:
                    stem_text = stem_text.rstrip() + "\n" + option_inline

        options = [
            ExtractedOption(
                option_key=option_key,
                option_text=self._truncate_option_text_bleed(
                    self._join_structured_block_texts(option_block_map.get(option_key) or []),
                ),
            )
            for option_key in option_order
            if (
                self._join_structured_block_texts(option_block_map.get(option_key) or [])
                # keep formula-image-only options even when plain text is empty
                or any(
                    self._block_has_visible_content(b)
                    for b in (option_block_map.get(option_key) or [])
                )
            )
        ]
        answer_text = self._sanitize_answer_text(
            self._join_structured_block_texts(role_blocks.get("answer") or []) or None,
        )
        analysis_text = self._join_structured_block_texts(role_blocks.get("analysis") or []) or None
        solution_text = self._join_structured_block_texts(role_blocks.get("solution") or []) or None
        comment_text = self._join_structured_block_texts(role_blocks.get("comment") or []) or None
        knowledge_points = [
            self._join_structured_block_texts([block])
            for block in role_blocks.get("knowledge") or []
            if self._join_structured_block_texts([block])
        ]
        topics = [
            self._join_structured_block_texts([block])
            for block in role_blocks.get("topic") or []
            if self._join_structured_block_texts([block])
        ]

        normalized_text = self._build_full_question_text(
            stem_text=stem_text,
            options=options,
            answer_text=answer_text,
            analysis_text=analysis_text,
            solution_text=solution_text,
            comment_text=comment_text,
            knowledge_points=knowledge_points,
            topics=topics,
        )
        question_type = self._infer_question_type(stem_text or normalized_text, options)
        correct_option_keys = self._extract_correct_option_keys(answer_text) if question_type == "choice" else set()
        for option in options:
            if correct_option_keys:
                option.is_correct = option.option_key in correct_option_keys

        render_payloads: Dict[str, Dict[str, object]] = {}
        if role_blocks.get("stem"):
            render_payloads["stem"] = self._build_render_group("stem", role_blocks["stem"])
        if raw_option_blocks:
            render_payloads["options"] = self._build_render_group("options", raw_option_blocks)
        for option_key in option_order:
            option_blocks = option_block_map.get(option_key) or []
            if option_blocks:
                render_payloads[f"option:{option_key}"] = self._build_render_group(
                    "option",
                    option_blocks,
                    option_key=option_key,
                )

        for role in ["answer", "analysis", "solution", "comment"]:
            blocks_for_role = role_blocks.get(role) or []
            if blocks_for_role:
                render_payloads[role] = self._build_render_group(
                    role,
                    blocks_for_role,
                    display_label=role_display_labels.get(role),
                )
        for index, block in enumerate(role_blocks.get("knowledge") or [], start=1):
            render_payloads[f"knowledge:{index}"] = self._build_render_group(
                "knowledge",
                [block],
                kind="knowledge_point",
                display_label=role_display_labels.get("knowledge"),
            )
        for index, block in enumerate(role_blocks.get("topic") or [], start=1):
            render_payloads[f"topic:{index}"] = self._build_render_group(
                "topic",
                [block],
                kind="topic",
                display_label=role_display_labels.get("topic"),
            )


        formulas = self._extract_question_formulas(
            stem_text=stem_text or normalized_text,
            options=options,
            answer_text=answer_text,
            analysis_text=analysis_text,
            solution_text=solution_text,
        )
        figure_markers = self._unique_preserve_order(FIGURE_MARKER_PATTERN.findall(normalized_text))

        return ExtractedQuestion(
            question_no=question_no,
            text=normalized_text,
            question_type=question_type,
            has_formula=bool(formulas),
            stem_text=stem_text or normalized_text,
            options=options,
            answer_text=answer_text,
            analysis_text=analysis_text,
            solution_text=solution_text,
            comment_text=comment_text,
            knowledge_points=knowledge_points,
            topics=topics,
            formulas=formulas,
            figure_markers=figure_markers,
            render_payloads=render_payloads,
        )

    @staticmethod
    def _block_has_image(block: Dict[str, Any]) -> bool:
        render = block.get("render")
        if not isinstance(render, dict):
            return False
        raw = str(render)
        return "'type': 'image'" in raw or '"type": "image"' in raw

    def _merge_structured_option_paragraphs(self, blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """将同一选项下多个段落块合并为单一 paragraph，保留原文顺序与图片。"""
        children: List[Dict[str, Any]] = []
        style: Dict[str, Any] = {}

        def absorb_render(render: Any) -> None:
            nonlocal style, children
            if not isinstance(render, dict):
                return
            if render.get("type") == "paragraph":
                if not style:
                    style = copy.deepcopy(render.get("style") or {})
                children.extend(copy.deepcopy(c) for c in (render.get("children") or []))
            elif render.get("type") == "block_group":
                for sub in render.get("blocks") or []:
                    absorb_render(sub)

        for block in blocks:
            absorb_render(block.get("render"))

        if not children:
            return []
        merged_text = self._normalize_text("".join(self._inline_node_text(c) for c in children))
        return [
            {
                "text": merged_text,
                "render": {
                    "type": "paragraph",
                    "style": style,
                    "children": children,
                },
            }
        ]

    def _normalize_structured_option_blocks(
        self,
        option_block_map: Dict[str, List[Dict[str, Any]]],
        option_order: List[str],
    ) -> Tuple[Dict[str, List[Dict[str, Any]]], List[str]]:
        normalized_map: Dict[str, List[Dict[str, Any]]] = {}
        normalized_order: List[str] = []

        for option_key in option_order:
            blocks = option_block_map.get(option_key) or []
            if not blocks:
                continue
            split_segments = self._split_inline_structured_options(option_key, blocks)
            for split_key, split_blocks in split_segments:
                if split_key not in normalized_map:
                    normalized_map[split_key] = []
                    normalized_order.append(split_key)
                normalized_map[split_key].extend(split_blocks)

        for key in list(normalized_map.keys()):
            blk_list = normalized_map[key]
            if len(blk_list) <= 1:
                continue
            with_img = [b for b in blk_list if self._block_has_image(b)]
            without_img = [b for b in blk_list if not self._block_has_image(b)]
            if with_img and without_img:
                # 仅保留图片会丢掉 “√2”“区间”等伴生文字，行内公式也会被压成一条线
                merged = self._merge_structured_option_paragraphs(blk_list)
                if merged:
                    normalized_map[key] = merged

        return normalized_map, normalized_order

    _INLINE_OPTION_LABEL_RE = re.compile(r"(?:^|\s{2,})(?P<key>[A-HＡ-Ｈ])[\.．、:：]\s*")

    def _looks_like_option_rich_paragraph(
        self,
        render: Any,
        option_keys: Sequence[str],
    ) -> bool:
        if not isinstance(render, dict) or render.get("type") != "paragraph":
            return False
        paragraph_text = self._normalize_text(
            "".join(self._inline_node_text(child) for child in (render.get("children") or []))
        )
        if not paragraph_text:
            return False
        found = {
            self._normalize_option_key(match.group("key"))
            for match in self._INLINE_OPTION_LABEL_RE.finditer(paragraph_text)
        }
        if not found:
            match = self._match_option_label(paragraph_text)
            if match:
                found = {self._normalize_option_key(match.group("key"))}
        option_key_set = {self._normalize_option_key(key) for key in option_keys if key}
        return bool(found and found.issubset(option_key_set))

    def _strip_trailing_option_paragraphs_from_stem_blocks(
        self,
        stem_blocks: List[Dict[str, Any]],
        option_order: Sequence[str],
    ) -> List[Dict[str, Any]]:
        if not stem_blocks or not option_order:
            return stem_blocks

        sanitized = [copy.deepcopy(block) for block in stem_blocks]
        while sanitized:
            tail = sanitized[-1]
            render = tail.get("render")
            if isinstance(render, dict) and render.get("type") == "block_group":
                sub_blocks = [copy.deepcopy(item) for item in (render.get("blocks") or [])]
                removed = 0
                while sub_blocks and self._looks_like_option_rich_paragraph(sub_blocks[-1], option_order):
                    sub_blocks.pop()
                    removed += 1
                if removed:
                    if sub_blocks:
                        render["blocks"] = sub_blocks
                        tail["render"] = render
                        tail["text"] = self._structured_block_text_from_render(render)
                        if self._block_has_visible_content(tail):
                            sanitized[-1] = tail
                            break
                    sanitized.pop()
                    continue

            if self._looks_like_option_rich_paragraph(render, option_order):
                sanitized.pop()
                continue
            break
        return sanitized

    def _strip_trailing_option_lines_from_stem_text(
        self,
        stem_text: str,
        option_order: Sequence[str],
    ) -> str:
        if not stem_text or not option_order:
            return stem_text
        lines = stem_text.splitlines()
        option_key_set = {self._normalize_option_key(key) for key in option_order if key}
        removed_keys: set[str] = set()
        while lines:
            candidate = self._normalize_text(lines[-1])
            if not candidate:
                lines.pop()
                continue
            found = {
                self._normalize_option_key(match.group("key"))
                for match in self._INLINE_OPTION_LABEL_RE.finditer(candidate)
            }
            if not found:
                match = self._match_option_label(candidate)
                if match:
                    found = {self._normalize_option_key(match.group("key"))}
            if not found or not found.issubset(option_key_set):
                break
            removed_keys.update(found)
            lines.pop()
        if len(removed_keys) < 2:
            return stem_text
        return "\n".join(lines).rstrip()

    def _find_grouped_subquestion_insert_pos(self, stem_text: str) -> int:
        markers: List[re.Match[str]] = []
        for match in re.finditer(r"\(\d{1,3}\)", stem_text or ""):
            line_start = stem_text.rfind("\n", 0, match.start()) + 1
            prefix = stem_text[line_start:match.start()]
            prefix_clean = prefix.strip()
            if not prefix_clean or _LEADING_IMAGE_PLACEHOLDER_RE.match(prefix_clean):
                markers.append(match)
        if len(markers) >= 2:
            return markers[1].start()
        return -1

    def _split_inline_structured_options(
        self,
        initial_option_key: str,
        blocks: List[Dict[str, Any]],
    ) -> List[Tuple[str, List[Dict[str, Any]]]]:
        all_children: List[Dict[str, Any]] = []
        style: Dict[str, Any] = {}
        for block in blocks:
            render = block.get("render")
            if isinstance(render, dict) and render.get("type") == "paragraph":
                if not style:
                    style = render.get("style") or {}
                for child in render.get("children") or []:
                    all_children.append(child)
            elif isinstance(render, dict) and render.get("type") == "block_group":
                for sub_block in render.get("blocks") or []:
                    if sub_block.get("type") == "paragraph":
                        if not style:
                            style = sub_block.get("style") or {}
                        for child in sub_block.get("children") or []:
                            all_children.append(child)

        if not all_children:
            return [(initial_option_key, blocks)]

        atomic: List[Tuple[str, Any]] = []
        for child in all_children:
            if child.get("type") != "text":
                atomic.append(("child", copy.deepcopy(child)))
                continue
            text = child.get("text") or ""
            pos = 0
            for m in self._INLINE_OPTION_LABEL_RE.finditer(text):
                before = text[pos:m.start()]
                if before.strip():
                    atomic.append(("child", {**copy.deepcopy(child), "text": before.rstrip()}))
                label_key = self._normalize_option_key(m.group("key"))
                atomic.append(("label", label_key))
                pos = m.end()
            remainder = text[pos:]
            if remainder.strip():
                atomic.append(("child", {**copy.deepcopy(child), "text": remainder.lstrip()}))

        segments: List[Tuple[str, List[Dict[str, Any]]]] = []
        current_key = initial_option_key
        current_children: List[Dict[str, Any]] = []

        for piece_type, content in atomic:
            if piece_type == "label":
                if current_children:
                    segments.append((current_key, current_children))
                current_key = content
                current_children = []
            else:
                current_children.append(content)

        if current_children:
            segments.append((current_key, current_children))

        if len(segments) <= 1:
            return [(initial_option_key, blocks)]

        result: List[Tuple[str, List[Dict[str, Any]]]] = []
        for seg_key, seg_children in segments:
            seg_text = self._normalize_text(
                "".join(self._inline_node_text(c) for c in seg_children)
            )
            result.append((seg_key, [{
                "text": seg_text,
                "render": {
                    "type": "paragraph",
                    "style": copy.deepcopy(style),
                    "children": seg_children,
                },
            }]))
        return result or [(initial_option_key, blocks)]

    def _extract_paragraph_render_hints(
        self,
        blocks: List[Dict[str, Any]],
    ) -> Tuple[Dict[str, Any], Dict[str, Any]]:
        for block in blocks:
            render = block.get("render")
            if not isinstance(render, dict) or render.get("type") != "paragraph":
                continue
            style = copy.deepcopy(render.get("style") or {})
            for child in render.get("children") or []:
                if child.get("type") in {"text", "formula"}:
                    return style, copy.deepcopy(child.get("marks") or {})
            return style, {}
        return {}, {}

    def _build_plain_structured_paragraph_block(
        self,
        text: str,
        style: Optional[Dict[str, Any]] = None,
        marks: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        normalized_text = self._normalize_text(text)
        return {
            "text": normalized_text,
            "render": {
                "type": "paragraph",
                "style": copy.deepcopy(style or {}),
                "children": [
                    {
                        "type": "text",
                        "text": normalized_text,
                        "marks": copy.deepcopy(marks or {}),
                    }
                ],
            },
        }

    def _build_render_group(

        self,
        role: str,
        blocks: List[Dict[str, Any]],
        option_key: Optional[str] = None,
        kind: Optional[str] = None,
        display_label: Optional[str] = None,
    ) -> Dict[str, object]:
        payload: Dict[str, object] = {
            "type": "block_group",
            "role": role,
            "blocks": [copy.deepcopy(block.get("render") or {}) for block in blocks if block.get("render")],
            "plain_text": self._join_structured_block_texts(blocks),
        }
        if option_key:
            payload["option_key"] = option_key
        if kind:
            payload["kind"] = kind
        if display_label:
            payload["display_label"] = display_label
        return payload


    def _join_structured_block_texts(self, blocks: List[Dict[str, Any]]) -> str:
        return self._normalize_text("\n".join(self._normalize_text(str(block.get("text") or "")) for block in blocks if block))

    def _match_option_label(self, text: str) -> Optional[re.Match]:
        return re.match(r"^\s*(?P<key>[A-HＡ-Ｈ])[\.．、:：]\s*", text or "")

    def _block_has_visible_content(self, block: Dict[str, Any]) -> bool:
        if self._normalize_text(str(block.get("text") or "")):
            return True
        return self._render_has_visible_content(block.get("render"))

    def _render_has_visible_content(self, render: Any) -> bool:
        if isinstance(render, list):
            return any(self._render_has_visible_content(item) for item in render)
        if not isinstance(render, dict):
            return False
        node_type = render.get("type")
        if node_type in {"image", "formula"}:
            return True
        if node_type == "text":
            return bool(str(render.get("text") or "").strip())
        for value in render.values():
            if self._render_has_visible_content(value):
                return True
        return False

    def _trim_structured_block_prefix(self, block: Dict[str, Any], prefix_length: int) -> Dict[str, Any]:
        trimmed_block = copy.deepcopy(block)
        trimmed_text = self._normalize_text(str(trimmed_block.get("text") or "")[prefix_length:])
        trimmed_block["text"] = trimmed_text
        render = trimmed_block.get("render")
        if isinstance(render, dict):
            trimmed_block["render"] = self._trim_render_block_prefix(render, prefix_length)
            trimmed_block["text"] = self._structured_block_text_from_render(trimmed_block.get("render")) or trimmed_text
        return trimmed_block

    def _trim_render_block_prefix(self, render: Dict[str, Any], prefix_length: int) -> Dict[str, Any]:
        render_copy = copy.deepcopy(render)
        if render_copy.get("type") != "paragraph":
            return render_copy
        children = render_copy.get("children") or []
        render_copy["children"] = self._slice_inline_children(children, prefix_length)
        return render_copy

    def _slice_inline_children(self, children: List[Dict[str, Any]], prefix_length: int) -> List[Dict[str, Any]]:
        remaining = max(prefix_length, 0)
        sliced: List[Dict[str, Any]] = []
        for child in children:
            child_text = self._inline_node_text(child)
            child_length = len(child_text)
            if remaining <= 0:
                sliced.append(copy.deepcopy(child))
                continue
            if child.get("type") == "text" and child_length:
                if remaining >= child_length:
                    remaining -= child_length
                    continue
                updated_child = copy.deepcopy(child)
                updated_child["text"] = child_text[remaining:]
                remaining = 0
                if updated_child.get("text"):
                    sliced.append(updated_child)
                continue
            if child_length and remaining >= child_length:
                remaining -= child_length
                continue
            if child.get("type") == "line_break":
                if remaining > 0:
                    remaining = max(remaining - 1, 0)
                continue
            sliced.append(copy.deepcopy(child))
            remaining = 0
        return sliced

    def _slice_inline_children_keep_prefix(self, children: List[Dict[str, Any]], max_len: int) -> List[Dict[str, Any]]:
        """保留行内树前 max_len 个「逻辑字符」（与 _inline_node_text 计数一致）。"""
        if max_len <= 0:
            return []
        out: List[Dict[str, Any]] = []
        taken = 0
        for child in children:
            if taken >= max_len:
                break
            ln = len(self._inline_node_text(child))
            if child.get("type") == "line_break":
                if taken >= max_len:
                    break
                out.append(copy.deepcopy(child))
                taken += 1
                continue
            if taken + ln <= max_len:
                out.append(copy.deepcopy(child))
                taken += ln
                continue
            need = max_len - taken
            if child.get("type") == "text" and need > 0:
                txt = self._inline_node_text(child)
                c2 = copy.deepcopy(child)
                c2["text"] = txt[:need]
                out.append(c2)
            else:
                out.append(copy.deepcopy(child))
            break
        return out

    def slice_paragraph_render_range(self, render: Dict[str, Any], start: int, end: int) -> Dict[str, Any]:
        """按「逻辑字符」下标截取 paragraph 的 render（用于同段解析+新题号合并块的拆分）。"""
        r = copy.deepcopy(render)
        if r.get("type") != "paragraph":
            return r
        children = r.get("children") or []
        span = max(0, end - start)
        after_start = self._slice_inline_children(copy.deepcopy(children), start)
        r["children"] = self._slice_inline_children_keep_prefix(after_start, span)
        return r

    _FORMULA_ALT_TEXTS = frozenset({"公式", "公式图片", "[公式]", "formula"})

    def _inline_node_text(self, node: Dict[str, Any]) -> str:
        node_type = node.get("type")
        if node_type == "text":
            return str(node.get("text") or "")
        if node_type == "formula":
            return str(node.get("text") or "")
        if node_type == "image":
            if node.get("omml_raster"):
                return ""
            alt = str(node.get("alt_text") or "")
            if alt in self._FORMULA_ALT_TEXTS:
                return ""
            return alt or "[图片]"
        if node_type == "line_break":
            return "\n"
        return ""

    def _structured_block_text_from_render(self, render: Any) -> str:
        if not isinstance(render, dict):
            return ""
        node_type = render.get("type")
        if node_type == "paragraph":
            return self._normalize_text("".join(self._inline_node_text(child) for child in render.get("children") or []))
        if node_type == "block_group":
            return self._normalize_text("\n".join(self._structured_block_text_from_render(block) for block in render.get("blocks") or []))
        if node_type == "table":
            rows: List[str] = []
            for row in render.get("rows") or []:
                cell_texts: List[str] = []
                for cell in row.get("cells") or []:
                    cell_blocks = cell.get("blocks") or []
                    text_value = self._normalize_text(
                        "\n".join(self._structured_block_text_from_render(block) for block in cell_blocks)
                    )
                    if text_value:
                        cell_texts.append(text_value)
                if cell_texts:
                    rows.append(" | ".join(cell_texts))
            return self._normalize_text("\n".join(rows))
        return ""

    def _upsert_paper(self, db: Session, source_document: models.SourceDocument) -> models.Paper:

        paper = db.query(models.Paper).filter(models.Paper.source_document_id == source_document.id).first()
        if paper:
            return paper

        content_source = source_document.content_source
        paper = models.Paper(
            source_document_id=source_document.id,
            title=source_document.title or Path(source_document.file_name or "paper").stem,
            subject=source_document.subject,
            grade=source_document.grade,
            year=source_document.year,
            region=source_document.region,
            exam_type="unknown",
            source_type=content_source.source_type if content_source else "unknown",
            is_canonical=(source_document.visibility_scope == "global"),
            review_status="draft",
            total_questions=0,
        )
        db.add(paper)
        db.commit()
        db.refresh(paper)
        return paper

    def _tear_down_paper_questions(self, db: Session, source_document_id: int, paper_ids: List[int]) -> None:
        if not paper_ids:
            return

        db.query(models.ExamSession).filter(models.ExamSession.matched_paper_id.in_(paper_ids)).update(
            {models.ExamSession.matched_paper_id: None},
            synchronize_session=False,
        )

        question_item_ids = list(
            dict.fromkeys(
                row[0]
                for row in db.query(models.PaperQuestion.question_item_id)
                .filter(models.PaperQuestion.paper_id.in_(paper_ids))
                .filter(models.PaperQuestion.question_item_id.isnot(None))
                .all()
            )
        )

        if question_item_ids:
            block_ids = list(
                dict.fromkeys(
                    row[0]
                    for row in db.query(models.QuestionBlock.id)
                    .filter(models.QuestionBlock.question_item_id.in_(question_item_ids))
                    .all()
                )
            )

            retrieval_documents = [
                retrieval_document
                for retrieval_document in db.query(models.RetrievalDocument)
                .filter(models.RetrievalDocument.is_active.is_(True))
                .all()
                if retrieval_document.entity_id in question_item_ids
                or (retrieval_document.metadata_json or {}).get("source_document_id") == source_document_id
            ]

            vector_ids = self._unique_preserve_order(
                [
                    self._resolve_vector_id_from_retrieval_document(
                        retrieval_document=retrieval_document,
                        source_document_id=source_document_id,
                    )
                    for retrieval_document in retrieval_documents
                ]
            )
            if vector_ids:
                try:
                    vector_db.db.delete_documents(ids=vector_ids)
                except Exception as exc:
                    logger.warning("Failed to delete stale vectors for source_document=%s: %s", source_document_id, exc)

            retrieval_document_ids = [item.id for item in retrieval_documents]
            if retrieval_document_ids:
                db.query(models.EmbeddingPoint).filter(
                    models.EmbeddingPoint.retrieval_document_id.in_(retrieval_document_ids)
                ).delete(synchronize_session=False)

            db.query(models.KnowledgePackageQuestion).filter(
                models.KnowledgePackageQuestion.question_item_id.in_(question_item_ids)
            ).delete(synchronize_session=False)
            db.query(models.KnowledgeQuestionLink).filter(
                models.KnowledgeQuestionLink.question_item_id.in_(question_item_ids)
            ).delete(synchronize_session=False)

            db.query(models.QuestionMatchResult).filter(
                models.QuestionMatchResult.candidate_question_id.in_(question_item_ids)
            ).delete(synchronize_session=False)
            db.query(models.StudentAttempt).filter(models.StudentAttempt.question_item_id.in_(question_item_ids)).update(
                {models.StudentAttempt.question_item_id: None},
                synchronize_session=False,
            )
            db.query(models.ExamSessionQuestion).filter(
                models.ExamSessionQuestion.question_item_id.in_(question_item_ids)
            ).update(
                {models.ExamSessionQuestion.question_item_id: None},
                synchronize_session=False,
            )
            db.query(models.QuestionRelation).filter(
                (models.QuestionRelation.source_question_id.in_(question_item_ids))
                | (models.QuestionRelation.target_question_id.in_(question_item_ids))
            ).delete(synchronize_session=False)
            if block_ids:
                db.query(models.QuestionTagLink).filter(
                    (models.QuestionTagLink.question_item_id.in_(question_item_ids))
                    | (models.QuestionTagLink.evidence_block_id.in_(block_ids))
                ).delete(synchronize_session=False)
                db.query(models.QuestionBlock).filter(models.QuestionBlock.id.in_(block_ids)).update(
                    {models.QuestionBlock.formula_id: None, models.QuestionBlock.parent_block_id: None},
                    synchronize_session=False,
                )
            else:
                db.query(models.QuestionTagLink).filter(
                    models.QuestionTagLink.question_item_id.in_(question_item_ids)
                ).delete(synchronize_session=False)
            db.query(models.QuestionOption).filter(models.QuestionOption.question_item_id.in_(question_item_ids)).update(
                {models.QuestionOption.formula_id: None},
                synchronize_session=False,
            )
            db.query(models.Formula).filter(models.Formula.question_item_id.in_(question_item_ids)).delete(
                synchronize_session=False
            )
            db.query(models.QuestionOption).filter(models.QuestionOption.question_item_id.in_(question_item_ids)).delete(
                synchronize_session=False
            )
            if block_ids:
                db.query(models.QuestionBlock).filter(models.QuestionBlock.id.in_(block_ids)).delete(
                    synchronize_session=False
                )
            db.query(models.Asset).filter(models.Asset.owner_type == "question_item").filter(
                models.Asset.owner_id.in_(question_item_ids)
            ).delete(synchronize_session=False)
            db.query(models.RetrievalDocument).filter(models.RetrievalDocument.id.in_(retrieval_document_ids)).delete(
                synchronize_session=False
            )
            db.query(models.PaperQuestion).filter(models.PaperQuestion.paper_id.in_(paper_ids)).delete(
                synchronize_session=False
            )
            db.query(models.QuestionItem).filter(models.QuestionItem.id.in_(question_item_ids)).delete(
                synchronize_session=False
            )

        else:
            db.query(models.PaperQuestion).filter(models.PaperQuestion.paper_id.in_(paper_ids)).delete(
                synchronize_session=False
            )

        db.query(models.PaperSection).filter(models.PaperSection.paper_id.in_(paper_ids)).delete(synchronize_session=False)
        db.query(models.Paper).filter(models.Paper.id.in_(paper_ids)).delete(synchronize_session=False)

    def clear_topic_material_papers_for_document(
        self,
        db: Session,
        source_document_id: int,
        package_ids: Optional[Sequence[int]] = None,
    ) -> int:
        """删除某文档下「专题材料卷」（Paper.knowledge_package_id 非空）及关联题目。"""
        query = db.query(models.Paper).filter(models.Paper.source_document_id == source_document_id)
        query = query.filter(models.Paper.knowledge_package_id.isnot(None))
        if package_ids is not None:
            id_list = [int(x) for x in package_ids]
            if not id_list:
                return 0
            query = query.filter(models.Paper.knowledge_package_id.in_(id_list))
        paper_ids = [row.id for row in query.all()]
        if not paper_ids:
            return 0
        self._tear_down_paper_questions(db, source_document_id, paper_ids)
        db.commit()
        return len(paper_ids)

    def ingest_topic_packages_questions(
        self,
        db: Session,
        source_document_id: int,
        progress_callback: Optional[Callable[[str], None]] = None,
    ) -> Dict[str, object]:
        """按 KnowledgePackage 的页码范围切 PDF 文本并切题落库；每专题一张「材料卷」并写入 KnowledgePackageQuestion。"""

        def notify(msg: str) -> None:
            if progress_callback:
                try:
                    progress_callback(msg)
                except Exception:
                    pass

        source_document = db.query(models.SourceDocument).filter(models.SourceDocument.id == source_document_id).first()
        if not source_document:
            raise ValueError(f"SourceDocument {source_document_id} 不存在")

        packages = (
            db.query(models.KnowledgePackage)
            .filter(models.KnowledgePackage.source_document_id == source_document_id)
            .order_by(models.KnowledgePackage.id.asc())
            .all()
        )
        if not packages:
            return {"status": "skipped", "reason": "no_knowledge_packages", "papers_created": 0, "question_count": 0}

        source_path = self._resolve_local_path(source_document.storage_url)
        if source_path.suffix.lower() != ".pdf":
            return {"status": "skipped", "reason": "not_pdf", "papers_created": 0, "question_count": 0}

        pkg_ids = [pkg.id for pkg in packages]
        notify(f"清理旧专题材料卷（共 {len(pkg_ids)} 个专题包）…")
        self.clear_topic_material_papers_for_document(db, source_document_id, package_ids=pkg_ids)

        notify("读取 PDF 目标页文本（结构化提取 + PUA 符号映射；仅纯图片页 fallback OCR）…")
        needed_pages: set[int] = set()
        for pkg in packages:
            page_range = pkg.page_range_json or {}
            start_p = int(page_range.get("start") or 1)
            end_p = int(page_range.get("end") or start_p)
            if start_p > end_p:
                start_p, end_p = end_p, start_p
            for pn in range(max(1, start_p), max(1, end_p) + 1):
                needed_pages.add(pn)

        page_text_by_no: Dict[int, str] = {}
        page_rich_by_no: Dict[int, List[Dict[str, Any]]] = {}
        if needed_pages:
            from . import question_bank_pdf_ocr_fallback as ocr_fb
            from .config import QUESTION_BANK_PDF_OCR_RENDER_SCALE
            from .pdf_structured_extractor import extract_page_structured, is_page_text_based

            ocr_ok = ocr_fb.pix2text_available()
            render_scale = QUESTION_BANK_PDF_OCR_RENDER_SCALE
            math_clip_dir = None
            if QUESTION_BANK_PDF_MATH_CLIP_IMAGES:
                math_clip_dir = Path(QUESTION_BANK_ASSET_DIR) / f"document_{source_document.id}" / "pdf_math_clips"
                math_clip_dir.mkdir(parents=True, exist_ok=True)

            with fitz.open(source_path) as doc:
                page_count = len(doc)
                target_pages = sorted(pn for pn in needed_pages if 1 <= pn <= page_count)
                total_targets = len(target_pages)
                for idx_pn, pn in enumerate(target_pages, start=1):
                    page = doc[pn - 1]
                    extraction_mode = "structured"

                    if is_page_text_based(page, min_spans=5):
                        content_obj = extract_page_structured(
                            page, page_no=pn, math_clip_output_dir=math_clip_dir,
                        )
                        page_text_by_no[pn] = content_obj.plain_text
                        page_rich_by_no[pn] = content_obj.rich_paragraphs
                    else:
                        extraction_mode = "ocr"
                        if ocr_ok:
                            txt = ocr_fb.ocr_pdf_page_to_text(source_path, pn - 1, render_scale)
                            page_text_by_no[pn] = txt if txt else ""
                        else:
                            page_text_by_no[pn] = (page.get_text(sort=True) or "").strip()

                    if idx_pn == 1 or idx_pn == total_targets or idx_pn % 10 == 0:
                        notify(f"  已读取页 {idx_pn}/{total_targets}（pn={pn} mode={extraction_mode}）")

        pages = [(pn, (page_text_by_no.get(pn) or "").strip()) for pn in sorted(page_text_by_no.keys()) if (page_text_by_no.get(pn) or "").strip()]
        notify(f"PDF 有效页块：{len(pages)} 段（目标页={len(page_text_by_no)}）")

        file_stem = Path(source_document.file_name or "document").stem
        papers_created = 0
        question_total = 0
        paper_summaries: List[Dict[str, object]] = []
        n_pkg = len(packages)

        for idx, pkg in enumerate(packages, start=1):
            title_short = (pkg.package_title or "")[:72]
            notify(f"专题题 [{idx}/{n_pkg}] package_id={pkg.id} «{title_short}» 页码切片 + 切题…")
            page_range = pkg.page_range_json or {}
            start_p = int(page_range.get("start") or 1)
            end_p = int(page_range.get("end") or start_p)
            slice_parts = [txt for pn, txt in pages if start_p <= pn <= end_p]
            slice_text = "\n".join(slice_parts)
            if not slice_text.strip():
                notify(f"  [跳过] 页码范围 {start_p}-{end_p} 无文本")
                continue
            extracted_questions = self.segment_questions(slice_text)
            if not extracted_questions:
                notify(f"  [跳过] 切题结果为 0（页 {start_p}-{end_p}）")
                continue

            # 将结构化 rich content 映射到每道切分出来的题目
            if page_rich_by_no:
                self._attach_rich_payloads_to_questions(
                    extracted_questions, slice_text,
                    pages, page_rich_by_no, start_p, end_p,
                )

            notify(f"  segment_questions → {len(extracted_questions)} 题，写入 Paper 与 QuestionItem…")
            title = f"{file_stem} · {pkg.package_title}"[:255]
            paper = models.Paper(
                source_document_id=source_document.id,
                knowledge_package_id=pkg.id,
                title=title,
                subject=source_document.subject,
                grade=source_document.grade,
                year=source_document.year,
                region=source_document.region,
                exam_type="topic_material",
                source_type="topic_material",
                is_canonical=False,
                review_status="draft",
                total_questions=0,
            )
            db.add(paper)
            db.flush()

            metrics = self.persist_questions(
                db=db,
                paper=paper,
                source_document=source_document,
                extracted_text=slice_text,
                extracted_questions=extracted_questions,
                document_asset_count=0,
                knowledge_package_id=pkg.id,
            )
            papers_created += 1
            qn = int(metrics.get("question_count") or 0)
            question_total += qn
            qids = metrics.get("question_item_ids") or []
            id_range = ""
            if isinstance(qids, list) and qids:
                id_range = f" id∈[{min(qids)}, {max(qids)}]"
            notify(f"  ✓ paper_id={paper.id} 题目数={qn}{id_range}")
            paper_summaries.append(
                {
                    "package_id": pkg.id,
                    "package_title": pkg.package_title,
                    "paper_id": paper.id,
                    "page_range": {"start": start_p, "end": end_p},
                    "question_count": qn,
                    "question_item_ids": qids,
                }
            )

        db.commit()

        notify(f"批量提交完成，开始向量索引（source_document_id={source_document_id}）…")
        try:
            index_metrics = self.index_document_questions(db, source_document_id)
            notify(f"向量索引完成：indexed_documents={index_metrics.get('indexed_documents', 0)}")
        except Exception as exc:
            logger.warning("Topic question vector index failed: source_document_id=%s err=%s", source_document_id, exc)
            index_metrics = {"indexed_documents": 0}
            notify(f"向量索引失败（已记录日志）：{exc}")

        return {
            "status": "success",
            "papers_created": papers_created,
            "question_count": question_total,
            "indexed_documents": index_metrics.get("indexed_documents", 0),
            "paper_summaries": paper_summaries,
            "source_document_id": source_document_id,
        }

    def _clear_existing_document_artifacts(self, db: Session, source_document_id: int) -> None:
        papers = db.query(models.Paper).filter(models.Paper.source_document_id == source_document_id).all()

        db.query(models.Asset).filter(models.Asset.owner_type == "source_document").filter(
            models.Asset.owner_id == source_document_id
        ).delete(synchronize_session=False)

        if not papers:
            db.commit()
            return

        paper_ids = [paper.id for paper in papers]
        self._tear_down_paper_questions(db, source_document_id, paper_ids)
        db.commit()

    def _start_job(
        self,
        db: Session,
        source_document_id: int,
        stage: str,
        tool_name: Optional[str] = None,
    ) -> models.DocumentParseJob:
        job = models.DocumentParseJob(
            source_document_id=source_document_id,
            job_stage=stage,
            tool_name=tool_name,
            status="running",
        )
        db.add(job)
        db.commit()
        db.refresh(job)
        return job

    def _finish_job(
        self,
        db: Session,
        job: models.DocumentParseJob,
        output_location: Optional[str] = None,
        metrics_json: Optional[Dict[str, object]] = None,
    ) -> None:
        job.status = "success"
        job.output_location = output_location
        job.metrics_json = metrics_json
        job.ended_at = datetime.utcnow()
        db.commit()

    def _fail_job(self, db: Session, job: models.DocumentParseJob, error_message: str) -> None:
        job.status = "failed"
        job.error_message = error_message[:2000]
        job.ended_at = datetime.utcnow()
        db.commit()

    def _resolve_local_path(self, storage_url: str) -> Path:
        if storage_url.startswith("file://"):
            storage_url = storage_url[7:]
        candidate = Path(storage_url)
        if candidate.is_absolute():
            return candidate
        return Path(os.getcwd()) / candidate

    def _normalize_extension(self, file_ext: str) -> str:
        normalized = file_ext.lower()
        if not normalized.startswith("."):
            normalized = f".{normalized}"
        return normalized

    def _path_to_storage_url(self, path: Optional[Path]) -> Optional[str]:
        return str(path) if path else None

    def _extract_text_from_docx(self, path: Path) -> str:
        document = Document(path)
        parts: List[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    def _extract_pdf_page_texts(self, path: Path) -> List[str]:
        """逐页取 PDF 文本：默认 PyMuPDF 文本层；按质量分与 QUESTION_BANK_PDF_OCR_MODE 可整页走 Pix2Text。"""
        from . import question_bank_pdf_ocr_fallback as ocr_fb
        from .config import QUESTION_BANK_PDF_OCR_MODE, QUESTION_BANK_PDF_OCR_RENDER_SCALE, QUESTION_BANK_PDF_OCR_THRESHOLD
        from .question_bank_pdf_text_quality import score_pdf_text_layer_quality

        mode = QUESTION_BANK_PDF_OCR_MODE
        if mode not in {"off", "auto", "force"}:
            mode = "auto"
        ocr_ok = ocr_fb.pix2text_available()
        threshold = QUESTION_BANK_PDF_OCR_THRESHOLD
        render_scale = QUESTION_BANK_PDF_OCR_RENDER_SCALE

        out: List[str] = []
        document = fitz.open(path)
        try:
            page_count = len(document)
            for i in range(page_count):
                page = document[i]
                raw = page.get_text(sort=True) or ""
                score = score_pdf_text_layer_quality(raw)
                use_ocr = False
                if mode == "off":
                    use_ocr = False
                elif not ocr_ok:
                    use_ocr = False
                elif mode == "force":
                    use_ocr = True
                else:
                    use_ocr = score >= threshold or not raw.strip()

                if use_ocr:
                    txt = ocr_fb.ocr_pdf_page_to_text(path, i, render_scale)
                    if txt:
                        logger.info(
                            "PDF OCR page %s/%s file=%s quality_score=%.3f threshold=%.3f mode=%s",
                            i + 1,
                            page_count,
                            path.name,
                            score,
                            threshold,
                            mode,
                        )
                        out.append(txt)
                    else:
                        logger.warning("PDF OCR page %s returned empty; fallback to text layer", i + 1)
                        out.append(raw)
                else:
                    if mode == "auto" and not ocr_ok and score >= threshold:
                        logger.warning(
                            "PDF page %s quality_score=%.3f >= %.3f but pix2text not installed; using text layer",
                            i + 1,
                            score,
                            threshold,
                        )
                    if mode == "force" and not ocr_ok and i == 0:
                        logger.warning("QUESTION_BANK_PDF_OCR_MODE=force but pix2text not installed; using text layer for all pages")
                    out.append(raw)
        finally:
            document.close()
        return out

    def _extract_text_from_pdf(self, path: Path) -> str:
        pieces = [p for p in self._extract_pdf_page_texts(path) if (p or "").strip()]
        return "\n".join(pieces)

    def _extract_text_from_txt(self, path: Path) -> str:
        return path.read_text(encoding="utf-8")

    def _extract_docx_assets(self, path: Path, output_dir: Path) -> List[ExtractedAsset]:
        document = Document(path)
        assets: List[ExtractedAsset] = []
        seen_hashes = set()
        asset_index = 1

        for rel in document.part.rels.values():
            if rel.reltype != RT.IMAGE:
                continue
            image_part = rel.target_part
            blob = image_part.blob
            file_hash = hashlib.sha256(blob).hexdigest()
            if file_hash in seen_hashes:
                continue
            seen_hashes.add(file_hash)
            suffix = Path(image_part.partname).suffix or ".png"
            storage_path = output_dir / f"docx_asset_{asset_index:03d}_{file_hash[:12]}{suffix}"
            storage_path.write_bytes(blob)
            width, height = self._get_image_dimensions(blob)
            assets.append(
                ExtractedAsset(
                    asset_role="embedded_image",
                    storage_url=str(storage_path),
                    file_hash=file_hash,
                    width=width,
                    height=height,
                    caption_text=f"docx_embedded_image_{asset_index}",
                )
            )
            asset_index += 1

        return assets

    def _extract_pdf_assets(self, path: Path, output_dir: Path) -> List[ExtractedAsset]:
        assets: List[ExtractedAsset] = []
        seen_keys = set()
        asset_index = 1
        document = fitz.open(path)
        try:
            for page_index, page in enumerate(document, start=1):
                for image_info in page.get_images(full=True):
                    xref = image_info[0]
                    extracted = document.extract_image(xref)
                    if not extracted:
                        continue
                    blob = extracted.get("image")
                    if not blob:
                        continue
                    file_hash = hashlib.sha256(blob).hexdigest()
                    dedupe_key = (page_index, file_hash)
                    if dedupe_key in seen_keys:
                        continue
                    seen_keys.add(dedupe_key)
                    suffix = f'.{extracted.get("ext") or "png"}'
                    storage_path = output_dir / f"pdf_p{page_index}_{asset_index:03d}_{file_hash[:12]}{suffix}"
                    storage_path.write_bytes(blob)
                    width, height = self._get_image_dimensions(blob)
                    assets.append(
                        ExtractedAsset(
                            asset_role="embedded_image",
                            storage_url=str(storage_path),
                            file_hash=file_hash,
                            page_no=page_index,
                            width=width,
                            height=height,
                            caption_text=f"pdf_embedded_image_page_{page_index}_{asset_index}",
                        )
                    )
                    asset_index += 1
        finally:
            document.close()

        return assets

    def _resolve_soffice_binary(self) -> Optional[str]:
        candidates = [
            shutil.which("soffice"),
            shutil.which("soffice.exe"),
            shutil.which("soffice.com"),
            r"D:\Program Files\LibreOffice\program\soffice.exe",
            r"D:\Program Files\LibreOffice\program\soffice.com",
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files\LibreOffice\program\soffice.com",
        ]
        for candidate in candidates:
            if candidate and Path(candidate).exists():
                return str(candidate)
        return None

    def _convert_with_libreoffice(
        self,
        source_path: Path,
        output_dir: Path,
        target_format: str,
        optional: bool = False,
    ) -> Optional[Path]:
        soffice_binary = self._resolve_soffice_binary()
        if not soffice_binary:
            logger.warning(
                "LibreOffice binary not found: source_path=%s target_format=%s optional=%s",
                source_path,
                target_format,
                optional,
            )
            if optional:
                return None
            raise DocumentNormalizationError("未找到 LibreOffice 可执行文件 soffice")

        output_dir.mkdir(parents=True, exist_ok=True)
        final_output_path = output_dir / f"normalized.{target_format.lower()}"

        with tempfile.TemporaryDirectory(prefix="qb_lo_convert_") as temp_dir:
            temp_root = Path(temp_dir)
            staged_source_path = temp_root / f"source{source_path.suffix.lower()}"
            staged_output_dir = temp_root / "out"
            staged_output_dir.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source_path, staged_source_path)

            command = [
                soffice_binary,
                "--headless",
                "--convert-to",
                target_format,
                "--outdir",
                str(staged_output_dir),
                str(staged_source_path),
            ]
            started_at = time.perf_counter()
            logger.info(
                "LibreOffice convert start: source_path=%s staged_source=%s target_format=%s optional=%s output_dir=%s command=%s",
                source_path,
                staged_source_path,
                target_format,
                optional,
                staged_output_dir,
                command,
            )
            completed = subprocess.run(command, capture_output=True, text=True, check=False)
            elapsed = f"{time.perf_counter() - started_at:.2f}s"
            logger.info(
                "LibreOffice convert finished: source_path=%s staged_source=%s target_format=%s optional=%s elapsed=%s returncode=%s stdout_tail=%s stderr_tail=%s",
                source_path,
                staged_source_path,
                target_format,
                optional,
                elapsed,
                completed.returncode,
                (completed.stdout or "")[-500:],
                (completed.stderr or "")[-500:],
            )

            converted_path = staged_output_dir / f"{staged_source_path.stem}.{target_format.lower()}"
            if not converted_path.exists():
                matches = sorted(staged_output_dir.glob(f"*.{target_format.lower()}"))
                if matches:
                    converted_path = matches[0]

            if not converted_path.exists():
                logger.warning(
                    "LibreOffice convert output missing: source_path=%s target_format=%s optional=%s staged_output_dir=%s returncode=%s",
                    source_path,
                    target_format,
                    optional,
                    staged_output_dir,
                    completed.returncode,
                )
                fallback_error = completed.stderr or completed.stdout or "LibreOffice 转换失败"
                fallback_path = self._convert_with_word(source_path, output_dir, target_format, optional=True)
                if fallback_path is not None:
                    logger.info("Word convert fallback output found: %s", fallback_path)
                    return fallback_path
                if optional:
                    logger.warning("LibreOffice optional convert failed without Word fallback: source_path=%s target_format=%s", source_path, target_format)
                    return None
                raise DocumentNormalizationError(fallback_error if completed.returncode != 0 else f"转换结果不存在: {final_output_path}")

            shutil.copy2(converted_path, final_output_path)
            if completed.returncode != 0:
                logger.warning(
                    "LibreOffice convert returned non-zero but output exists: source_path=%s target_format=%s returncode=%s output=%s",
                    source_path,
                    target_format,
                    completed.returncode,
                    final_output_path,
                )
            logger.info("LibreOffice convert output found: %s", final_output_path)
            return final_output_path

    def _convert_with_word(
        self,
        source_path: Path,
        output_dir: Path,
        target_format: str,
        optional: bool = False,
    ) -> Optional[Path]:
        if pythoncom is None or (DispatchEx is None and gencache is None):
            logger.warning(
                "Word convert unavailable: source_path=%s target_format=%s optional=%s reason=pythoncom_unavailable",
                source_path,
                target_format,
                optional,
            )
            return None if optional else None

        target_format_normalized = str(target_format or "").lower()
        if target_format_normalized not in {"docx", "pdf"}:
            logger.warning("Word convert skipped: source_path=%s target_format=%s reason=unsupported_format", source_path, target_format)
            return None if optional else None

        output_dir.mkdir(parents=True, exist_ok=True)
        final_output_path = output_dir / f"normalized.{target_format_normalized}"
        app = None
        document = None
        started_at = time.perf_counter()
        try:
            pythoncom.CoInitialize()
            app = self._create_word_application()
            app.Visible = False
            app.DisplayAlerts = 0
            app.ScreenUpdating = False
            document = self._word_call(lambda: app.Documents.Open(str(source_path), ReadOnly=True, AddToRecentFiles=False))
            time.sleep(1.0)
            if final_output_path.exists():
                final_output_path.unlink()
            logger.info(
                "Word convert start: source_path=%s target_format=%s optional=%s output_path=%s",
                source_path,
                target_format_normalized,
                optional,
                final_output_path,
            )
            if target_format_normalized == "docx":
                self._word_call(lambda: document.SaveAs2(str(final_output_path), FileFormat=16))
            else:
                self._word_call(lambda: document.ExportAsFixedFormat(str(final_output_path), 17))
            logger.info(
                "Word convert finished: source_path=%s target_format=%s optional=%s elapsed=%s output_exists=%s",
                source_path,
                target_format_normalized,
                optional,
                f"{time.perf_counter() - started_at:.2f}s",
                final_output_path.exists(),
            )
            if final_output_path.exists():
                return final_output_path
            return None
        except Exception:
            logger.exception(
                "Word convert failed: source_path=%s target_format=%s optional=%s elapsed=%s",
                source_path,
                target_format_normalized,
                optional,
                f"{time.perf_counter() - started_at:.2f}s",
            )
            return None
        finally:
            if document is not None:
                try:
                    document.Close(False)
                except Exception:
                    pass
            if app is not None:
                try:
                    app.Quit()
                except Exception:
                    pass
            if pythoncom is not None:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

    def _create_word_application(self):
        last_error = None
        factories = []
        if gencache is not None:
            factories.append(lambda: gencache.EnsureDispatch("Word.Application"))
        if DispatchEx is not None:
            factories.append(lambda: DispatchEx("Word.Application"))
        for factory in factories:
            try:
                return factory()
            except Exception as exc:
                last_error = exc
        if last_error is not None:
            raise last_error
        raise RuntimeError("Word COM is unavailable")

    def _word_call(self, func, retries: int = 30, delay_seconds: float = 1.0):
        last_error = None
        for _ in range(retries):
            try:
                return func()
            except Exception as exc:
                last_error = exc
                if pythoncom is not None:
                    try:
                        pythoncom.PumpWaitingMessages()
                    except Exception:
                        pass
                time.sleep(delay_seconds)
        raise last_error


    def _append_question_block(
        self,
        db: Session,
        question_item_id: int,
        block_order: int,
        block_role: str,
        text_content: str,
        rich_content_json: Optional[Dict[str, object]] = None,
        content_format: str = "plain_text",
        parent_block_id: Optional[int] = None,
        is_primary: bool = False,
    ) -> Tuple[models.QuestionBlock, int]:
        block = models.QuestionBlock(
            question_item_id=question_item_id,
            block_order=block_order,
            block_role=block_role,
            content_format=content_format,
            text_content=text_content,
            rich_content_json=rich_content_json,
            parent_block_id=parent_block_id,
            source_origin="explicit",
            confidence=1.0,
            is_primary=is_primary,
        )

        db.add(block)
        db.flush()
        return block, block_order + 1

    def _persist_formulas_for_question(
        self,
        db: Session,
        question_item: models.QuestionItem,
        extracted_question: ExtractedQuestion,
        role_blocks: Dict[Tuple[str, Optional[str]], models.QuestionBlock],
        option_rows: Dict[str, models.QuestionOption],
    ) -> int:
        count = 0
        for formula in extracted_question.formulas:
            target_block = role_blocks.get((formula.block_role, formula.option_key)) or role_blocks.get((formula.block_role, None))
            formula_row = models.Formula(
                question_item_id=question_item.id,
                block_id=target_block.id if target_block else None,
                source_type=formula.source_type,
                latex_text=formula.source_text if "\\" in formula.source_text or "$" in formula.source_text else None,
                linear_text=formula.source_text,
                normalized_signature=formula.normalized_signature,
                parse_confidence=0.68,
            )
            db.add(formula_row)
            db.flush()

            if target_block and not target_block.formula_id:
                target_block.formula_id = formula_row.id
            if formula.option_key and formula.option_key in option_rows and not option_rows[formula.option_key].formula_id:
                option_rows[formula.option_key].formula_id = formula_row.id
            count += 1
        return count

    def _persist_question_assets(
        self,
        db: Session,
        source_document: models.SourceDocument,
        question_item: models.QuestionItem,
        extracted_question: ExtractedQuestion,
    ) -> Tuple[int, Dict[str, int]]:
        base_storage_url = (
            source_document.normalized_pdf_url
            or source_document.normalized_docx_url
            or source_document.storage_url
        )
        asset_count = 0
        asset_lookup: Dict[str, int] = {}
        for marker in self._unique_preserve_order(extracted_question.figure_markers):
            db.add(
                models.Asset(
                    tenant_id=source_document.tenant_id,
                    owner_type="question_item",
                    owner_id=question_item.id,
                    asset_role="figure_reference",
                    storage_url=base_storage_url,
                    ocr_text=marker,
                    caption_text=f"detected_figure_reference:{marker}",
                )
            )
            asset_count += 1

        for image_node in self._collect_render_image_nodes(extracted_question.render_payloads):
            storage_url = str(image_node.get("storage_url") or "").strip()
            file_hash = str(image_node.get("file_hash") or "").strip() or None
            lookup_key = file_hash or storage_url
            if not lookup_key or lookup_key in asset_lookup:
                continue
            asset = models.Asset(
                tenant_id=source_document.tenant_id,
                owner_type="question_item",
                owner_id=question_item.id,
                asset_role="question_inline_image",
                storage_url=storage_url,
                width=self._safe_int(image_node.get("width")),
                height=self._safe_int(image_node.get("height")),
                ocr_text=image_node.get("alt_text"),
                caption_text=f"question_inline_image:{image_node.get('role') or 'stem'}",
                file_hash=file_hash,
            )
            db.add(asset)
            db.flush()
            asset_lookup[lookup_key] = asset.id
            asset_count += 1
        return asset_count, asset_lookup

    def _collect_render_image_nodes(
        self,
        render_payloads: Dict[str, Dict[str, object]],
    ) -> List[Dict[str, object]]:
        image_nodes: List[Dict[str, object]] = []
        for payload_key, payload in render_payloads.items():
            role = payload_key.split(":", 1)[0]
            for node in self._walk_render_nodes(payload):
                if node.get("type") != "image":
                    continue
                image_node = copy.deepcopy(node)
                image_node.setdefault("role", role)
                image_nodes.append(image_node)
        return image_nodes

    def _attach_asset_ids_to_render_payloads(
        self,
        render_payloads: Dict[str, Dict[str, object]],
        asset_lookup: Dict[str, int],
    ) -> Dict[str, Dict[str, object]]:
        payload_copy = copy.deepcopy(render_payloads)
        for payload in payload_copy.values():
            self._attach_asset_ids_in_node(payload, asset_lookup)
        return payload_copy

    def _attach_asset_ids_in_node(self, node: Any, asset_lookup: Dict[str, int]) -> None:
        if isinstance(node, list):
            for item in node:
                self._attach_asset_ids_in_node(item, asset_lookup)
            return
        if not isinstance(node, dict):
            return

        if node.get("type") == "image":
            lookup_key = str(node.get("file_hash") or "").strip() or str(node.get("storage_url") or "").strip()
            if lookup_key and lookup_key in asset_lookup:
                node["asset_id"] = asset_lookup[lookup_key]

        for value in node.values():
            self._attach_asset_ids_in_node(value, asset_lookup)

    def _walk_render_nodes(self, node: Any) -> Iterable[Dict[str, object]]:
        if isinstance(node, list):
            for item in node:
                yield from self._walk_render_nodes(item)
            return
        if not isinstance(node, dict):
            return

        yield node
        for value in node.values():
            yield from self._walk_render_nodes(value)

    def _persist_retrieval_documents(

        self,
        db: Session,
        question_item: models.QuestionItem,
        source_document: models.SourceDocument,
        paper: models.Paper,
        extracted_question: ExtractedQuestion,
        stem_hash: str,
    ) -> None:
        base_metadata: Dict[str, Any] = {
            "source_document_id": source_document.id,
            "paper_id": paper.id,
            "question_no": extracted_question.question_no,
            "subject": source_document.subject,
            "grade": source_document.grade,
        }
        if getattr(paper, "knowledge_package_id", None):
            base_metadata["knowledge_package_id"] = paper.knowledge_package_id

        retrieval_payloads = [
            ("question_stem", extracted_question.stem_text or extracted_question.text, "stem", stem_hash),
            (
                "question_answer",
                extracted_question.answer_text,
                "answer",
                self._hash_text(extracted_question.answer_text),
            ),
            (
                "question_analysis",
                extracted_question.analysis_text,
                "analysis",
                self._hash_text(extracted_question.analysis_text),
            ),
            (
                "question_solution",
                extracted_question.solution_text,
                "solution",
                self._hash_text(extracted_question.solution_text),
            ),
            (
                "question_comment",
                extracted_question.comment_text,
                "comment",
                self._hash_text(extracted_question.comment_text),
            ),
            (
                "question_knowledge",
                "；".join(extracted_question.knowledge_points) if extracted_question.knowledge_points else None,
                "knowledge",
                self._hash_text("；".join(extracted_question.knowledge_points)),
            ),
            (
                "question_topic",
                "；".join(extracted_question.topics) if extracted_question.topics else None,
                "topic",
                self._hash_text("；".join(extracted_question.topics)),
            ),
        ]

        for entity_type, text_value, block_role, content_hash in retrieval_payloads:
            normalized_value = self._normalize_text(text_value or "")
            if not normalized_value:
                continue
            metadata = {
                **base_metadata,
                "entity_type": entity_type,
                "block_role": block_role,
                "vector_id": self._build_vector_id(
                    source_document.id,
                    entity_type,
                    extracted_question.question_no,
                ),
            }
            db.add(
                models.RetrievalDocument(
                    tenant_id=source_document.tenant_id,
                    entity_type=entity_type,
                    entity_id=question_item.id,
                    text_for_bm25=normalized_value,
                    text_for_embedding=normalized_value,
                    metadata_json=metadata,
                    is_active=True,
                    content_hash=content_hash,
                )
            )

    @staticmethod
    def _trim_topic_comment_bleed(text: Optional[str]) -> Optional[str]:
        """裁掉专家解读/点评后误并入的图书大节标题（如「一、考向分析」「（1）确定性」）。"""
        if not text or not text.strip():
            return text
        m = re.search(
            r"\n(?:[一二三四五六七八九十百千]{1,4}[、．]|（\s*\d{1,2}\s*[)）])",
            text,
        )
        if m and m.start() > 30:
            return text[: m.start()].strip()
        return text

    @staticmethod
    def _locate_field_span_for_rich(
        slice_text: str,
        field_text: Optional[str],
        role: str,
        lo: int,
        hi: int,
    ) -> Optional[Tuple[int, int]]:
        """在 slice_text[lo:hi] 内定位 field_text 的 [start,end)，避免短串误匹配到选项区。"""
        if not field_text or not field_text.strip() or lo < 0 or hi > len(slice_text) or lo >= hi:
            return None
        region = slice_text[lo:hi]
        ft = field_text.strip()
        if not ft:
            return None
        n = len(ft)

        role_prefixes: Dict[str, Tuple[str, ...]] = {
            "answer": ("【答案】", "【参考答案】"),
            "analysis": ("【解析】", "【分析】", "【答案解析】"),
            "solution": ("【详解】", "【解答】", "【思路导引】", "【思路引导】"),
            "comment": ("【专家解读】", "【点评】", "【评注】", "【点拨】", "【点睛】"),
        }
        prefixes = role_prefixes.get(role, ())
        for pfx in prefixes:
            j = region.find(pfx)
            if j < 0:
                continue
            after = region[j + len(pfx) :]
            k = after.find(ft)
            if k >= 0:
                gap = after[:k].strip()
                if k == 0 or not gap or re.fullmatch(r"[:：\s]+", gap):
                    start = lo + j + len(pfx) + k
                    return (start, start + n)

        if role == "answer" and n <= 24:
            m = re.search(
                r"(?:【答案】|【参考答案】)\s*[:：]?\s*(" + re.escape(ft) + r")",
                region,
            )
            if m:
                start = lo + m.start(1)
                return (start, start + n)

        idx = region.find(ft)
        if idx >= 0:
            return (lo + idx, lo + idx + n)
        return None

    def _split_labeled_sections(self, text: str) -> Tuple[str, Dict[str, str]]:
        prepared_text = self._prepare_section_parsing_text(text)
        matches = list(SECTION_LABEL_PATTERN.finditer(prepared_text))
        if not matches:
            return prepared_text, {}

        body_text = self._normalize_text(prepared_text[: matches[0].start()])
        sections: Dict[str, List[str]] = {}
        for index, match in enumerate(matches):
            normalized_role = SECTION_LABEL_ALIASES.get(match.group("label"), "comment")
            content_start = match.end()
            content_end = matches[index + 1].start() if index + 1 < len(matches) else len(prepared_text)
            content = self._normalize_text(prepared_text[content_start:content_end])
            if not content:
                continue
            sections.setdefault(normalized_role, []).append(content)

        return body_text, {role: self._normalize_text("\n".join(values)) for role, values in sections.items()}

    def _prepare_section_parsing_text(self, text: str) -> str:
        prepared = self._normalize_text(text)
        # 无【】仅「标签：」换行时，短标签极易与正文撞车（如「专题：」「解析：」误切题干），仅保留「答案」类
        colon_only_short_labels = frozenset({"答案", "参考答案"})
        for label in sorted(SECTION_LABEL_ALIASES, key=len, reverse=True):
            escaped_label = re.escape(label)
            prepared = re.sub(
                rf"(?<!^)(?<!\n)\s*(【{escaped_label}】)",
                r"\n\1",
                prepared,
                flags=re.IGNORECASE,
            )
            prepared = re.sub(
                rf"(?<!^)(?<!\n)\s*(\[{escaped_label}\])",
                r"\n\1",
                prepared,
                flags=re.IGNORECASE,
            )
            if len(label) >= 3 or label in colon_only_short_labels:
                prepared = re.sub(
                    rf"(?<!^)(?<!\n)\s*({escaped_label}\s*[:：])",
                    r"\n\1",
                    prepared,
                    flags=re.IGNORECASE,
                )
        # 「答案：A 解析：…」同段时，SECTION_LABEL 无法识别行中的「解析」——强制换行
        prepared = re.sub(
            r"((?:答案|参考答案)\s*[:：]\s*[^\n]{0,120}?)\s+((?:解析|分析|详解|解答|解法|过程|点评)\s*[:：])",
            r"\1\n\2",
            prepared,
            flags=re.IGNORECASE,
        )
        return self._normalize_text(prepared)

    def _sanitize_answer_text(self, text: Optional[str]) -> Optional[str]:
        """裁掉误并入答案字段的「解析」「下一题」等尾部（常见于 Word 单行排版）。"""
        if not text or not str(text).strip():
            return text
        s = str(text).strip()
        cut = len(s)
        for pat in (
            r"(?<=[\s\u3000,，;；。])"
            r"(?:解析|分析|详解|解答|解法|过程|点评|【解析】|【分析】|【详解】|【解答】|【答案解析】)\s*[:：]?",
            r"(?<=\S)(?:解析|分析|详解|解答)\s*[:：]",  # 如「A解析：」
            r"\s+(?:第\s*\d{1,3}\s*题)",
        ):
            m = re.search(pat, s, re.IGNORECASE)
            if m and m.start() > 0 and m.start() < cut:
                cut = m.start()
        m2 = re.search(
            r"\s+\d{1,3}[\.．、](?!\d)(?=[^\n]{0,240}?(?:\(多选\)|\(单选\)|已知|若|设|求|下列))",
            s,
        )
        if m2 and m2.start() > 8 and m2.start() < cut:
            cut = m2.start()
        if cut < len(s):
            return s[:cut].strip()
        return s

    def _truncate_option_text_bleed(self, text: str) -> str:
        """选项正文后误并入答案区、解析或下一题题号时截断。"""
        if not text:
            return text
        m = _OPTION_TAIL_BLEED.search(text)
        if m and m.start() > 0:
            text = text[: m.start()].strip()
        # 同一行内「…15 答案：」无换行时
        m2 = re.search(r"(?:\s|^)(?:答案|参考答案)\s*[:：]", text, re.IGNORECASE)
        if m2 and m2.start() > 0:
            text = text[: m2.start()].strip()
        return text.strip()

    def _reflow_horizontal_options(self, text: str) -> str:
        """将「A. -4 B. -2 C. 2 D. 4」同一行横向选项拆成多行，供 OPTION_BLOCK_PATTERN 使用。"""
        if not text or not text.strip():
            return text
        out_lines: List[str] = []
        for line in text.split("\n"):
            raw = line.rstrip("\r")
            if not raw.strip():
                out_lines.append(raw)
                continue
            markers = list(
                re.finditer(
                    r"(?<![A-Za-z])[A-HＡ-Ｈ][\.．、:：]",
                    raw,
                )
            )
            if len(markers) < 2:
                out_lines.append(raw)
                continue
            for i, _m in enumerate(markers):
                start = markers[i].start()
                end = markers[i + 1].start() if i + 1 < len(markers) else len(raw)
                piece = raw[start:end].strip()
                if piece:
                    out_lines.append(piece)
        return "\n".join(out_lines)

    def _extract_options_loose(self, text: str) -> Tuple[str, List[ExtractedOption]]:
        """题干后选项为「A -4」「A) -4」等带点可有可无、键与值间有空格时的兜底。"""
        if not text:
            return "", []
        work = self._reflow_horizontal_options(text)
        work = re.sub(r"([)）\]】])\s*([A-HＡ-Ｈ])", r"\1\n\2", work)
        work = re.sub(r"(?<!^)(?<!\n)\s+([A-HＡ-Ｈ])(?=\s*[-−０-９0-9\(（])", r"\n\1", work)
        loose_block = re.compile(
            r"(?m)(?:^|\n)\s*(?P<key>[A-HＡ-Ｈ])\s*[\.．、:：）)]?\s+"
            r"(?P<content>[^\n]+?)"
            r"(?=(?:(?:^|\n)\s*[A-HＡ-Ｈ]\s*[\.．、:：）)]?\s+)|\Z)"
        )
        matches = list(loose_block.finditer(work))
        if len(matches) < 2:
            return "", []
        stem_text = self._normalize_text(work[: matches[0].start()])
        options: List[ExtractedOption] = []
        for match in matches:
            option_key = self._normalize_option_key(match.group("key"))
            option_text = self._truncate_option_text_bleed(self._normalize_text(match.group("content")))
            if not option_text:
                continue
            options.append(ExtractedOption(option_key=option_key, option_text=option_text))
        return stem_text, options

    def _extract_options(self, text: str) -> Tuple[str, List[ExtractedOption]]:
        if not text:
            return "", []

        original = text
        text = self._reflow_horizontal_options(text)
        text = re.sub(r"([)）\]】])\s*([A-HＡ-Ｈ][\.．、:：])", r"\1\n\2", text)
        prepared = re.sub(r"(?<!^)(?<!\n)\s+([A-HＡ-Ｈ][\.．、:：])", r"\n\1", text)
        matches = list(OPTION_BLOCK_PATTERN.finditer(prepared))
        options: List[ExtractedOption] = []
        if matches:
            stem_text = self._normalize_text(prepared[: matches[0].start()])
            for match in matches:
                option_key = self._normalize_option_key(match.group("key"))
                option_text = self._truncate_option_text_bleed(self._normalize_text(match.group("content")))
                if not option_text:
                    continue
                options.append(ExtractedOption(option_key=option_key, option_text=option_text))
            if options:
                return stem_text, options

        stem_loose, options_loose = self._extract_options_loose(original)
        if options_loose:
            return stem_loose, options_loose

        if matches:
            return self._normalize_text(prepared[: matches[0].start()]), []

        return self._normalize_text(original), []

    def _extract_question_formulas(
        self,
        stem_text: str,
        options: List[ExtractedOption],
        answer_text: Optional[str],
        analysis_text: Optional[str],
        solution_text: Optional[str],
    ) -> List[ExtractedFormula]:
        formulas: List[ExtractedFormula] = []
        formulas.extend(self._extract_formulas_from_text(stem_text, block_role="stem"))
        for option in options:
            formulas.extend(
                self._extract_formulas_from_text(
                    option.option_text,
                    block_role="option",
                    option_key=option.option_key,
                )
            )
        formulas.extend(self._extract_formulas_from_text(answer_text, block_role="answer"))
        formulas.extend(self._extract_formulas_from_text(analysis_text, block_role="analysis"))
        formulas.extend(self._extract_formulas_from_text(solution_text, block_role="solution"))
        return formulas

    def _extract_formulas_from_text(
        self,
        text: Optional[str],
        block_role: str,
        option_key: Optional[str] = None,
    ) -> List[ExtractedFormula]:
        normalized_text = self._normalize_text(text or "")
        if not normalized_text or not self._has_formula(normalized_text):
            return []

        candidates: List[str] = []
        candidates.extend(match.group(1) for match in LATEX_FORMULA_PATTERN.finditer(normalized_text))
        for line in normalized_text.splitlines():
            stripped_line = line.strip()
            if stripped_line and FORMULA_HINT_PATTERN.search(stripped_line):
                candidates.append(stripped_line[:200])

        formulas: List[ExtractedFormula] = []
        seen_signatures = set()
        for candidate in candidates:
            normalized_signature = self._normalize_formula_signature(candidate)
            if not normalized_signature or normalized_signature in seen_signatures:
                continue
            seen_signatures.add(normalized_signature)
            formulas.append(
                ExtractedFormula(
                    source_text=candidate.strip(),
                    normalized_signature=normalized_signature,
                    block_role=block_role,
                    option_key=option_key,
                )
            )
        return formulas

    def _extract_correct_option_keys(self, answer_text: Optional[str]) -> set:
        if not answer_text:
            return set()
        normalized_answer = self._normalize_option_key(answer_text)
        return set(re.findall(r"[A-H]", normalized_answer))

    def _normalize_option_key(self, raw_value: str) -> str:
        translation_table = str.maketrans("ＡＢＣＤＥＦＧＨａｂｃｄｅｆｇｈ", "ABCDEFGHabcdefgh")
        return raw_value.translate(translation_table).upper()

    def _infer_question_type(self, text: str, options: List[ExtractedOption]) -> str:
        if options or OPTION_PATTERN.search(text):
            return "choice"
        if "填空" in text:
            return "fill"
        return "subjective"

    def _extract_inline_metadata(self, text: str, patterns: List[re.Pattern]) -> List[str]:
        items: List[str] = []
        for pattern in patterns:
            for match in pattern.finditer(text):
                items.extend(self._split_metadata_values(match.group("value")))
        return self._unique_preserve_order(items)

    def _extract_inline_comment(self, text: str) -> Optional[str]:
        comments = self._extract_inline_metadata(text, INLINE_COMMENT_PATTERNS)
        if not comments:
            return None
        return self._normalize_text("；".join(comments))

    def _split_metadata_values(self, value: Optional[str]) -> List[str]:
        normalized_value = self._normalize_text(value or "")
        if not normalized_value:
            return []
        parts = re.split(r"[；;、,/，]\s*", normalized_value)
        return self._unique_preserve_order([part.strip(" ：:。") for part in parts if part.strip(" ：:。")])

    def _coalesce_metadata_items(self, *groups: List[str]) -> List[str]:
        merged: List[str] = []
        for group in groups:
            merged.extend(group)
        return self._unique_preserve_order(merged)

    def _build_full_question_text(
        self,
        stem_text: str,
        options: List[ExtractedOption],
        answer_text: Optional[str],
        analysis_text: Optional[str],
        solution_text: Optional[str],
        comment_text: Optional[str],
        knowledge_points: List[str],
        topics: List[str],
    ) -> str:
        parts = [self._normalize_text(stem_text)]
        if options:
            parts.extend(f"{option.option_key}. {option.option_text}" for option in options)
        if answer_text:
            parts.append(f"答案：{self._normalize_text(answer_text)}")
        if analysis_text:
            parts.append(f"解析：{self._normalize_text(analysis_text)}")
        if solution_text:
            parts.append(f"解法：{self._normalize_text(solution_text)}")
        if comment_text:
            parts.append(f"点评：{self._normalize_text(comment_text)}")
        if knowledge_points:
            parts.append(f"考点：{'；'.join(knowledge_points)}")
        if topics:
            parts.append(f"专题：{'；'.join(topics)}")
        return self._normalize_text("\n".join(part for part in parts if part))

    def _build_solution_summary(self, extracted_question: ExtractedQuestion) -> Optional[str]:
        for candidate in [
            extracted_question.solution_text,
            extracted_question.analysis_text,
            extracted_question.comment_text,
        ]:
            normalized_candidate = self._normalize_text(candidate or "")
            if normalized_candidate:
                return normalized_candidate[:1000]
        return None

    def _normalize_backend_vector_id(self, raw_value: Optional[str]) -> Optional[str]:
        candidate = str(raw_value or "").strip()
        if not candidate:
            return None
        try:
            return str(UUID(candidate))
        except (TypeError, ValueError, AttributeError):
            return str(uuid5(NAMESPACE_URL, candidate))

    def _build_vector_id(self, source_document_id: int, entity_type: str, question_no: str) -> str:
        safe_question_no = re.sub(r"[^0-9A-Za-z_-]+", "-", str(question_no)).strip("-") or "q"
        return self._normalize_backend_vector_id(f"qb-source-{source_document_id}-{entity_type}-{safe_question_no}") or str(
            uuid5(NAMESPACE_URL, f"qb-source-{source_document_id}-{entity_type}-{safe_question_no}")
        )

    def _resolve_vector_id_from_retrieval_document(
        self,
        retrieval_document: models.RetrievalDocument,
        source_document_id: int,
    ) -> Optional[str]:
        metadata = retrieval_document.metadata_json or {}
        vector_id = self._normalize_backend_vector_id(metadata.get("vector_id"))
        if vector_id:
            return vector_id

        question_no = metadata.get("question_no")
        if question_no:
            entity_type = metadata.get("entity_type") or retrieval_document.entity_type
            return self._build_vector_id(source_document_id, str(entity_type), str(question_no))
        return None


    def _hash_text(self, text: Optional[str]) -> Optional[str]:
        normalized_text = self._normalize_text(text or "")
        if not normalized_text:
            return None
        return hashlib.sha256(normalized_text.encode("utf-8")).hexdigest()

    def _sanitize_pdf_extract_noise(self, text: str) -> str:
        """弱化 PDF 内嵌数学字体映射到私用区导致的「小方块」与零宽字符。"""
        if not text:
            return ""
        s = text
        s = re.sub(r"[\u200b-\u200f\u202a-\u202e\u2060-\u206f\ufeff]", "", s)
        s = re.sub(r"[\ue000-\uf8ff]", " ", s)
        s = re.sub(r"[\uFFF0-\uFFFF]", "", s)
        s = re.sub(r"[ \t\u00a0\u1680\u2000-\u200a\u202f\u205f\u3000]+", " ", s)
        return s

    def _normalize_text(self, text: str) -> str:
        cleaned = re.sub(r"\r", "\n", text or "")
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        cleaned = self._sanitize_pdf_extract_noise(cleaned)
        return cleaned.strip()

    def _normalize_formula_signature(self, text: str) -> str:
        return re.sub(r"\s+", "", text or "").lower()[:500]

    def _has_formula(self, text: str) -> bool:
        return bool(FORMULA_HINT_PATTERN.search(text))

    def _get_image_dimensions(self, blob: bytes) -> Tuple[Optional[int], Optional[int]]:
        try:
            with Image.open(BytesIO(blob)) as image:
                return image.width, image.height
        except Exception:
            return None, None

    def _safe_int(self, value: Any) -> Optional[int]:
        try:
            parsed = int(value)
        except (TypeError, ValueError):
            return None
        return parsed if parsed >= 0 else None

    def _unique_preserve_order(self, values: Iterable[Optional[str]]) -> List[str]:

        seen = set()
        unique_values: List[str] = []
        for value in values:
            normalized_value = self._normalize_text(value or "")
            if not normalized_value or normalized_value in seen:
                continue
            seen.add(normalized_value)
            unique_values.append(normalized_value)
        return unique_values
